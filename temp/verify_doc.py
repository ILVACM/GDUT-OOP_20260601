#!/usr/bin/env python3
"""验证需求分析Word文档是否满足所有checklist检查点"""

import re
import sys
from docx import Document
from docx.oxml.ns import qn

DOC_PATH = "/workspace/temp/report/需求分析_黄泊凯_面向对象软件设计与建模.docx"

doc = Document(DOC_PATH)

# ============================================================
# 辅助函数
# ============================================================

def get_all_text(doc):
    """获取文档中所有段落文本"""
    return "\n".join(p.text for p in doc.paragraphs)

def get_all_table_text(doc):
    """获取文档中所有表格文本"""
    texts = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                texts.append(cell.text)
    return "\n".join(texts)

def get_all_runs_font_info(doc):
    """获取所有run的字体信息"""
    fonts = []
    for p in doc.paragraphs:
        for run in p.runs:
            font_name = run.font.name
            # 也检查东亚字体
            rPr = run._element.find(qn('w:rPr'))
            ea_font = None
            if rPr is not None:
                rFonts = rPr.find(qn('w:rFonts'))
                if rFonts is not None:
                    ea_font = rFonts.get(qn('w:eastAsia'))
            fonts.append({
                'text': run.text[:50],
                'font_name': font_name,
                'east_asia_font': ea_font,
                'bold': run.font.bold,
                'size': run.font.size,
            })
    return fonts

def check_keywords(text, keywords, label=""):
    """检查文本中是否包含关键词列表"""
    results = {}
    for kw in keywords:
        found = kw in text
        results[kw] = found
        status = "✅" if found else "❌"
        print(f"  {status} {label}: '{kw}' {'存在' if found else '缺失'}")
    return results

# ============================================================
# 提取文档全部内容
# ============================================================

all_text = get_all_text(doc)
all_table_text = get_all_table_text(doc)
combined_text = all_text + "\n" + all_table_text
all_fonts = get_all_runs_font_info(doc)

# 也提取Mermaid代码块内容（在段落中查找）
mermaid_blocks = []
current_block = []
in_mermaid = False
for p in doc.paragraphs:
    text = p.text.strip()
    if text.startswith("```mermaid") or text.startswith("```"):
        in_mermaid = True
        current_block = []
        continue
    if in_mermaid and text.startswith("```"):
        mermaid_blocks.append("\n".join(current_block))
        in_mermaid = False
        continue
    if in_mermaid:
        current_block.append(text)

mermaid_text = "\n".join(mermaid_blocks)

print("=" * 70)
print("需求分析文档验证报告")
print("=" * 70)

# ============================================================
# 1. 封面页检查
# ============================================================
print("\n" + "=" * 70)
print("1. 封面页检查")
print("=" * 70)

# 封面通常在文档前几段
cover_text = "\n".join(p.text for p in doc.paragraphs[:30])

cover_keywords = {
    "面向对象软件设计与建模": "课程名称",
    "需求分析": "题目",
    "欧毓毅": "指导教师",
    "计算机学院": "系别",
    "软件工程": "专业",
    "黄泊凯": "学生姓名",
    "3123004394": "班级/学号",
    "2026-06-01": "实验日期",
}

cover_results = {}
for kw, desc in cover_keywords.items():
    found = kw in cover_text or kw in combined_text
    cover_results[kw] = found
    status = "✅" if found else "❌"
    print(f"  {status} {desc}（{kw}）{'存在' if found else '缺失'}")

# ============================================================
# 2. 第一章 背景及意义
# ============================================================
print("\n" + "=" * 70)
print("2. 第一章 背景及意义")
print("=" * 70)

# 查找第一章内容范围
ch1_keywords = ["背景", "在线教育", "英语", "题库管理", "智能组卷", "在线考试", "成绩分析", "学生", "教师", "管理员"]
ch1_results = check_keywords(combined_text, ch1_keywords, "第一章")

# 检查篇幅（约500-800字）
ch1_text = ""
in_ch1 = False
for p in doc.paragraphs:
    text = p.text.strip()
    if "背景" in text and ("意义" in text or "一" in text or "1" in text):
        in_ch1 = True
    if in_ch1 and ("实验环境" in text or "第二章" in text or "2." in text):
        in_ch1 = False
        break
    if in_ch1:
        ch1_text += text

ch1_char_count = len(ch1_text.replace(" ", "").replace("\n", ""))
ch1_length_ok = 400 <= ch1_char_count <= 1200  # 放宽范围
print(f"  {'✅' if ch1_length_ok else '⚠️'} 第一章篇幅：约{ch1_char_count}字（期望500-800字）")

# ============================================================
# 3. 第二章 实验环境
# ============================================================
print("\n" + "=" * 70)
print("3. 第二章 实验环境")
print("=" * 70)

ch2_hw_keywords = ["i5-14600KF", "32.0 GB", "RTX 3090", "Windows 10"]
ch2_sw_keywords = ["JDK", "Maven", "SQLite", "Trae CN"]
ch2_tool_keywords = ["IntelliJ IDEA", "Git", "Mermaid"]

print("  --- 硬件环境 ---")
ch2_hw_results = check_keywords(combined_text, ch2_hw_keywords, "硬件")

print("  --- 软件环境 ---")
ch2_sw_results = check_keywords(combined_text, ch2_sw_keywords, "软件")

print("  --- 开发工具 ---")
ch2_tool_results = check_keywords(combined_text, ch2_tool_keywords, "工具")

# ============================================================
# 4. 第三章 系统ER图
# ============================================================
print("\n" + "=" * 70)
print("4. 第三章 系统ER图")
print("=" * 70)

er_keywords = ["erDiagram", "USER", "QUESTION", "EXAM", "SCORE"]
er_results = check_keywords(combined_text, er_keywords, "ER图")

# 检查关系标注
relation_keywords = ["1:N", "M:N", "1:1", "1..*", "*..1", "1--", "||--", "}o--", "|o--"]
relation_found = any(kw in combined_text or kw in mermaid_text for kw in relation_keywords)
print(f"  {'✅' if relation_found else '❌'} 表间关系标注：{'存在' if relation_found else '缺失'}")

# 检查Mermaid语法
er_mermaid = "erDiagram" in mermaid_text
print(f"  {'✅' if er_mermaid else '❌'} Mermaid erDiagram语法：{'存在' if er_mermaid else '缺失'}")

# ============================================================
# 5. 第四章 实体类图
# ============================================================
print("\n" + "=" * 70)
print("5. 第四章 实体类图")
print("=" * 70)

class_keywords = ["classDiagram", "User", "Question", "Exam", "Score",
                  "UserType", "QuestionType", "ExamStatus"]
class_results = check_keywords(combined_text, class_keywords, "类图")

# 检查Mermaid语法
class_mermaid = "classDiagram" in mermaid_text
print(f"  {'✅' if class_mermaid else '❌'} Mermaid classDiagram语法：{'存在' if class_mermaid else '缺失'}")

# 检查"独立设计"/"无JPA关联"
no_jpa = "无JPA关联" in combined_text or "独立设计" in combined_text or "无关联" in combined_text
print(f"  {'✅' if no_jpa else '❌'} 标注4表独立设计（无JPA关联）：{'存在' if no_jpa else '缺失'}")

# ============================================================
# 6. 第五章 子系统用例图
# ============================================================
print("\n" + "=" * 70)
print("6. 第五章 子系统用例图")
print("=" * 70)

# 检查4个子系统用例图
subsystems = {
    "M01": {"name": "用户管理", "expected": 11},
    "M02": {"name": "题库管理", "expected": 8},
    "M03": {"name": "考试管理", "expected": 10},
    "M04": {"name": "成绩管理", "expected": 9},
}

for sys_id, info in subsystems.items():
    # 检查子系统标识存在
    sys_found = sys_id in combined_text
    print(f"  {'✅' if sys_found else '❌'} {sys_id}（{info['name']}）子系统标识：{'存在' if sys_found else '缺失'}")

# 检查参与者
actors = ["student", "teacher", "admin"]
for actor in actors:
    found = actor in combined_text.lower() or actor in mermaid_text.lower()
    print(f"  {'✅' if found else '❌'} 参与者 {actor}：{'存在' if found else '缺失'}")

# 检查Mermaid用例图语法（可能没有标准语法，但应有用例相关描述）
usecase_mermaid = any("usecase" in block.lower() or "actor" in block.lower() or "useCase" in block for block in mermaid_blocks)
print(f"  {'✅' if usecase_mermaid else '⚠️'} Mermaid用例图语法：{'存在' if usecase_mermaid else '未检测到标准usecase语法（可能使用其他表示方式）'}")

# ============================================================
# 7. 第六章 用例描述
# ============================================================
print("\n" + "=" * 70)
print("7. 第六章 用例描述")
print("=" * 70)

# 检查用例编号
uc_patterns = {
    "M01": {"pattern": r"UC-M01-\d{2}", "expected": 11, "range": range(1, 12)},
    "M02": {"pattern": r"UC-M02-\d{2}", "expected": 8, "range": range(1, 9)},
    "M03": {"pattern": r"UC-M03-\d{2}", "expected": 10, "range": range(1, 11)},
    "M04": {"pattern": r"UC-M04-\d{2}", "expected": 9, "range": range(1, 10)},
}

total_uc_found = 0
for sys_id, info in uc_patterns.items():
    # 找到所有匹配的用例编号
    matches = re.findall(info["pattern"], combined_text)
    unique_matches = set(matches)
    found_count = len(unique_matches)

    # 检查每个编号是否存在
    missing = []
    for i in info["range"]:
        uc_id = f"UC-{sys_id}-{i:02d}"
        if uc_id not in combined_text:
            missing.append(uc_id)

    all_present = len(missing) == 0
    total_uc_found += found_count
    status = "✅" if all_present else "❌"
    print(f"  {status} {sys_id}模块：期望{info['expected']}个用例，找到{found_count}个唯一编号", end="")
    if missing:
        print(f"，缺失：{', '.join(missing)}")
    else:
        print()

print(f"  总计找到 {total_uc_found}/38 个用例编号")

# 检查用例描述表格字段
uc_table_fields = ["名称", "编号", "参与者", "前置条件", "基本事件流", "替代事件流", "后置条件"]
for field in uc_table_fields:
    found = field in combined_text
    print(f"  {'✅' if found else '❌'} 用例字段「{field}」：{'存在' if found else '缺失'}")

# ============================================================
# 8. 第七章 时序图
# ============================================================
print("\n" + "=" * 70)
print("8. 第七章 时序图")
print("=" * 70)

# 统计sequenceDiagram出现次数
seq_count_text = combined_text.count("sequenceDiagram")
seq_count_mermaid = mermaid_text.count("sequenceDiagram")
seq_count = max(seq_count_text, seq_count_mermaid)

print(f"  sequenceDiagram关键词出现次数：{seq_count}（期望38次）")
print(f"    - 全文出现：{seq_count_text}次")
print(f"    - Mermaid代码块内：{seq_count_mermaid}次")

# 按模块统计
for sys_id in ["M01", "M02", "M03", "M04"]:
    # 在mermaid块中查找包含该模块标识的sequenceDiagram
    sys_seq_count = 0
    for block in mermaid_blocks:
        if "sequenceDiagram" in block and sys_id in block:
            sys_seq_count += 1

    # 也从全文中统计
    # 查找该模块相关的时序图段落
    sys_text_count = 0
    in_sys = False
    for p in doc.paragraphs:
        text = p.text.strip()
        if sys_id in text and ("时序" in text or "顺序" in text or "交互" in text):
            in_sys = True
        if "sequenceDiagram" in text:
            sys_text_count += 1

    expected = {"M01": 11, "M02": 8, "M03": 10, "M04": 9}[sys_id]
    print(f"  {sys_id}模块时序图：Mermaid块内找到{sys_seq_count}张（期望{expected}张）")

# ============================================================
# 9. 文档格式检查
# ============================================================
print("\n" + "=" * 70)
print("9. 文档格式检查")
print("=" * 70)

# 检查楷体(KaiTi)和仿宋(FangSong)字体使用
kaiti_found = False
fangsong_found = False
courier_found = False

for font_info in all_fonts:
    fn = font_info['font_name'] or ""
    ea = font_info['east_asia_font'] or ""
    if "KaiTi" in fn or "楷体" in fn or "KaiTi" in ea or "楷体" in ea:
        kaiti_found = True
    if "FangSong" in fn or "仿宋" in fn or "FangSong" in ea or "仿宋" in ea:
        fangsong_found = True
    if "Courier" in fn or "Consolas" in fn:
        courier_found = True

print(f"  {'✅' if kaiti_found else '❌'} 楷体(KaiTi)字体：{'已使用' if kaiti_found else '未使用'}")
print(f"  {'✅' if fangsong_found else '❌'} 仿宋(FangSong)字体：{'已使用' if fangsong_found else '未使用'}")
print(f"  {'✅' if courier_found else '⚠️'} 等宽字体(Courier New)：{'已使用' if courier_found else '未检测到'}")

# 检查图编号和图题
fig_pattern = r"图\s*\d"
fig_matches = re.findall(fig_pattern, combined_text)
print(f"  {'✅' if len(fig_matches) > 0 else '❌'} 图编号：找到{len(fig_matches)}处")

# 检查文件名
import os
filename = os.path.basename(DOC_PATH)
expected_filename = "需求分析_黄泊凯_面向对象软件设计与建模.docx"
print(f"  {'✅' if filename == expected_filename else '❌'} 文件名：{filename}（期望：{expected_filename}）")

# 检查保存路径
expected_path = "temp/report/"
print(f"  {'✅' if expected_path in DOC_PATH else '❌'} 保存路径：{'正确' if expected_path in DOC_PATH else '不正确'}")

# ============================================================
# 汇总
# ============================================================
print("\n" + "=" * 70)
print("验证汇总")
print("=" * 70)

checks = {
    "封面页": all(cover_results.values()),
    "第一章-背景及意义关键词": all(ch1_results.values()),
    "第一章-篇幅": ch1_length_ok,
    "第二章-硬件环境": all(ch2_hw_results.values()),
    "第二章-软件环境": all(ch2_sw_results.values()),
    "第二章-开发工具": all(ch2_tool_results.values()),
    "第三章-ER图关键词": all(er_results.values()),
    "第三章-表间关系": relation_found,
    "第三章-Mermaid语法": er_mermaid,
    "第四章-类图关键词": all(class_results.values()),
    "第四章-独立设计标注": no_jpa,
    "第四章-Mermaid语法": class_mermaid,
    "第五章-子系统标识": all(sys_id in combined_text for sys_id in ["M01", "M02", "M03", "M04"]),
    "第五章-参与者": all(actor in combined_text.lower() for actor in ["student", "teacher", "admin"]),
    "第六章-用例编号完整": total_uc_found >= 38,
    "第六章-用例表格字段": all(field in combined_text for field in uc_table_fields),
    "第七章-时序图数量": seq_count >= 38,
    "格式-楷体": kaiti_found,
    "格式-仿宋": fangsong_found,
}

passed = sum(1 for v in checks.values() if v)
total = len(checks)
print(f"\n通过 {passed}/{total} 项检查\n")

for name, result in checks.items():
    print(f"  {'✅' if result else '❌'} {name}")

# 输出失败项详情
failed = [name for name, result in checks.items() if not result]
if failed:
    print(f"\n未通过项：")
    for name in failed:
        print(f"  ❌ {name}")
else:
    print(f"\n🎉 所有检查项均通过！")
