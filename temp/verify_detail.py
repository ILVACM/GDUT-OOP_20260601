#!/usr/bin/env python3
"""检查用例图和更详细的文档结构"""

import re
from docx import Document
from docx.oxml.ns import qn

DOC_PATH = "/workspace/temp/report/需求分析_黄泊凯_面向对象软件设计与建模.docx"
doc = Document(DOC_PATH)

all_text = "\n".join(p.text for p in doc.paragraphs)

# 检查用例图相关内容
print("=" * 70)
print("用例图相关段落")
print("=" * 70)

usecase_related = []
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if any(kw in text.lower() for kw in ["用例图", "usecase", "use case", "actor", "participant", "M01", "M02", "M03", "M04"]):
        if "用例" in text or "子系统" in text or "M0" in text:
            usecase_related.append((i, text[:150]))

for idx, (i, text) in enumerate(usecase_related):
    print(f"  段落#{i}: {text}")

# 检查章节标题结构
print("\n" + "=" * 70)
print("章节标题结构")
print("=" * 70)

for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if not text:
        continue
    # 检查是否是标题（通过字体大小或样式判断）
    is_heading = False
    heading_level = 0
    for run in p.runs:
        if run.font.size and run.font.size.pt:
            pt = run.font.size.pt
            if pt >= 18:  # 小二号=18pt
                is_heading = True
                heading_level = 1
            elif pt >= 15:  # 小三号=15pt
                is_heading = True
                heading_level = 2
            elif pt >= 14:  # 四号=14pt
                is_heading = True
                heading_level = 3
        if run.font.bold and run.font.size and run.font.size.pt and run.font.size.pt >= 14:
            is_heading = True

    if is_heading and len(text) < 100:
        font_sizes = []
        for run in p.runs:
            if run.font.size and run.font.size.pt:
                font_sizes.append(f"{run.font.size.pt}pt")
        print(f"  段落#{i} [{'/'.join(font_sizes)}]: {text}")

# 检查图编号
print("\n" + "=" * 70)
print("图编号统计")
print("=" * 70)
fig_pattern = r"图\s*\d+[\.-]\d*"
fig_matches = re.findall(fig_pattern, all_text)
unique_figs = sorted(set(fig_matches))
print(f"  找到{len(unique_figs)}个唯一图编号:")
for fig in unique_figs[:30]:
    print(f"    {fig}")

# 检查第五章用例图中的具体用例数量
print("\n" + "=" * 70)
print("用例图中的用例数量检查")
print("=" * 70)

# 查找用例图段落区域
in_usecase_section = False
usecase_sections = {}
current_section = None

for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if "第五章" in text or "子系统用例" in text:
        in_usecase_section = True
    if "第六章" in text or "用例描述" in text:
        in_usecase_section = False

    if in_usecase_section:
        for sys_id in ["M01", "M02", "M03", "M04"]:
            if sys_id in text and ("用例" in text or "子系统" in text):
                current_section = sys_id
                usecase_sections[sys_id] = []

        if current_section and text:
            usecase_sections.setdefault(current_section, []).append(text)

for sys_id, texts in usecase_sections.items():
    # 统计用例数量（通过(UC-xxx)或用例名称判断）
    uc_count = 0
    for t in texts:
        uc_matches = re.findall(r'\([^)]*\)', t)
        for m in uc_matches:
            if 'UC-' in m or re.match(r'\([A-Z]', m):
                uc_count += 1
    print(f"  {sys_id}: 找到{len(texts)}个相关段落，{uc_count}个用例标识")

# 检查时序图按模块分布
print("\n" + "=" * 70)
print("时序图按模块分布")
print("=" * 70)

# 通过时序图段落前后的标题/标识来判断模块归属
seq_diagrams = []
for i, p in enumerate(doc.paragraphs):
    if p.text.strip().startswith("sequenceDiagram"):
        # 向上查找最近的模块标识
        module = "未知"
        for j in range(i-1, max(0, i-20), -1):
            prev_text = doc.paragraphs[j].text.strip()
            if "M01" in prev_text:
                module = "M01"
                break
            elif "M02" in prev_text:
                module = "M02"
                break
            elif "M03" in prev_text:
                module = "M03"
                break
            elif "M04" in prev_text:
                module = "M04"
                break
        seq_diagrams.append((i, module))

module_counts = {"M01": 0, "M02": 0, "M03": 0, "M04": 0, "未知": 0}
for _, m in seq_diagrams:
    module_counts[m] = module_counts.get(m, 0) + 1

for m, c in module_counts.items():
    expected = {"M01": 11, "M02": 8, "M03": 10, "M04": 9}.get(m, "?")
    print(f"  {m}: {c}张时序图（期望{expected}张）")

# 检查字体格式详情
print("\n" + "=" * 70)
print("字体格式详情（标题段落）")
print("=" * 70)

for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if not text or len(text) > 80:
        continue
    for run in p.runs:
        if run.font.bold and run.font.size and run.font.size.pt and run.font.size.pt >= 14:
            fn = run.font.name or ""
            rPr = run._element.find(qn('w:rPr'))
            ea = ""
            if rPr is not None:
                rFonts = rPr.find(qn('w:rFonts'))
                if rFonts is not None:
                    ea = rFonts.get(qn('w:eastAsia'), "")
            print(f"  段落#{i}: [{fn}/{ea}] {run.font.size.pt}pt 加粗: {text[:60]}")
            break
