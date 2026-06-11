#!/usr/bin/env python3
"""检查用例图具体内容和正文格式"""

import re
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

DOC_PATH = "/workspace/temp/report/需求分析_黄泊凯_面向对象软件设计与建模.docx"
doc = Document(DOC_PATH)

# 检查用例图段落内容（段落44-59之间）
print("=" * 70)
print("用例图段落内容")
print("=" * 70)

for i in range(44, 60):
    if i < len(doc.paragraphs):
        text = doc.paragraphs[i].text.strip()
        if text:
            print(f"  段落#{i}: {text[:200]}")

# 检查正文段落字体
print("\n" + "=" * 70)
print("正文段落字体检查（仿宋小四号=12pt）")
print("=" * 70)

fangsong_count = 0
kaiti_count = 0
other_count = 0
size_12pt_count = 0

for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if not text or len(text) < 5:
        continue

    for run in p.runs:
        if not run.text.strip():
            continue
        fn = run.font.name or ""
        rPr = run._element.find(qn('w:rPr'))
        ea = ""
        if rPr is not None:
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is not None:
                ea = rFonts.get(qn('w:eastAsia'), "")

        font_id = ea or fn

        if "FangSong" in font_id or "仿宋" in font_id:
            fangsong_count += 1
        elif "KaiTi" in font_id or "楷体" in font_id:
            kaiti_count += 1
        else:
            other_count += 1

        if run.font.size and run.font.size.pt == 12:
            size_12pt_count += 1

print(f"  仿宋(FangSong)字体run数: {fangsong_count}")
print(f"  楷体(KaiTi)字体run数: {kaiti_count}")
print(f"  其他字体run数: {other_count}")
print(f"  12pt(小四号)字体run数: {size_12pt_count}")

# 检查行距
print("\n" + "=" * 70)
print("行距检查（1.5倍行距）")
print("=" * 70)

line_spacing_15 = 0
line_spacing_other = 0

for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if not text or len(text) < 5:
        continue
    pf = p.paragraph_format
    if pf.line_spacing is not None:
        # line_spacing可以是float(倍数)或Pt值
        if pf.line_spacing == 1.5:
            line_spacing_15 += 1
        elif isinstance(pf.line_spacing, float) and abs(pf.line_spacing - 1.5) < 0.01:
            line_spacing_15 += 1
        else:
            line_spacing_other += 1
    elif pf.line_spacing_rule is not None:
        line_spacing_other += 1

print(f"  1.5倍行距段落数: {line_spacing_15}")
print(f"  其他行距段落数: {line_spacing_other}")

# 检查用例图中的用例数量 - 通过查找用例图标题后的内容
print("\n" + "=" * 70)
print("用例图用例数量详细检查")
print("=" * 70)

# 查找每个用例图区域的内容
usecase_sections = {
    "M01": {"start": 44, "end": 48, "expected": 11},
    "M02": {"start": 48, "end": 52, "expected": 8},
    "M03": {"start": 52, "end": 56, "expected": 10},
    "M04": {"start": 56, "end": 60, "expected": 9},
}

for sys_id, info in usecase_sections.items():
    section_text = ""
    for i in range(info["start"], min(info["end"], len(doc.paragraphs))):
        section_text += doc.paragraphs[i].text + "\n"

    # 统计用例（通过括号内的用例名称）
    # 在Mermaid用例图中，用例通常用(用例名)表示
    usecase_matches = re.findall(r'\([^)]+\)', section_text)
    # 过滤掉非用例的括号内容
    usecase_names = [m for m in usecase_matches if len(m) > 3 and not m.startswith('(图')]

    print(f"  {sys_id}（段落{info['start']}-{info['end']}）：")
    print(f"    期望{info['expected']}个用例，找到{len(usecase_names)}个括号用例标识")
    for uc in usecase_names[:15]:
        print(f"      {uc}")
    if len(usecase_names) > 15:
        print(f"      ... 还有{len(usecase_names)-15}个")

# 检查三级标题
print("\n" + "=" * 70)
print("三级标题检查（楷体四号=14pt加粗）")
print("=" * 70)

for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if not text or len(text) > 80:
        continue
    for run in p.runs:
        if run.font.bold and run.font.size and run.font.size.pt == 14:
            fn = run.font.name or ""
            rPr = run._element.find(qn('w:rPr'))
            ea = ""
            if rPr is not None:
                rFonts = rPr.find(qn('w:rFonts'))
                if rFonts is not None:
                    ea = rFonts.get(qn('w:eastAsia'), "")
            print(f"  段落#{i}: [{fn}/{ea}] 14pt 加粗: {text[:60]}")
            break

# 检查用例描述表格的结构
print("\n" + "=" * 70)
print("用例描述表格结构检查")
print("=" * 70)

table_count = len(doc.tables)
print(f"  文档共{table_count}个表格")

# 检查前几个表格的字段
for t_idx, table in enumerate(doc.tables[:5]):
    if len(table.rows) > 0:
        first_row = [cell.text.strip() for cell in table.rows[0].cells]
        print(f"  表格#{t_idx} 首行: {first_row}")

# 检查所有表格是否包含用例字段
uc_field_tables = 0
for table in doc.tables:
    table_text = " ".join(cell.text for row in table.rows for cell in row.cells)
    if "用例名称" in table_text or "用例编号" in table_text or "参与者" in table_text:
        uc_field_tables += 1

print(f"  包含用例字段的表格数: {uc_field_tables}")
