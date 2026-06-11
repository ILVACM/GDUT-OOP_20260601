#!/usr/bin/env python3
"""生成「需求分析」实验报告Word文档"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ============================================================
# 辅助函数
# ============================================================

def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}>'
                          f'<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                          f'<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                          f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                          f'<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                          f'</w:tcBorders>')
    tcPr.append(tcBorders)


def set_table_borders(table):
    """设置表格所有单元格边框"""
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell)


def set_run_font(run, font_name, size_pt, bold=False):
    """设置run的字体、大小和加粗"""
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.name = font_name
    # 设置中文字体
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{font_name}"/>')
        rPr.append(rFonts)
    else:
        rFonts.set(qn('w:eastAsia'), font_name)


def add_heading1(doc, text):
    """添加一级标题：楷体，小二号(18pt)，加粗"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_run_font(run, 'KaiTi', 18, bold=True)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_heading2(doc, text):
    """添加二级标题：楷体，小三号(15pt)，加粗"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_run_font(run, 'KaiTi', 15, bold=True)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_heading3(doc, text):
    """添加三级标题：楷体，四号(14pt)，加粗"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_run_font(run, 'KaiTi', 14, bold=True)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_body_text(doc, text):
    """添加正文：仿宋，小四号(12pt)，1.5倍行距"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_run_font(run, 'FangSong', 12)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Cm(0.74)  # 两个字符缩进
    return p


def add_body_text_no_indent(doc, text):
    """添加正文（无缩进）：仿宋，小四号(12pt)，1.5倍行距"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_run_font(run, 'FangSong', 12)
    p.paragraph_format.line_spacing = 1.5
    return p


def add_mermaid_code(doc, code, figure_title):
    """添加Mermaid代码块和图题"""
    # 代码块
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(code)
    set_run_font(run, 'Courier New', 12)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    # 设置段落底纹（浅灰色背景）
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5" w:val="clear"/>')
    p.paragraph_format.element.get_or_add_pPr().append(shd)

    # 图题：居中，仿宋五号，1.5倍行距
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(figure_title)
    set_run_font(run2, 'FangSong', 10.5)
    p2.paragraph_format.line_spacing = 1.5
    p2.paragraph_format.space_before = Pt(4)
    p2.paragraph_format.space_after = Pt(8)
    return p, p2


def add_table(doc, headers, rows):
    """添加表格：仿宋五号，1.0倍行距"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(h)
        set_run_font(run, 'FangSong', 10.5, bold=True)
    # 数据行
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if c_idx > 0 else WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing = 1.0
            run = p.add_run(str(val))
            set_run_font(run, 'FangSong', 10.5)
    set_table_borders(table)
    return table


def add_use_case_table(doc, name, code, actors, precondition, basic_flow, alt_flow, postcondition):
    """添加用例描述表格（7列），1.0倍行距"""
    table = doc.add_table(rows=2, cols=7)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['用例名称', '用例编号', '参与者', '前置条件', '基本事件流', '替代事件流', '后置条件']
    values = [name, code, actors, precondition, basic_flow, alt_flow, postcondition]

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(h)
        set_run_font(run, 'FangSong', 10.5, bold=True)

    for i, v in enumerate(values):
        cell = table.rows[1].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(v)
        set_run_font(run, 'FangSong', 10.5)

    set_table_borders(table)
    # 表格后加空行
    doc.add_paragraph()
    return table


# ============================================================
# 主函数
# ============================================================

def generate_report():
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = 'FangSong'
    font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'FangSong')

    # 设置页面边距
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    # ============================================================
    # 封面页
    # ============================================================
    # 添加空行使内容居中偏下
    for _ in range(4):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

    cover_items = [
        ('广东工业大学', 26, True),
        ('面向对象软件设计与建模实验报告', 22, True),
        ('', 14, False),
        ('题目：需求分析', 18, False),
        ('', 14, False),
        ('指导教师：欧毓毅', 16, False),
        ('系别：计算机学院', 16, False),
        ('专业：软件工程', 16, False),
        ('学生姓名：黄泊凯', 16, False),
        ('班级/学号：软工3/3123004394', 16, False),
        ('实验日期：2026-06-01', 16, False),
    ]

    for text, size, bold in cover_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if text:
            run = p.add_run(text)
            set_run_font(run, 'KaiTi', size, bold=bold)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)

    # 封面后分页符
    doc.add_page_break()

    # ============================================================
    # 第一章 在线学习系统的背景及意义
    # ============================================================
    add_heading1(doc, '第一章 在线学习系统的背景及意义')

    paragraphs_ch1 = [
        '随着互联网技术的飞速发展和全球化进程的加速，英语作为国际通用语言的重要性日益凸显。传统的线下英语教学模式面临着诸多挑战：教学资源分配不均、考试组织效率低下、纸质试卷管理困难、成绩统计分析耗时耗力。在线教育平台的出现为解决这些问题提供了新的思路。',
        '本系统——英语在线学习系统（Online English Learning System），正是一个面向英语学科的在线学习与考试平台。系统以"题库管理→智能组卷→在线考试→成绩分析"为核心闭环，旨在为高校英语教学提供一站式的数字化解决方案。',
        '系统定义了三种核心角色，满足不同用户群体的需求：',
        '（1）学生（Student）：可自助注册账号，浏览可参加的考试，在线答题提交答卷，查看个人成绩和错题集，通过数据分析发现知识薄弱点。',
        '（2）教师（Teacher）：可管理题库（支持单选、多选、判断、填空、简答五种题型），通过手动或自动方式组卷，发布和管理考试，对简答题进行评卷，查看考试统计报表和题目质量分析。',
        '（3）管理员（Admin）：拥有系统全部权限，除教师功能外还可管理用户账号（创建、更新、禁用、删除），保障系统安全稳定运行。',
        '本系统的核心价值在于：通过数字化题库管理提升出题效率，通过智能组卷减少人工组卷耗时，通过自动判分降低评分错误率，通过数据统计分析辅助教学决策，最终实现英语教学与考试的全面信息化。',
    ]

    for text in paragraphs_ch1:
        add_body_text(doc, text)

    # ============================================================
    # 第二章 实验环境
    # ============================================================
    add_heading1(doc, '第二章 实验环境')

    # 2.1 硬件环境
    add_heading2(doc, '2.1 硬件环境')
    hw_headers = ['类别', '配置']
    hw_rows = [
        ['设备型号', '自定义高性能开发主机'],
        ['处理器', 'Intel(R) Core(TM) i5-14600KF (3.50 GHz)'],
        ['内存', '32.0 GB (31.8 GB 可用)'],
        ['显卡', 'NVIDIA GeForce RTX 3090 (24 GB 显存)'],
        ['存储设备', '1.82 TB HDD (WDC WD20EARZ-00C5XB0) + 932 GB SSD (Samsung SSD 990 PRO 1TB) + 8 GB 虚拟磁盘 (Msft Virtual Disk)'],
    ]
    add_table(doc, hw_headers, hw_rows)
    doc.add_paragraph()

    # 2.2 软件环境
    add_heading2(doc, '2.2 软件环境')
    sw_headers = ['类别', '名称', '版本/说明']
    sw_rows = [
        ['操作系统', 'Windows 10 IoT 企业版 LTSC', '21H2 (OS 内部版本 19044.7291)'],
        ['运行环境', 'JDK', '21'],
        ['后端框架', 'Spring Boot', '4.0.6'],
        ['ORM框架', 'Spring Data JPA + Hibernate', '—'],
        ['数据库', 'SQLite', '3.x (嵌入式)'],
        ['认证机制', 'JWT (jjwt)', '0.12.x'],
        ['构建工具', 'Maven', '3.9+'],
        ['前端框架', 'Vue 3 + Vite', '设计规划中'],
        ['AI辅助', 'Trae CN', 'Qwen3.6-Plus / GLM5.1'],
    ]
    add_table(doc, sw_headers, sw_rows)
    doc.add_paragraph()

    # 2.3 开发工具
    add_heading2(doc, '2.3 开发工具')
    dt_headers = ['工具', '用途']
    dt_rows = [
        ['IntelliJ IDEA', '后端开发IDE'],
        ['Trae CN', 'AI辅助编程'],
        ['Git + GitHub', '版本控制与代码托管'],
        ['Postman / curl', 'API接口测试'],
        ['Mermaid / PlantUML', 'UML图绘制'],
    ]
    add_table(doc, dt_headers, dt_rows)
    doc.add_paragraph()

    # ============================================================
    # 第三章 系统ER图
    # ============================================================
    add_heading1(doc, '第三章 系统ER图')

    add_body_text(doc, '本系统采用4张核心数据库表的设计方案，遵循"单表+JSON"的灵活存储策略。下图展示了系统的实体关系图，包含user（用户表）、question（题目表）、exam（考试表）和score（分数表）四张表及其关系。')

    er_code = """erDiagram
    USER {
        INTEGER id PK
        TEXT name UK
        TEXT password
        TEXT type
        INTEGER status
    }
    QUESTION {
        INTEGER id PK
        TEXT type
        TEXT context
        INTEGER img
        TEXT answer
        INTEGER use
        INTEGER correct
    }
    EXAM {
        INTEGER id PK
        TEXT exam
        TEXT status
        TEXT starttime
        TEXT endtime
        TEXT question_sum
    }
    SCORE {
        INTEGER id PK
        INTEGER user FK
        INTEGER exam FK
        INTEGER all
        TEXT detail
    }
    USER ||--o{ SCORE : "1:N 考生-成绩"
    EXAM ||--o{ SCORE : "1:N 考试-成绩"
    QUESTION }o--o{ EXAM : "M:N 快照引用" """

    add_mermaid_code(doc, er_code, '图3-1 系统ER图')

    relation_texts = [
        '（1）user → score：一对多关系，score.user为物理外键指向user.id，一个用户可有多条成绩记录。',
        '（2）exam → score：一对多关系，score.exam为物理外键指向exam.id，一个考试可有多条考生成绩。',
        '（3）exam → question：多对多关系（快照式），由exam.question_sum JSON字段表达，组卷时一次性快照，题目后续修改不影响已组卷考试。',
        '（4）user → question：隐式一对多关系，question表不存储creator_id外键，通过score表反向追溯出题教师。',
    ]
    for text in relation_texts:
        add_body_text(doc, text)

    # ============================================================
    # 第四章 实体类图
    # ============================================================
    add_heading1(doc, '第四章 实体类图')

    add_body_text(doc, '下图展示了系统4个实体类和3个枚举类的类图。本系统采用4表独立设计，无JPA @ManyToOne/@OneToMany关联关系，实体间通过ID引用而非对象引用。')

    class_code = """classDiagram
    class User {
        -Integer id
        -String name
        -String password
        -UserType type
        -Integer status
        +getId() Integer
        +getName() String
        +setName(String) void
        +getPassword() String
        +setPassword(String) void
        +getType() UserType
        +setType(UserType) void
        +getStatus() Integer
        +setStatus(Integer) void
    }
    class Question {
        -Integer id
        -QuestionType type
        -String context
        -Integer img
        -String answer
        -Integer use
        -Integer correct
        +getId() Integer
        +getType() QuestionType
        +setType(QuestionType) void
        +getContext() String
        +setContext(String) void
        +getImg() Integer
        +setImg(Integer) void
        +getAnswer() String
        +setAnswer(String) void
        +getUse() Integer
        +setUse(Integer) void
        +getCorrect() Integer
        +setCorrect(Integer) void
    }
    class Exam {
        -Integer id
        -String exam
        -ExamStatus status
        -String starttime
        -String endtime
        -String questionSum
        +getId() Integer
        +getExam() String
        +setExam(String) void
        +getStatus() ExamStatus
        +setStatus(ExamStatus) void
        +getStarttime() String
        +setStarttime(String) void
        +getEndtime() String
        +setEndtime(String) void
        +getQuestionSum() String
        +setQuestionSum(String) void
    }
    class Score {
        -Integer id
        -Integer user
        -Integer exam
        -Integer all
        -String detail
        +getId() Integer
        +getUser() Integer
        +setUser(Integer) void
        +getExam() Integer
        +setExam(Integer) void
        +getAll() Integer
        +setAll(Integer) void
        +getDetail() String
        +setDetail(String) void
    }
    class UserType {
        <<enumeration>>
        student 学生
        teacher 教师
        admin 管理员
    }
    class QuestionType {
        <<enumeration>>
        SingleChoice 单选题
        MultipleChoice 多选题
        Judge 判断题
        Fill 填空题
        Essay 简答题
    }
    class ExamStatus {
        <<enumeration>>
        draft 草稿
        publish 已发布
        running 进行中
        done 已结束
    }
    User --> UserType
    Question --> QuestionType
    Exam --> ExamStatus """

    add_mermaid_code(doc, class_code, '图4-1 系统实体类图')

    # ============================================================
    # 第五章 子系统用例图
    # ============================================================
    add_heading1(doc, '第五章 子系统用例图')

    # 5.1 M01
    add_heading3(doc, '5.1 M01 用户认证与管理子系统')
    add_body_text(doc, 'M01模块负责用户认证与权限管理，包含11个用例。参与者包括学生、教师和管理员。')

    m01_code = """graph TD
    subgraph M01用户认证与管理子系统
        UC01[UC-M01-01 用户登录]
        UC02[UC-M01-02 用户注册]
        UC03[UC-M01-03 用户登出]
        UC04[UC-M01-04 获取当前用户信息]
        UC05[UC-M01-05 修改密码]
        UC06[UC-M01-06 分页查询用户列表]
        UC07[UC-M01-07 创建用户]
        UC08[UC-M01-08 更新用户]
        UC09[UC-M01-09 更新用户状态]
        UC10[UC-M01-10 删除用户]
        UC11[UC-M01-11 批量删除用户]
    end

    Student((学生)) --> UC01
    Student --> UC02
    Student --> UC03
    Student --> UC04
    Student --> UC05

    Teacher((教师)) --> UC01
    Teacher --> UC03
    Teacher --> UC04
    Teacher --> UC05

    Admin((管理员)) --> UC01
    Admin --> UC03
    Admin --> UC04
    Admin --> UC05
    Admin --> UC06
    Admin --> UC07
    Admin --> UC08
    Admin --> UC09
    Admin --> UC10
    Admin --> UC11 """

    add_mermaid_code(doc, m01_code, '图5-1 M01用户认证与管理子系统用例图')

    # 5.2 M02
    add_heading3(doc, '5.2 M02 题库管理子系统')
    add_body_text(doc, 'M02模块负责题库管理，包含8个用例。参与者包括教师和管理员。')

    m02_code = """graph TD
    subgraph M02题库管理子系统
        UC01[UC-M02-01 创建题目]
        UC02[UC-M02-02 批量导入题目]
        UC03[UC-M02-03 查询题目详情]
        UC04[UC-M02-04 分页查询题目列表]
        UC05[UC-M02-05 更新题目]
        UC06[UC-M02-06 删除题目]
        UC07[UC-M02-07 批量删除题目]
        UC08[UC-M02-08 随机获取题目]
    end

    Teacher((教师)) --> UC01
    Teacher --> UC02
    Teacher --> UC03
    Teacher --> UC04
    Teacher --> UC05
    Teacher --> UC06
    Teacher --> UC07
    Teacher --> UC08

    Admin((管理员)) --> UC01
    Admin --> UC02
    Admin --> UC03
    Admin --> UC04
    Admin --> UC05
    Admin --> UC06
    Admin --> UC07
    Admin --> UC08 """

    add_mermaid_code(doc, m02_code, '图5-2 M02题库管理子系统用例图')

    # 5.3 M03
    add_heading3(doc, '5.3 M03 组卷与考试管理子系统')
    add_body_text(doc, 'M03模块负责组卷与考试管理，包含10个用例。参与者包括教师、管理员和学生。')

    m03_code = """graph TD
    subgraph M03组卷与考试管理子系统
        UC01[UC-M03-01 创建手动组卷考试]
        UC02[UC-M03-02 创建自动组卷考试]
        UC03[UC-M03-03 获取可参加考试列表]
        UC04[UC-M03-04 获取考试详情]
        UC05[UC-M03-05 学生预览考试]
        UC06[UC-M03-06 修改考试]
        UC07[UC-M03-07 发布考试]
        UC08[UC-M03-08 撤回考试]
        UC09[UC-M03-09 删除考试]
        UC10[UC-M03-10 分页查询考试列表]
    end

    Student((学生)) --> UC03
    Student --> UC05

    Teacher((教师)) --> UC01
    Teacher --> UC02
    Teacher --> UC04
    Teacher --> UC06
    Teacher --> UC07
    Teacher --> UC08
    Teacher --> UC09
    Teacher --> UC10

    Admin((管理员)) --> UC01
    Admin --> UC02
    Admin --> UC04
    Admin --> UC06
    Admin --> UC07
    Admin --> UC08
    Admin --> UC09
    Admin --> UC10 """

    add_mermaid_code(doc, m03_code, '图5-3 M03组卷与考试管理子系统用例图')

    # 5.4 M04
    add_heading3(doc, '5.4 M04 成绩统计子系统')
    add_body_text(doc, 'M04模块负责成绩统计与分析，包含9个用例。参与者包括学生、教师和管理员。')

    m04_code = """graph TD
    subgraph M04成绩统计子系统
        UC01[UC-M04-01 提交答卷]
        UC02[UC-M04-02 教师评卷]
        UC03[UC-M04-03 查询我的成绩]
        UC04[UC-M04-04 查询我的错题集]
        UC05[UC-M04-05 查询分数详情]
        UC06[UC-M04-06 查询考试所有考生分数]
        UC07[UC-M04-07 考试统计报表]
        UC08[UC-M04-08 题目统计列表]
        UC09[UC-M04-09 单题详细统计]
    end

    Student((学生)) --> UC01
    Student --> UC03
    Student --> UC04
    Student --> UC05

    Teacher((教师)) --> UC02
    Teacher --> UC03
    Teacher --> UC05
    Teacher --> UC06
    Teacher --> UC07
    Teacher --> UC08
    Teacher --> UC09

    Admin((管理员)) --> UC02
    Admin --> UC03
    Admin --> UC05
    Admin --> UC06
    Admin --> UC07
    Admin --> UC08
    Admin --> UC09 """

    add_mermaid_code(doc, m04_code, '图5-4 M04成绩统计子系统用例图')

    # ============================================================
    # 第六章 用例描述
    # ============================================================
    add_heading1(doc, '第六章 用例描述')

    # 6.1 M01模块用例描述
    add_heading3(doc, '6.1 M01模块用例描述')

    m01_use_cases = [
        {
            'name': '用户登录',
            'code': 'UC-M01-01',
            'actors': '学生、教师、管理员',
            'precondition': '用户已注册，账号未被禁用',
            'basic_flow': '1.用户输入用户名和密码\n2.系统查询用户记录(UserRepository.findByName)\n3.系统校验密码(BCrypt比对)\n4.系统检查账号状态(status!=0)\n5.系统签发JWT Token(JwtUtil.generateToken)\n6.返回Token和用户信息(UserVO)',
            'alt_flow': '2a.用户名不存在，提示"用户名或密码错误"\n3a.密码错误，提示"用户名或密码错误"\n4a.账号被禁用，提示"账号已被禁用"',
            'postcondition': '用户获得有效Token，可访问受保护接口',
        },
        {
            'name': '用户注册',
            'code': 'UC-M01-02',
            'actors': '学生（仅学生可自助注册）',
            'precondition': '用户名未被占用',
            'basic_flow': '1.用户输入用户名、密码、用户类型\n2.系统校验用户名唯一性(UserRepository.existsByName)\n3.系统校验仅学生可自助注册\n4.密码BCrypt加密\n5.创建用户记录(status=1)\n6.返回用户信息(UserVO)',
            'alt_flow': '2a.用户名已存在，提示"用户名已存在"\n3a.非学生类型注册，提示"仅支持学生自助注册"',
            'postcondition': '新用户记录已创建，状态为启用',
        },
        {
            'name': '用户登出',
            'code': 'UC-M01-03',
            'actors': '学生、教师、管理员',
            'precondition': '用户已登录',
            'basic_flow': '1.用户请求登出\n2.系统返回成功(JWT无状态，前端删除Token即可)',
            'alt_flow': '无',
            'postcondition': '前端清除Token，用户需重新登录',
        },
        {
            'name': '获取当前用户信息',
            'code': 'UC-M01-04',
            'actors': '学生、教师、管理员',
            'precondition': '用户已登录(携带有效Token)',
            'basic_flow': '1.系统从请求属性获取userId(JwtAuthenticationInterceptor解析)\n2.查询用户信息(UserService.getCurrentUser)\n3.转换为UserVO返回(不含密码)',
            'alt_flow': '2a.用户不存在，提示"用户不存在"',
            'postcondition': '返回当前用户的基本信息',
        },
        {
            'name': '修改密码',
            'code': 'UC-M01-05',
            'actors': '学生、教师、管理员',
            'precondition': '用户已登录，知道旧密码',
            'basic_flow': '1.用户输入旧密码和新密码\n2.系统校验旧密码(BCrypt比对)\n3.新密码BCrypt加密\n4.更新用户密码(UserRepository.save)\n5.返回成功',
            'alt_flow': '2a.旧密码错误，提示"旧密码错误"',
            'postcondition': '用户密码已更新',
        },
        {
            'name': '分页查询用户列表',
            'code': 'UC-M01-06',
            'actors': '管理员',
            'precondition': '管理员已登录',
            'basic_flow': '1.管理员输入分页参数和筛选条件(type/status)\n2.系统按条件查询用户列表(UserRepository.findAll/findByTypeAndStatus)\n3.转换为UserVO分页结果返回(PageResult)',
            'alt_flow': '无',
            'postcondition': '返回用户列表分页数据',
        },
        {
            'name': '创建用户',
            'code': 'UC-M01-07',
            'actors': '管理员',
            'precondition': '管理员已登录，用户名未被占用',
            'basic_flow': '1.管理员输入用户名、密码、用户类型\n2.系统校验用户名唯一性(UserRepository.existsByName)\n3.密码BCrypt加密\n4.创建用户记录(可创建任意角色，status=1)\n5.返回用户信息(UserVO)',
            'alt_flow': '2a.用户名已存在，提示"用户名已存在"',
            'postcondition': '新用户记录已创建',
        },
        {
            'name': '更新用户',
            'code': 'UC-M01-08',
            'actors': '管理员',
            'precondition': '管理员已登录，目标用户存在',
            'basic_flow': '1.管理员输入用户ID和更新信息\n2.系统查询目标用户(UserRepository.findById)\n3.若修改name则校验唯一性\n4.更新用户信息\n5.返回更新后的用户信息(UserVO)',
            'alt_flow': '2a.用户不存在，提示"用户不存在"\n3a.新用户名已存在，提示"用户名已存在"',
            'postcondition': '用户信息已更新',
        },
        {
            'name': '更新用户状态',
            'code': 'UC-M01-09',
            'actors': '管理员',
            'precondition': '管理员已登录，目标用户存在',
            'basic_flow': '1.管理员输入用户ID和目标状态\n2.系统查询目标用户(UserRepository.findById)\n3.若禁用admin则检查是否为最后一个启用管理员\n4.更新用户状态(UserRepository.save)\n5.返回成功',
            'alt_flow': '3a.禁用最后一个管理员，提示"不能禁用最后一个管理员"',
            'postcondition': '用户状态已更新',
        },
        {
            'name': '删除用户',
            'code': 'UC-M01-10',
            'actors': '管理员',
            'precondition': '管理员已登录，目标用户存在，不能删除自己',
            'basic_flow': '1.管理员输入用户ID\n2.系统校验不能删除自己\n3.系统校验不能删除最后一个admin\n4.若用户有考试记录(ScoreRepository.findByUser)则禁用而非删除\n5.否则硬删除用户(UserRepository.delete)\n6.返回成功',
            'alt_flow': '2a.删除自己，提示"不能删除自己"\n3a.删除最后一个管理员，提示"不能删除最后一个管理员"\n4a.有考试记录，提示"该用户存在考试记录，已禁用而非删除"',
            'postcondition': '用户被删除或禁用',
        },
        {
            'name': '批量删除用户',
            'code': 'UC-M01-11',
            'actors': '管理员',
            'precondition': '管理员已登录',
            'basic_flow': '1.管理员输入用户ID列表\n2.系统逐个调用deleteUser逻辑\n3.返回成功',
            'alt_flow': '同UC-M01-10的替代事件流',
            'postcondition': '指定用户被删除或禁用',
        },
    ]

    for uc in m01_use_cases:
        add_use_case_table(doc, uc['name'], uc['code'], uc['actors'],
                           uc['precondition'], uc['basic_flow'], uc['alt_flow'], uc['postcondition'])

    # 6.2 M02模块用例描述
    add_heading3(doc, '6.2 M02模块用例描述')

    m02_use_cases = [
        {
            'name': '创建题目',
            'code': 'UC-M02-01',
            'actors': '教师、管理员',
            'precondition': '用户已登录且具有teacher或admin角色',
            'basic_flow': '1.用户输入题目类型、题干、是否带图、答案JSON\n2.系统校验答案JSON与题型匹配\n3.创建题目记录(use=0, correct=0)\n4.返回题目信息(QuestionVO)',
            'alt_flow': '2a.答案JSON与题型不匹配，提示校验错误',
            'postcondition': '新题目记录已创建',
        },
        {
            'name': '批量导入题目',
            'code': 'UC-M02-02',
            'actors': '教师、管理员',
            'precondition': '用户已登录且具有teacher或admin角色',
            'basic_flow': '1.用户提交题目数组\n2.系统逐题校验并创建\n3.统计成功数和失败数\n4.返回批量导入结果(BatchImportResult)',
            'alt_flow': '2a.某题校验失败，跳过该题并记录错误信息',
            'postcondition': '成功的题目记录已创建',
        },
        {
            'name': '查询题目详情',
            'code': 'UC-M02-03',
            'actors': '教师、管理员',
            'precondition': '用户已登录且具有teacher或admin角色，题目存在',
            'basic_flow': '1.用户输入题目ID\n2.系统查询题目记录(QuestionRepository.findById)\n3.转换为QuestionVO返回(含答案和统计)',
            'alt_flow': '2a.题目不存在，提示"题目不存在"',
            'postcondition': '返回题目完整信息',
        },
        {
            'name': '分页查询题目列表',
            'code': 'UC-M02-04',
            'actors': '教师、管理员',
            'precondition': '用户已登录且具有teacher或admin角色',
            'basic_flow': '1.用户输入筛选条件(type/keyword)和分页参数\n2.系统按条件查询题目列表(QuestionRepository.searchByKeyword)\n3.转换为QuestionVO分页结果返回',
            'alt_flow': '无',
            'postcondition': '返回题目列表分页数据',
        },
        {
            'name': '更新题目',
            'code': 'UC-M02-05',
            'actors': '教师、管理员',
            'precondition': '用户已登录且具有teacher或admin角色，题目存在',
            'basic_flow': '1.用户输入题目ID和更新信息\n2.系统查询目标题目(QuestionRepository.findById)\n3.校验答案JSON与题型匹配\n4.更新题目信息(QuestionRepository.save)\n5.返回更新后的题目信息(QuestionVO)',
            'alt_flow': '2a.题目不存在\n3a.答案JSON与题型不匹配',
            'postcondition': '题目信息已更新',
        },
        {
            'name': '删除题目',
            'code': 'UC-M02-06',
            'actors': '教师、管理员',
            'precondition': '用户已登录且具有teacher或admin角色，题目存在',
            'basic_flow': '1.用户输入题目ID\n2.系统查询题目(QuestionRepository.findById)\n3.删除题目记录(QuestionRepository.delete)\n4.返回成功',
            'alt_flow': '2a.题目不存在',
            'postcondition': '题目记录已删除',
        },
        {
            'name': '批量删除题目',
            'code': 'UC-M02-07',
            'actors': '教师、管理员',
            'precondition': '用户已登录且具有teacher或admin角色',
            'basic_flow': '1.用户输入题目ID列表\n2.系统逐个调用删除题目逻辑\n3.返回成功',
            'alt_flow': '同UC-M02-06',
            'postcondition': '指定题目被删除',
        },
        {
            'name': '随机获取题目',
            'code': 'UC-M02-08',
            'actors': '教师、管理员',
            'precondition': '用户已登录且具有teacher或admin角色',
            'basic_flow': '1.用户输入题型过滤和排除ID列表\n2.系统按条件随机查询一道题目(QuestionRepository.findRandomByType)\n3.返回题目预览信息(QuestionPreviewVO，不含答案)',
            'alt_flow': '2a.无符合条件的题目，提示"未找到符合条件的题目"',
            'postcondition': '返回随机题目预览',
        },
    ]

    for uc in m02_use_cases:
        add_use_case_table(doc, uc['name'], uc['code'], uc['actors'],
                           uc['precondition'], uc['basic_flow'], uc['alt_flow'], uc['postcondition'])

    # 6.3 M03模块用例描述
    add_heading3(doc, '6.3 M03模块用例描述')

    m03_use_cases = [
        {
            'name': '创建手动组卷考试',
            'code': 'UC-M03-01',
            'actors': '教师、管理员',
            'precondition': '用户已登录且具有teacher或admin角色，所选题目均存在',
            'basic_flow': '1.用户输入考试名称、开始/结束时间、题目项列表[{questionId, score}]\n2.系统校验所有questionId存在(QuestionRepository.findById)\n3.构造question_sum JSON快照\n4.创建考试记录(status=draft)\n5.为每个被抽中题目执行use+=1(QuestionRepository.incrementUse)\n6.返回考试信息(ExamVO)',
            'alt_flow': '2a.某个questionId不存在，提示错误',
            'postcondition': '新考试记录已创建，状态为草稿',
        },
        {
            'name': '创建自动组卷考试',
            'code': 'UC-M03-02',
            'actors': '教师、管理员',
            'precondition': '用户已登录且具有teacher或admin角色，题库中有足够题目',
            'basic_flow': '1.用户输入考试名称、开始/结束时间、自动组卷规则{totalQuestions, totalScore, typeFilter, usePenalty}\n2.系统按规则随机抽题(QuestionRepository.findByType + 加权随机)\n3.构造question_sum JSON快照\n4.创建考试记录(status=draft)\n5.为每个被抽中题目执行use+=1\n6.返回考试信息(ExamVO)',
            'alt_flow': '2a.候选题目不足，提示"题库中符合条件的题目不足"',
            'postcondition': '新考试记录已创建，状态为草稿',
        },
        {
            'name': '获取可参加考试列表',
            'code': 'UC-M03-03',
            'actors': '学生',
            'precondition': '学生已登录',
            'basic_flow': '1.学生请求可参加的考试列表\n2.系统查询publish/running状态的考试(ExamRepository.findByStatus)\n3.实时计算考试状态(按时间窗)\n4.转换为ExamForStudentVO返回(脱敏，剔除答案)',
            'alt_flow': '无',
            'postcondition': '返回可参加的考试列表',
        },
        {
            'name': '获取考试详情',
            'code': 'UC-M03-04',
            'actors': '教师、管理员',
            'precondition': '用户已登录且具有teacher或admin角色，考试存在',
            'basic_flow': '1.用户输入考试ID\n2.系统查询考试记录(ExamRepository.findById)\n3.解析question_sum JSON\n4.转换为ExamVO返回(含完整题目信息和答案)',
            'alt_flow': '2a.考试不存在',
            'postcondition': '返回考试完整信息',
        },
        {
            'name': '学生预览考试',
            'code': 'UC-M03-05',
            'actors': '学生',
            'precondition': '学生已登录，考试存在',
            'basic_flow': '1.学生输入考试ID\n2.系统查询考试记录(ExamRepository.findById)\n3.转换为ExamForStudentVO返回(脱敏，剔除答案，仅保留选项)',
            'alt_flow': '2a.考试不存在',
            'postcondition': '返回考试预览信息(不含答案)',
        },
        {
            'name': '修改考试',
            'code': 'UC-M03-06',
            'actors': '教师、管理员',
            'precondition': '用户已登录且具有teacher或admin角色，考试处于draft状态',
            'basic_flow': '1.用户输入考试ID和更新信息\n2.系统校验考试状态为draft\n3.更新考试信息(ExamRepository.save)\n4.返回更新后的考试信息(ExamVO)',
            'alt_flow': '2a.考试非draft状态，提示"仅草稿状态可修改"',
            'postcondition': '考试信息已更新',
        },
        {
            'name': '发布考试',
            'code': 'UC-M03-07',
            'actors': '教师、管理员',
            'precondition': '用户已登录且具有teacher或admin角色，考试处于draft状态',
            'basic_flow': '1.用户输入考试ID\n2.系统校验考试状态为draft\n3.更新状态为publish(ExamRepository.save)\n4.返回成功',
            'alt_flow': '2a.考试非draft状态，提示"仅草稿状态可发布"',
            'postcondition': '考试状态变为publish',
        },
        {
            'name': '撤回考试',
            'code': 'UC-M03-08',
            'actors': '教师、管理员',
            'precondition': '用户已登录且具有teacher或admin角色，考试处于publish状态',
            'basic_flow': '1.用户输入考试ID\n2.系统校验考试状态为publish\n3.更新状态为draft(ExamRepository.save)\n4.返回成功',
            'alt_flow': '2a.考试非publish状态，提示"仅已发布状态可撤回"',
            'postcondition': '考试状态变为draft',
        },
        {
            'name': '删除考试',
            'code': 'UC-M03-09',
            'actors': '教师、管理员',
            'precondition': '用户已登录且具有teacher或admin角色，考试处于draft状态',
            'basic_flow': '1.用户输入考试ID\n2.系统校验考试状态为draft\n3.删除考试记录(ExamRepository.delete)\n4.返回成功',
            'alt_flow': '2a.考试非draft状态，提示"仅草稿状态可删除"',
            'postcondition': '考试记录已删除',
        },
        {
            'name': '分页查询考试列表',
            'code': 'UC-M03-10',
            'actors': '教师、管理员',
            'precondition': '用户已登录且具有teacher或admin角色',
            'basic_flow': '1.用户输入筛选条件(status)和分页参数\n2.系统按条件查询考试列表(ExamRepository.findByStatusNot)\n3.转换为ExamVO分页结果返回',
            'alt_flow': '无',
            'postcondition': '返回考试列表分页数据',
        },
    ]

    for uc in m03_use_cases:
        add_use_case_table(doc, uc['name'], uc['code'], uc['actors'],
                           uc['precondition'], uc['basic_flow'], uc['alt_flow'], uc['postcondition'])

    # 6.4 M04模块用例描述
    add_heading3(doc, '6.4 M04模块用例描述')

    m04_use_cases = [
        {
            'name': '提交答卷',
            'code': 'UC-M04-01',
            'actors': '学生',
            'precondition': '学生已登录，考试处于running状态，未重复提交',
            'basic_flow': '1.学生输入考试ID和答题列表[{questionId, userAnswer}]\n2.系统校验考试状态为running\n3.读取exam.question_sum确定题序与分值\n4.逐题判分(客观题自动比对question.answer，简答题暂存待评)\n5.计算总分\n6.构造score.detail JSON\n7.UPSERT score表(ScoreRepository.upsertScore)\n8.为每个isCorrect=true的题目执行correct+=1(QuestionRepository.incrementCorrect)\n9.返回分数信息(ScoreVO)',
            'alt_flow': '2a.考试非running状态\n2b.重复提交',
            'postcondition': '成绩记录已创建/更新',
        },
        {
            'name': '教师评卷',
            'code': 'UC-M04-02',
            'actors': '教师、管理员',
            'precondition': '用户已登录且具有teacher或admin角色，分数记录存在',
            'basic_flow': '1.教师输入分数ID和评卷信息[{questionId, score}]\n2.系统查询分数记录(ScoreRepository.findById)\n3.更新简答题得分和isCorrect\n4.重新计算总分和准确率\n5.若isCorrect变更则更新question.correct(QuestionRepository.incrementCorrect)\n6.返回更新后的分数信息(ScoreVO)',
            'alt_flow': '2a.分数记录不存在',
            'postcondition': '简答题评分已更新',
        },
        {
            'name': '查询我的成绩',
            'code': 'UC-M04-03',
            'actors': '学生、教师、管理员',
            'precondition': '用户已登录',
            'basic_flow': '1.用户请求个人成绩列表\n2.系统从请求属性获取userId\n3.查询该用户的所有成绩记录(ScoreRepository.findByUser)\n4.转换为ScoreListVO分页结果返回',
            'alt_flow': '无',
            'postcondition': '返回个人成绩列表',
        },
        {
            'name': '查询我的错题集',
            'code': 'UC-M04-04',
            'actors': '学生',
            'precondition': '学生已登录',
            'basic_flow': '1.学生请求错题集\n2.系统从请求属性获取userId\n3.查询该用户所有成绩记录(ScoreRepository.findByUser)\n4.从detail JSON中筛选isCorrect=false的题目\n5.转换为MistakeItemVO分页结果返回',
            'alt_flow': '无',
            'postcondition': '返回错题列表',
        },
        {
            'name': '查询分数详情',
            'code': 'UC-M04-05',
            'actors': '学生、教师、管理员',
            'precondition': '用户已登录，分数记录存在',
            'basic_flow': '1.用户输入分数ID\n2.系统查询分数记录(ScoreRepository.findById)\n3.解析detail JSON\n4.转换为ScoreVO返回(含逐题明细)',
            'alt_flow': '2a.分数记录不存在',
            'postcondition': '返回分数完整信息',
        },
        {
            'name': '查询考试所有考生分数',
            'code': 'UC-M04-06',
            'actors': '教师、管理员',
            'precondition': '用户已登录且具有teacher或admin角色',
            'basic_flow': '1.用户输入考试ID和分页参数\n2.系统查询该考试的所有成绩记录(ScoreRepository.findByExam)\n3.转换为ScoreListVO分页结果返回',
            'alt_flow': '无',
            'postcondition': '返回考试成绩列表',
        },
        {
            'name': '考试统计报表',
            'code': 'UC-M04-07',
            'actors': '教师、管理员',
            'precondition': '用户已登录且具有teacher或admin角色',
            'basic_flow': '1.用户输入考试ID\n2.系统查询该考试的所有成绩记录(ScoreRepository.findByExam)\n3.计算参与人数、平均分、最高分、最低分、通过率\n4.转换为ExamStatisticsVO返回',
            'alt_flow': '无',
            'postcondition': '返回考试统计报表',
        },
        {
            'name': '题目统计列表',
            'code': 'UC-M04-08',
            'actors': '教师、管理员',
            'precondition': '用户已登录且具有teacher或admin角色',
            'basic_flow': '1.用户输入分页参数和排序字段\n2.系统查询所有题目的统计信息(QuestionRepository.findAll)\n3.计算每题的使用次数、正确次数、正确率(correct/use)\n4.转换为QuestionStatisticsVO分页结果返回',
            'alt_flow': '无',
            'postcondition': '返回题目统计列表',
        },
        {
            'name': '单题详细统计',
            'code': 'UC-M04-09',
            'actors': '教师、管理员',
            'precondition': '用户已登录且具有teacher或admin角色，题目存在',
            'basic_flow': '1.用户输入题目ID\n2.系统查询题目统计信息(QuestionRepository.findById)\n3.计算正确率(correct/use)\n4.返回单题详细统计(QuestionStatisticsVO)',
            'alt_flow': '2a.题目不存在',
            'postcondition': '返回单题统计信息',
        },
    ]

    for uc in m04_use_cases:
        add_use_case_table(doc, uc['name'], uc['code'], uc['actors'],
                           uc['precondition'], uc['basic_flow'], uc['alt_flow'], uc['postcondition'])

    # ============================================================
    # 第七章 时序图
    # ============================================================
    add_heading1(doc, '第七章 时序图')

    # 7.1 M01模块时序图
    add_heading3(doc, '7.1 M01模块时序图')

    m01_sequences = [
        {
            'title': '图7-1 UC-M01-01 用户登录时序图',
            'code': """sequenceDiagram
    participant U as 用户
    participant C as UserController
    participant S as UserService
    participant R as UserRepository
    participant J as JwtUtil
    U->>C: POST /auth/login {name, password}
    C->>S: login(LoginReq)
    S->>R: findByName(name)
    R-->>S: User
    S->>S: BCrypt密码校验
    S->>S: 检查status!=0
    S-->>C: LoginResp(userVO)
    C->>J: generateToken(userId, name, type)
    J-->>C: token
    C-->>U: Result<LoginResp>{token, userVO} """,
        },
        {
            'title': '图7-2 UC-M01-02 用户注册时序图',
            'code': """sequenceDiagram
    participant U as 用户
    participant C as UserController
    participant S as UserService
    participant R as UserRepository
    U->>C: POST /auth/register {name, password, type}
    C->>S: register(RegisterReq)
    S->>R: existsByName(name)
    R-->>S: boolean
    S->>S: 校验仅学生可注册
    S->>S: BCrypt加密密码
    S->>R: save(User)
    R-->>S: User
    S-->>C: UserVO
    C-->>U: Result<UserVO> """,
        },
        {
            'title': '图7-3 UC-M01-03 用户登出时序图',
            'code': """sequenceDiagram
    participant U as 用户
    participant C as UserController
    U->>C: POST /auth/logout
    C-->>U: Result<Void> (JWT无状态,前端删Token) """,
        },
        {
            'title': '图7-4 UC-M01-04 获取当前用户信息时序图',
            'code': """sequenceDiagram
    participant U as 用户
    participant C as UserController
    participant S as UserService
    participant R as UserRepository
    U->>C: GET /auth/me
    C->>C: 从request获取userId
    C->>S: getCurrentUser(userId)
    S->>R: findById(userId)
    R-->>S: User
    S-->>C: UserVO
    C-->>U: Result<UserVO> """,
        },
        {
            'title': '图7-5 UC-M01-05 修改密码时序图',
            'code': """sequenceDiagram
    participant U as 用户
    participant C as UserController
    participant S as UserService
    participant R as UserRepository
    U->>C: POST /auth/password {oldPassword, newPassword}
    C->>C: 从request获取userId
    C->>S: changePassword(userId, req)
    S->>R: findById(userId)
    R-->>S: User
    S->>S: BCrypt校验旧密码
    S->>S: BCrypt加密新密码
    S->>R: save(User)
    S-->>C: void
    C-->>U: Result<Void> """,
        },
        {
            'title': '图7-6 UC-M01-06 分页查询用户列表时序图',
            'code': """sequenceDiagram
    participant A as 管理员
    participant C as UserController
    participant S as UserService
    participant R as UserRepository
    A->>C: GET /users?page=&size=&type=&status=
    C->>S: listUsers(page, size, type, status)
    S->>R: findAll/findByTypeAndStatus/findByType
    R-->>S: Page<User>
    S->>S: 转换为UserVO
    S-->>C: PageResult<UserVO>
    C-->>A: Result<PageResult<UserVO>> """,
        },
        {
            'title': '图7-7 UC-M01-07 创建用户时序图',
            'code': """sequenceDiagram
    participant A as 管理员
    participant C as UserController
    participant S as UserService
    participant R as UserRepository
    A->>C: POST /users {name, password, type}
    C->>S: createUser(RegisterReq)
    S->>R: existsByName(name)
    R-->>S: boolean
    S->>S: BCrypt加密密码
    S->>R: save(User)
    R-->>S: User
    S-->>C: UserVO
    C-->>A: Result<UserVO> """,
        },
        {
            'title': '图7-8 UC-M01-08 更新用户时序图',
            'code': """sequenceDiagram
    participant A as 管理员
    participant C as UserController
    participant S as UserService
    participant R as UserRepository
    A->>C: PUT /users/{id} {name, password, type}
    C->>S: updateUser(id, req)
    S->>R: findById(id)
    R-->>S: User
    S->>R: existsByName(name) [若name变更]
    S->>S: 更新字段
    S->>R: save(User)
    R-->>S: User
    S-->>C: UserVO
    C-->>A: Result<UserVO> """,
        },
        {
            'title': '图7-9 UC-M01-09 更新用户状态时序图',
            'code': """sequenceDiagram
    participant A as 管理员
    participant C as UserController
    participant S as UserService
    participant R as UserRepository
    A->>C: PATCH /users/{id}/status {status}
    C->>S: updateUserStatus(id, req)
    S->>R: findById(id)
    R-->>S: User
    S->>S: 检查最后一个admin保护
    S->>R: save(User)
    S-->>C: void
    C-->>A: Result<Void> """,
        },
        {
            'title': '图7-10 UC-M01-10 删除用户时序图',
            'code': """sequenceDiagram
    participant A as 管理员
    participant C as UserController
    participant S as UserService
    participant R as UserRepository
    participant SR as ScoreRepository
    A->>C: DELETE /users/{id}
    C->>C: 从request获取currentUserId
    C->>S: deleteUser(id, currentUserId)
    S->>S: 校验不能删除自己
    S->>R: findById(id)
    R-->>S: User
    S->>S: 检查最后一个admin保护
    S->>SR: findByUser(id)
    SR-->>S: List<Score>
    alt 有考试记录
        S->>R: save(User{status=0})
    else 无考试记录
        S->>R: delete(User)
    end
    S-->>C: void
    C-->>A: Result<Void> """,
        },
        {
            'title': '图7-11 UC-M01-11 批量删除用户时序图',
            'code': """sequenceDiagram
    participant A as 管理员
    participant C as UserController
    participant S as UserService
    A->>C: DELETE /users/batch {ids}
    C->>C: 从request获取currentUserId
    C->>S: batchDeleteUsers(ids, currentUserId)
    loop 逐个ID
        S->>S: deleteUser(id, currentUserId)
    end
    S-->>C: void
    C-->>A: Result<Void> """,
        },
    ]

    for seq in m01_sequences:
        add_mermaid_code(doc, seq['code'], seq['title'])

    # 7.2 M02模块时序图
    add_heading3(doc, '7.2 M02模块时序图')

    m02_sequences = [
        {
            'title': '图7-12 UC-M02-01 创建题目时序图',
            'code': """sequenceDiagram
    participant T as 教师/管理员
    participant C as QuestionController
    participant S as QuestionService
    participant R as QuestionRepository
    T->>C: POST /questions {type, context, img, answer}
    C->>S: create(QuestionCreateReq)
    S->>S: 校验答案JSON与题型匹配
    S->>R: save(Question{use=0, correct=0})
    R-->>S: Question
    S-->>C: QuestionVO
    C-->>T: Result<QuestionVO> """,
        },
        {
            'title': '图7-13 UC-M02-02 批量导入题目时序图',
            'code': """sequenceDiagram
    participant T as 教师/管理员
    participant C as QuestionController
    participant S as QuestionService
    participant R as QuestionRepository
    T->>C: POST /questions/batch [QuestionCreateReq...]
    C->>S: batchCreate(reqs)
    loop 逐题
        S->>S: 校验答案JSON
        alt 校验通过
            S->>R: save(Question)
        else 校验失败
            S->>S: 记录错误信息
        end
    end
    S-->>C: BatchImportResult
    C-->>T: Result<BatchImportResult> """,
        },
        {
            'title': '图7-14 UC-M02-03 查询题目详情时序图',
            'code': """sequenceDiagram
    participant T as 教师/管理员
    participant C as QuestionController
    participant S as QuestionService
    participant R as QuestionRepository
    T->>C: GET /questions/{id}
    C->>S: findById(id)
    S->>R: findById(id)
    R-->>S: Question
    S->>S: 转换为QuestionVO
    S-->>C: QuestionVO
    C-->>T: Result<QuestionVO> """,
        },
        {
            'title': '图7-15 UC-M02-04 分页查询题目列表时序图',
            'code': """sequenceDiagram
    participant T as 教师/管理员
    participant C as QuestionController
    participant S as QuestionService
    participant R as QuestionRepository
    T->>C: GET /questions?type=&keyword=&page=&size=
    C->>S: search(QuestionQueryReq)
    S->>R: searchByKeyword/findByType/findAll
    R-->>S: Page<Question>
    S->>S: 转换为QuestionVO
    S-->>C: Page<QuestionVO>
    C-->>T: Result<PageResult<QuestionVO>> """,
        },
        {
            'title': '图7-16 UC-M02-05 更新题目时序图',
            'code': """sequenceDiagram
    participant T as 教师/管理员
    participant C as QuestionController
    participant S as QuestionService
    participant R as QuestionRepository
    T->>C: PUT /questions/{id} {context, img, answer}
    C->>S: update(id, QuestionUpdateReq)
    S->>R: findById(id)
    R-->>S: Question
    S->>S: 校验答案JSON
    S->>R: save(Question)
    R-->>S: Question
    S-->>C: QuestionVO
    C-->>T: Result<QuestionVO> """,
        },
        {
            'title': '图7-17 UC-M02-06 删除题目时序图',
            'code': """sequenceDiagram
    participant T as 教师/管理员
    participant C as QuestionController
    participant S as QuestionService
    participant R as QuestionRepository
    T->>C: DELETE /questions/{id}
    C->>S: delete(id)
    S->>R: findById(id)
    R-->>S: Question
    S->>R: delete(Question)
    S-->>C: void
    C-->>T: Result<Void> """,
        },
        {
            'title': '图7-18 UC-M02-07 批量删除题目时序图',
            'code': """sequenceDiagram
    participant T as 教师/管理员
    participant C as QuestionController
    participant S as QuestionService
    T->>C: DELETE /questions/batch [ids]
    loop 逐个ID
        C->>S: delete(id)
    end
    C-->>T: Result<Void> """,
        },
        {
            'title': '图7-19 UC-M02-08 随机获取题目时序图',
            'code': """sequenceDiagram
    participant T as 教师/管理员
    participant C as QuestionController
    participant S as QuestionService
    participant R as QuestionRepository
    T->>C: GET /questions/random?type=&excludedIds=
    C->>S: getRandomQuestion(type, excludedIds)
    S->>R: findRandomByType(type, excludedIds)
    R-->>S: Question
    S->>S: 转换为QuestionPreviewVO
    S-->>C: QuestionPreviewVO
    C-->>T: Result<QuestionPreviewVO> """,
        },
    ]

    for seq in m02_sequences:
        add_mermaid_code(doc, seq['code'], seq['title'])

    # 7.3 M03模块时序图
    add_heading3(doc, '7.3 M03模块时序图')

    m03_sequences = [
        {
            'title': '图7-20 UC-M03-01 创建手动组卷考试时序图',
            'code': """sequenceDiagram
    participant T as 教师/管理员
    participant C as ExamController
    participant S as ExamService
    participant ER as ExamRepository
    participant QR as QuestionRepository
    T->>C: POST /exams/manual {exam, starttime, endtime, items[]}
    C->>S: createManualExam(ExamCreateManualReq)
    S->>QR: findById(questionId) [逐题]
    QR-->>S: Question
    S->>S: 构造question_sum JSON快照
    S->>QR: incrementUse(questionId) [逐题]
    S->>ER: save(Exam{status=draft})
    ER-->>S: Exam
    S-->>C: ExamVO
    C-->>T: Result<ExamVO> """,
        },
        {
            'title': '图7-21 UC-M03-02 创建自动组卷考试时序图',
            'code': """sequenceDiagram
    participant T as 教师/管理员
    participant C as ExamController
    participant S as ExamService
    participant ER as ExamRepository
    participant QR as QuestionRepository
    T->>C: POST /exams/auto {exam, starttime, endtime, rule}
    C->>S: createAutoExam(ExamCreateAutoReq)
    S->>QR: findByType(typeFilter)
    QR-->>S: List<Question>
    S->>S: 按rule随机抽题(usePenalty加权)
    S->>S: 构造question_sum JSON快照
    S->>QR: incrementUse(questionId) [逐题]
    S->>ER: save(Exam{status=draft})
    ER-->>S: Exam
    S-->>C: ExamVO
    C-->>T: Result<ExamVO> """,
        },
        {
            'title': '图7-22 UC-M03-03 获取可参加考试列表时序图',
            'code': """sequenceDiagram
    participant S as 学生
    participant C as ExamController
    participant ES as ExamService
    participant R as ExamRepository
    S->>C: GET /exams/available
    C->>ES: listAvailableExams()
    ES->>R: findByStatus(publish/running)
    R-->>ES: List<Exam>
    ES->>ES: 实时计算状态(时间窗)
    ES->>ES: 转换为ExamForStudentVO(脱敏)
    ES-->>C: List<ExamForStudentVO>
    C-->>S: Result<List<ExamForStudentVO>> """,
        },
        {
            'title': '图7-23 UC-M03-04 获取考试详情时序图',
            'code': """sequenceDiagram
    participant T as 教师/管理员
    participant C as ExamController
    participant S as ExamService
    participant R as ExamRepository
    T->>C: GET /exams/{id}
    C->>S: getExamById(id)
    S->>R: findById(id)
    R-->>S: Exam
    S->>S: 解析question_sum JSON
    S-->>C: ExamVO
    C-->>T: Result<ExamVO> """,
        },
        {
            'title': '图7-24 UC-M03-05 学生预览考试时序图',
            'code': """sequenceDiagram
    participant S as 学生
    participant C as ExamController
    participant ES as ExamService
    participant R as ExamRepository
    S->>C: GET /exams/{id}/preview
    C->>ES: getExamForStudent(id)
    ES->>R: findById(id)
    R-->>ES: Exam
    ES->>ES: 转换为ExamForStudentVO(脱敏)
    ES-->>C: ExamForStudentVO
    C-->>S: Result<ExamForStudentVO> """,
        },
        {
            'title': '图7-25 UC-M03-06 修改考试时序图',
            'code': """sequenceDiagram
    participant T as 教师/管理员
    participant C as ExamController
    participant S as ExamService
    participant R as ExamRepository
    T->>C: PUT /exams/{id} {exam, starttime, endtime, items[]}
    C->>S: updateExam(id, req)
    S->>R: findById(id)
    R-->>S: Exam
    S->>S: 校验status=draft
    S->>R: save(Exam)
    R-->>S: Exam
    S-->>C: ExamVO
    C-->>T: Result<ExamVO> """,
        },
        {
            'title': '图7-26 UC-M03-07 发布考试时序图',
            'code': """sequenceDiagram
    participant T as 教师/管理员
    participant C as ExamController
    participant S as ExamService
    participant R as ExamRepository
    T->>C: POST /exams/{id}/publish
    C->>S: publishExam(id)
    S->>R: findById(id)
    R-->>S: Exam
    S->>S: 校验status=draft
    S->>R: save(Exam{status=publish})
    S-->>C: void
    C-->>T: Result<Void> """,
        },
        {
            'title': '图7-27 UC-M03-08 撤回考试时序图',
            'code': """sequenceDiagram
    participant T as 教师/管理员
    participant C as ExamController
    participant S as ExamService
    participant R as ExamRepository
    T->>C: POST /exams/{id}/withdraw
    C->>S: withdrawExam(id)
    S->>R: findById(id)
    R-->>S: Exam
    S->>S: 校验status=publish
    S->>R: save(Exam{status=draft})
    S-->>C: void
    C-->>T: Result<Void> """,
        },
        {
            'title': '图7-28 UC-M03-09 删除考试时序图',
            'code': """sequenceDiagram
    participant T as 教师/管理员
    participant C as ExamController
    participant S as ExamService
    participant R as ExamRepository
    T->>C: DELETE /exams/{id}
    C->>S: deleteExam(id)
    S->>R: findById(id)
    R-->>S: Exam
    S->>S: 校验status=draft
    S->>R: delete(Exam)
    S-->>C: void
    C-->>T: Result<Void> """,
        },
        {
            'title': '图7-29 UC-M03-10 分页查询考试列表时序图',
            'code': """sequenceDiagram
    participant T as 教师/管理员
    participant C as ExamController
    participant S as ExamService
    participant R as ExamRepository
    T->>C: GET /exams?status=&page=&size=
    C->>S: listExams(page, size, status)
    S->>R: findByStatusNot(status, pageable)
    R-->>S: Page<Exam>
    S->>S: 转换为ExamVO
    S-->>C: PageResult<ExamVO>
    C-->>T: Result<PageResult<ExamVO>> """,
        },
    ]

    for seq in m03_sequences:
        add_mermaid_code(doc, seq['code'], seq['title'])

    # 7.4 M04模块时序图
    add_heading3(doc, '7.4 M04模块时序图')

    m04_sequences = [
        {
            'title': '图7-30 UC-M04-01 提交答卷时序图',
            'code': """sequenceDiagram
    participant S as 学生
    participant C as ScoreController
    participant SS as ScoreService
    participant SR as ScoreRepository
    participant ER as ExamRepository
    participant QR as QuestionRepository
    S->>C: POST /exams/{examId}/submit {answers[]}
    C->>C: 从request获取userId
    C->>SS: submitExam(ExamSubmitReq, userId)
    SS->>ER: findById(examId)
    ER-->>SS: Exam
    SS->>SS: 校验考试状态running
    SS->>SS: 解析question_sum确定题序分值
    loop 逐题判分
        SS->>QR: findById(questionId)
        QR-->>SS: Question
        SS->>SS: 解析answer JSON并判分
    end
    SS->>SS: 计算总分+构造detail JSON
    SS->>SR: upsertScore(userId, examId, all, detail)
    loop isCorrect=true的题目
        SS->>QR: incrementCorrect(questionId)
    end
    SS-->>C: ScoreVO
    C-->>S: Result<ScoreVO> """,
        },
        {
            'title': '图7-31 UC-M04-02 教师评卷时序图',
            'code': """sequenceDiagram
    participant T as 教师/管理员
    participant C as ScoreController
    participant S as ScoreService
    participant SR as ScoreRepository
    participant QR as QuestionRepository
    T->>C: POST /scores/{scoreId}/grade-essay {items[]}
    C->>S: gradeEssay(scoreId, EssayGradeReq)
    S->>SR: findById(scoreId)
    SR-->>S: Score
    S->>S: 解析detail JSON
    loop 逐题评卷
        S->>S: 更新简答题score和isCorrect
        alt isCorrect变为true
            S->>QR: incrementCorrect(questionId)
        end
    end
    S->>S: 重新计算总分和准确率
    S->>SR: save(Score)
    S-->>C: ScoreVO
    C-->>T: Result<ScoreVO> """,
        },
        {
            'title': '图7-32 UC-M04-03 查询我的成绩时序图',
            'code': """sequenceDiagram
    participant U as 用户
    participant C as ScoreController
    participant S as ScoreService
    participant R as ScoreRepository
    U->>C: GET /scores/me?page=&size=
    C->>C: 从request获取userId
    C->>S: getMyScores(userId, page, size)
    S->>R: findByUser(userId, pageable)
    R-->>S: Page<Score>
    S-->>C: PageResult<ScoreListVO>
    C-->>U: Result<PageResult<ScoreListVO>> """,
        },
        {
            'title': '图7-33 UC-M04-04 查询我的错题集时序图',
            'code': """sequenceDiagram
    participant S as 学生
    participant C as ScoreController
    participant SS as ScoreService
    participant SR as ScoreRepository
    S->>C: GET /scores/me/mistakes?page=&size=
    C->>C: 从request获取userId
    C->>SS: getMyMistakes(userId, page, size)
    SS->>SR: findByUser(userId)
    SR-->>SS: List<Score>
    SS->>SS: 解析detail,筛选isCorrect=false
    SS-->>C: PageResult<MistakeItemVO>
    C-->>S: Result<PageResult<MistakeItemVO>> """,
        },
        {
            'title': '图7-34 UC-M04-05 查询分数详情时序图',
            'code': """sequenceDiagram
    participant U as 用户
    participant C as ScoreController
    participant S as ScoreService
    participant R as ScoreRepository
    U->>C: GET /scores/{id}
    C->>S: findById(id)
    S->>R: findById(id)
    R-->>S: Score
    S->>S: 解析detail JSON
    S-->>C: ScoreVO
    C-->>U: Result<ScoreVO> """,
        },
        {
            'title': '图7-35 UC-M04-06 查询考试所有考生分数时序图',
            'code': """sequenceDiagram
    participant T as 教师/管理员
    participant C as ScoreController
    participant S as ScoreService
    participant R as ScoreRepository
    T->>C: GET /exams/{examId}/scores?page=&size=
    C->>S: getExamScores(examId, page, size)
    S->>R: findByExam(examId, pageable)
    R-->>S: Page<Score>
    S-->>C: PageResult<ScoreListVO>
    C-->>T: Result<PageResult<ScoreListVO>> """,
        },
        {
            'title': '图7-36 UC-M04-07 考试统计报表时序图',
            'code': """sequenceDiagram
    participant T as 教师/管理员
    participant C as ScoreController
    participant S as ScoreService
    participant SR as ScoreRepository
    participant ER as ExamRepository
    T->>C: GET /statistics/exams/{examId}
    C->>S: getExamStatistics(examId)
    S->>ER: findById(examId)
    ER-->>S: Exam
    S->>SR: findByExam(examId)
    SR-->>S: List<Score>
    S->>S: 计算参与人数/平均分/最高分/最低分/通过率
    S-->>C: ExamStatisticsVO
    C-->>T: Result<ExamStatisticsVO> """,
        },
        {
            'title': '图7-37 UC-M04-08 题目统计列表时序图',
            'code': """sequenceDiagram
    participant T as 教师/管理员
    participant C as ScoreController
    participant S as ScoreService
    participant R as QuestionRepository
    T->>C: GET /statistics/questions?page=&size=&sortBy=
    C->>S: getQuestionStatisticsPaginated(page, size, sortBy)
    S->>R: findAll(pageable)
    R-->>S: Page<Question>
    S->>S: 计算每题正确率=correct/use
    S-->>C: PageResult<QuestionStatisticsVO>
    C-->>T: Result<PageResult<QuestionStatisticsVO>> """,
        },
        {
            'title': '图7-38 UC-M04-09 单题详细统计时序图',
            'code': """sequenceDiagram
    participant T as 教师/管理员
    participant C as ScoreController
    participant S as ScoreService
    participant R as QuestionRepository
    T->>C: GET /statistics/questions/{id}
    C->>S: getQuestionStatisticById(id)
    S->>R: findById(id)
    R-->>S: Question
    S->>S: 计算正确率=correct/use
    S-->>C: QuestionStatisticsVO
    C-->>T: Result<QuestionStatisticsVO> """,
        },
    ]

    for seq in m04_sequences:
        add_mermaid_code(doc, seq['code'], seq['title'])

    # ============================================================
    # 添加页码（底部居中）
    # ============================================================
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 添加页码字段
        run = p.add_run()
        set_run_font(run, 'FangSong', 10.5)

        # 使用XML添加页码
        fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run._element.append(fldChar1)

        run2 = p.add_run()
        set_run_font(run2, 'FangSong', 10.5)
        instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run2._element.append(instrText)

        run3 = p.add_run()
        set_run_font(run3, 'FangSong', 10.5)
        fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run3._element.append(fldChar2)

    # ============================================================
    # 保存文档
    # ============================================================
    output_dir = '/workspace/temp/report'
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, '需求分析_黄泊凯_面向对象软件设计与建模.docx')
    doc.save(output_path)
    print(f'文档已生成: {output_path}')


if __name__ == '__main__':
    generate_report()
