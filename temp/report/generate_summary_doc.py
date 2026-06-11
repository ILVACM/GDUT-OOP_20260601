#!/usr/bin/env python3
"""生成「实验总结和感想」实验报告Word文档"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ============================================================
# 辅助函数（复用自 generate_design_doc.py）
# ============================================================

def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}>'
                          f'<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                          f'<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                          f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                          f'<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                          f'</w:tcBorders>')
    tcPr.append(tcBorders)


def set_table_borders(table):
    """设置表格所有单元格边框"""
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell)


def set_run_font(run, font_name, size_pt, bold=False):
    """设置run的字体、大小和加粗"""
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.name = font_name
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{font_name}"/>')
        rPr.append(rFonts)
    else:
        rFonts.set(qn('w:eastAsia'), font_name)


def add_heading1(doc, text):
    """添加一级标题：楷体，小二号(18pt)，加粗"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_run_font(run, 'KaiTi', 18, bold=True)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_heading2(doc, text):
    """添加二级标题：楷体，小三号(15pt)，加粗"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_run_font(run, 'KaiTi', 15, bold=True)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_heading3(doc, text):
    """添加三级标题：楷体，四号(14pt)，加粗"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_run_font(run, 'KaiTi', 14, bold=True)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_body_text(doc, text):
    """添加正文：仿宋，小四号(12pt)，1.5倍行距"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_run_font(run, 'FangSong', 12)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Cm(0.74)
    return p


def add_body_text_no_indent(doc, text):
    """添加正文（无缩进）：仿宋，小四号(12pt)，1.5倍行距"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_run_font(run, 'FangSong', 12)
    p.paragraph_format.line_spacing = 1.5
    return p


def add_code_block(doc, code):
    """添加代码片段：Courier New，小四号(12pt)，1.0倍行距，浅灰色背景"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(code)
    set_run_font(run, 'Courier New', 12)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5" w:val="clear"/>')
    p.paragraph_format.element.get_or_add_pPr().append(shd)
    return p


def add_mermaid_code(doc, code, figure_title):
    """添加Mermaid代码块和图题"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(code)
    set_run_font(run, 'Courier New', 12)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5" w:val="clear"/>')
    p.paragraph_format.element.get_or_add_pPr().append(shd)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(figure_title)
    set_run_font(run2, 'FangSong', 10.5)
    p2.paragraph_format.line_spacing = 1.5
    p2.paragraph_format.space_before = Pt(4)
    p2.paragraph_format.space_after = Pt(8)
    return p, p2


def add_table(doc, headers, rows):
    """添加表格：仿宋五号，1.0倍行距"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(h)
        set_run_font(run, 'FangSong', 10.5, bold=True)
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if c_idx > 0 else WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing = 1.0
            run = p.add_run(str(val))
            set_run_font(run, 'FangSong', 10.5)
    set_table_borders(table)
    return table


def add_figure_title(doc, text):
    """添加图题：仿宋，五号(10.5pt)，1.5倍行距，居中"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, 'FangSong', 10.5)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    return p


# ============================================================
# 主函数
# ============================================================

def generate_report():
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = 'FangSong'
    font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'FangSong')

    # 设置页面边距
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    # ============================================================
    # 封面页
    # ============================================================
    for _ in range(4):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

    cover_items = [
        ('广东工业大学', 26, True),
        ('面向对象软件设计与建模实验报告', 22, True),
        ('', 14, False),
        ('题目：实验总结和感想', 18, False),
        ('', 14, False),
        ('指导教师：欧毓毅', 16, False),
        ('系别：计算机学院', 16, False),
        ('专业：软件工程', 16, False),
        ('学生姓名：黄泊凯', 16, False),
        ('班级/学号：软工3/3123004394', 16, False),
        ('实验日期：2026-06-01', 16, False),
    ]

    for text, size, bold in cover_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if text:
            run = p.add_run(text)
            set_run_font(run, 'KaiTi', size, bold=bold)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)

    doc.add_page_break()

    # ============================================================
    # 第一章 系统扩展性和灵活性设计
    # ============================================================
    add_heading1(doc, '第一章 系统扩展性和灵活性设计')

    # 1.1 题型扩展设计——JSON多态答案字段
    add_heading2(doc, '1.1 题型扩展设计——JSON多态答案字段')

    add_heading3(doc, '设计问题')
    add_body_text(doc, '系统需支持5种题型（单选/多选/判断/填空/简答），且未来可能扩展更多题型。如果为每种题型建表，将导致表结构膨胀和查询复杂度急剧上升。')

    add_heading3(doc, '设计方案')
    add_body_text(doc, 'question表的answer字段采用JSON文本存储，按QuestionType枚举自适应不同的JSON结构。每种题型对应一个Answer DTO实现类（SingleChoiceAnswer、MultipleChoiceAnswer、JudgeAnswer、FillAnswer、EssayAnswer），它们共同实现Answer接口，形成多态设计。')

    add_heading3(doc, '代码体现')
    add_body_text(doc, 'Question.java（M02questionbank.entity）：answer字段声明为String类型，使用@Column(columnDefinition="TEXT")注解')
    add_body_text(doc, 'QuestionType.java（M02questionbank.enums）：定义5种枚举值')
    add_body_text(doc, 'Answer.java及5个实现类（M02questionbank.dto）：多态DTO设计')
    add_body_text(doc, 'QuestionService.create()：按type校验answer JSON结构')
    add_body_text(doc, 'ScoreService.gradeOne()：按type执行不同判分逻辑（switch表达式）')

    add_heading3(doc, '扩展方式')
    add_body_text(doc, '新增题型只需三步：\u2460在QuestionType枚举添加值 \u2461创建对应Answer DTO实现类 \u2462在QuestionService和ScoreService的switch表达式中添加对应分支——无需修改数据库表结构。')

    add_heading3(doc, 'answer JSON的5种结构示例')
    add_code_block(doc, '单选题：{"version":1, "correctOption":"A", "options":["A","B","C","D"]}')
    add_code_block(doc, '多选题：{"version":1, "correctOptions":["A","C"], "options":["A","B","C","D"]}')
    add_code_block(doc, '判断题：{"version":1, "correct":true}')
    add_code_block(doc, '填空题：{"version":1, "blanks":["went","goes"]}')
    add_code_block(doc, '简答题：{"version":1, "reference":"参考答案"}')

    # 1.2 组卷快照设计——考试与题目的解耦
    add_heading2(doc, '1.2 组卷快照设计——考试与题目的解耦')

    add_heading3(doc, '设计问题')
    add_body_text(doc, '组卷后题目可能被修改或删除，如何保证已组卷考试的数据不受影响？传统方案使用中间表exam_question建立M:N关系，但中间表增加了数据模型复杂度，且题目修改会直接影响已组卷考试。')

    add_heading3(doc, '设计方案')
    add_body_text(doc, 'exam表的question_sum字段采用JSON文本存储，在组卷时一次性快照题目ID、类型和分值。后续题目修改不回溯，已组卷考试始终使用快照数据。')

    add_heading3(doc, '代码体现')
    add_body_text(doc, 'Exam.java（M03examassembly.entity）：questionSum字段声明为String类型')
    add_body_text(doc, 'QuestionSum.java（M03examassembly.dto）：快照DTO，含version/items/totalQuestions/totalScore')
    add_body_text(doc, 'QuestionSumItem.java（M03examassembly.dto）：快照项，含questionId/score/type')
    add_body_text(doc, 'ExamService.createManualExam()：构造快照JSON并落库')
    add_body_text(doc, 'ExamService.createAutoExam()：自动抽题后构造快照')

    add_heading3(doc, '优势')
    add_body_text(doc, '消除中间表、简化数据模型、天然支持题目版本独立演进。代价是快照中的questionId可能指向已被删除的题目，但ScoreService在判分时通过批量加载（findAllById）优雅处理了这种情况。')

    # 1.3 考试状态机设计——4态生命周期管理
    add_heading2(doc, '1.3 考试状态机设计——4态生命周期管理')

    add_heading3(doc, '设计问题')
    add_body_text(doc, '考试从创建到结束经历多个状态，需要严格的状态流转控制，防止非法操作（如删除已发布的考试、修改进行中的考试等）。')

    add_heading3(doc, '设计方案')
    add_body_text(doc, 'ExamStatus枚举定义4个状态（draft/publish/running/done），Service层实现状态流转校验，系统按时间窗自动判定状态转换。')

    add_heading3(doc, '代码体现')
    add_body_text(doc, 'ExamStatus.java（M03examassembly.enums）：定义draft/publish/running/done四种状态')
    add_body_text(doc, 'ExamService.publishExam()：校验draft→publish')
    add_body_text(doc, 'ExamService.withdrawExam()：校验publish→draft')
    add_body_text(doc, 'ExamService.deleteExam()：仅draft状态可删除')
    add_body_text(doc, 'ExamService.updateExam()：仅draft状态可修改')
    add_body_text(doc, 'ExamService.resolveCurrentStatus()：按时间窗实时判定publish→running→done')

    add_heading3(doc, 'Mermaid状态机图')

    state_code = """stateDiagram-v2
    [*] --> draft : 创建考试
    draft --> publish : 发布(publishExam)
    publish --> draft : 撤回(withdrawExam)
    publish --> running : 到达开始时间(自动)
    running --> done : 到达结束时间(自动)
    draft --> [*] : 删除(deleteExam)"""

    add_mermaid_code(doc, state_code, '图1-1 考试状态机')

    add_heading3(doc, '扩展可能')
    add_body_text(doc, '可轻松添加archived归档态（done→archived），只需扩展枚举和添加状态转换方法。')

    # 1.4 Entity-DTO完全隔离——数据安全与接口稳定
    add_heading2(doc, '1.4 Entity-DTO完全隔离——数据安全与接口稳定')

    add_heading3(doc, '设计问题')
    add_body_text(doc, 'Entity直接暴露给API层会导致三个风险：\u2460JPA懒加载引发的JSON循环引用 \u2461password等敏感字段泄露 \u2462数据库结构变更直接传导到API契约。')

    add_heading3(doc, '设计方案')
    add_body_text(doc, '所有API响应使用独立的DTO/VO Record类，Service层负责Entity↔DTO转换。本系统共定义了42个DTO/VO/Req类，覆盖4个模块的所有数据传输需求。')

    add_heading3(doc, '代码体现')
    add_body_text(doc, '各模块dto包：M01(7个)、M02(12个含Answer接口及5种实现)、M03(10个)、M04(13个)，共42个DTO/VO/Req类')
    add_body_text(doc, 'UserService.toVO()：User→UserVO（剔除password字段）')
    add_body_text(doc, 'QuestionService.toVO()：Question→QuestionVO（含accuracy计算）')
    add_body_text(doc, 'ExamService.toVO()：Exam→ExamVO（解析question_sum JSON）')
    add_body_text(doc, 'ScoreService.toScoreVO()：Score→ScoreVO（解析detail JSON）')

    add_heading3(doc, '优势')
    add_body_text(doc, '修改数据库结构不影响API契约，密码等敏感字段永远不会泄露到前端，DTO的Record定义天然不可变，保证了数据传输的安全性。')

    # 1.5 统一响应契约与全局异常处理
    add_heading2(doc, '1.5 统一响应契约与全局异常处理')

    add_heading3(doc, '设计问题')
    add_body_text(doc, 'API返回格式不统一导致前端处理复杂，异常堆栈可能泄露给客户端，不同开发者的错误处理方式不一致。')

    add_heading3(doc, '设计方案')
    add_body_text(doc, 'Result<T>统一包装所有响应，GlobalExceptionHandler统一捕获异常并转换为标准Result格式。')

    add_heading3(doc, '代码体现')
    add_body_text(doc, 'Result.java（common.api）：code/message/data三字段，提供success()/error()静态工厂方法')
    add_body_text(doc, 'PageResult.java（common.api）：分页扩展，content/totalElements/totalPages/page/size')
    add_body_text(doc, 'GlobalExceptionHandler.java（common.exception）：3个@ExceptionHandler方法')
    add_body_text(doc, 'handleBusiness()：处理BusinessException，返回业务错误码')
    add_body_text(doc, 'handleIllegalArgument()：处理IllegalArgumentException，返回参数错误码4000')
    add_body_text(doc, 'handleAny()：兜底处理Exception，返回500服务器错误')
    add_body_text(doc, 'BusinessException.java（common.exception）：code+message，业务异常基类')

    add_heading3(doc, '业务码体系')
    add_body_text(doc, '200：成功')
    add_body_text(doc, '4100-4199：认证相关（4101用户名已存在、4102密码错误、4103账号被禁用）')
    add_body_text(doc, '4200-4299：题库相关（4201题目不存在、4202答案格式错误）')
    add_body_text(doc, '4300-4399：考试相关（4301考试不存在、4302考试非draft状态）')
    add_body_text(doc, '4400-4499：分数相关（4401分数记录不存在、4402重复提交）')
    add_body_text(doc, '5000-5099：系统异常（5000 JSON解析失败）')

    # 1.6 注解式权限控制——@RequireRole
    add_heading2(doc, '1.6 注解式权限控制——@RequireRole')

    add_heading3(doc, '设计问题')
    add_body_text(doc, '不同角色（学生/教师/管理员）有不同的API访问权限，硬编码权限检查导致代码冗余且难以维护。')

    add_heading3(doc, '设计方案')
    add_body_text(doc, '自定义@RequireRole注解声明方法所需角色，JwtAuthenticationInterceptor在请求拦截阶段检查权限。')

    add_heading3(doc, '代码体现')
    add_body_text(doc, 'RequireRole.java（common.security）：@Target(METHOD) + @Retention(RUNTIME)，value为UserType[]数组')
    add_body_text(doc, 'JwtAuthenticationInterceptor.java（common.security）：preHandle()方法从HandlerMethod读取@RequireRole注解，与Token中的type比对')
    add_body_text(doc, 'WebMvcConfig.java（common.config）：注册拦截器，配置排除路径（/auth/login、/auth/register等）')
    add_body_text(doc, 'Controller方法示例：@RequireRole({UserType.teacher, UserType.admin})')

    add_heading3(doc, '扩展方式')
    add_body_text(doc, '新增角色只需在UserType枚举添加值，在Controller方法上修改@RequireRole参数即可，无需修改拦截器逻辑。')

    # 1.7 4表独立架构——消除N+1查询风险
    add_heading2(doc, '1.7 4表独立架构——消除N+1查询风险')

    add_heading3(doc, '设计问题')
    add_body_text(doc, 'JPA关联关系（@ManyToOne、@OneToMany）在懒加载场景下容易导致N+1查询性能问题，尤其在列表查询时可能触发大量额外SQL。')

    add_heading3(doc, '设计方案')
    add_body_text(doc, '4张核心表完全独立，不建立JPA关联关系。跨表查询使用findAllById批量加载，替代在循环中逐个查询。')

    add_heading3(doc, '代码体现')
    add_body_text(doc, 'Score.java（M04scorestatistics.entity）：user/exam字段使用@Column(name="user")/@Column(name="exam")而非@JoinColumn')
    add_body_text(doc, 'ScoreService.submitExam()：使用questionRepository.findAllById(questionIds)批量加载题目，避免N+1')
    add_body_text(doc, 'ScoreService.toScoreVO()：批量加载题目后构建Map<Integer, Question>，O(1)查找')

    add_heading3(doc, '优势')
    add_body_text(doc, '零N+1风险、查询性能可控、代码逻辑清晰。代价是Service层需要手动协调跨表查询，但通过批量加载+Map查找的模式，代码复杂度可控。')

    # 1.8 题内统计自维护——冗余字段设计
    add_heading2(doc, '1.8 题内统计自维护——冗余字段设计')

    add_heading3(doc, '设计问题')
    add_body_text(doc, '题目使用次数和正确率需要实时展示，但每次查询时聚合计算（COUNT+JOIN）性能差，尤其在题目数量大时。')

    add_heading3(doc, '设计方案')
    add_body_text(doc, 'question表的use/correct字段在组卷和判分时同步维护，查询时直接读取，无需聚合计算。')

    add_heading3(doc, '代码体现')
    add_body_text(doc, 'Question.java（M02questionbank.entity）：use/correct字段，DDL CHECK约束保证0≤correct≤use')
    add_body_text(doc, 'QuestionRepository.incrementUse()：@Modifying+@Query原子自增use')
    add_body_text(doc, 'QuestionRepository.incrementCorrect()：@Modifying+@Query原子自增correct')
    add_body_text(doc, 'ExamService组卷时：为每个被抽中题目调用incrementUse()')
    add_body_text(doc, 'ScoreService判分时：为每个isCorrect=true的题目调用incrementCorrect()')
    add_body_text(doc, 'ScoreService评卷时：isCorrect从非true变为true时才调用incrementCorrect()')

    add_heading3(doc, '不变量约束')
    add_body_text(doc, 'DDL定义CHECK(correct >= 0 AND correct <= use)，保证数据一致性。')

    # ============================================================
    # 第二章 AI工具使用、系统环境与实验问题
    # ============================================================
    add_heading1(doc, '第二章 AI工具使用、系统环境与实验问题')

    # 2.1 使用的AI工具概述
    add_heading2(doc, '2.1 使用的AI工具概述')

    add_body_text(doc, '本项目的开发过程深度使用了AI辅助编程工具，形成了"人作为架构师和决策者，AI作为代码执行者"的协作模式。')

    add_heading3(doc, 'Trae CN')
    add_body_text(doc, '作为主要代码生成器，承担了从Entity类定义到Service业务逻辑实现的全部代码编写工作。Trae CN基于Qwen3.7 MAX和GLM 5.1模型，能够理解Wiki文档中的设计规范并生成符合约束的代码。整个代码管线——71个Java文件、41个Controller端点、42个DTO/VO类、73个单元测试——均由Trae CN作为主要代码生成器执行，最终73个单元测试全部通过。')

    add_heading3(doc, 'Qwen3.7 MAX')
    add_body_text(doc, '在方案设计和文档生成阶段提供辅助，特别是在数据模型设计和API接口规范制定时，帮助梳理业务逻辑和边界条件。')

    add_heading3(doc, 'GLM 5.1')
    add_body_text(doc, '在问题排查和代码审查阶段提供辅助，特别是在遇到JPA类型映射问题和SQLite方言兼容性问题时，帮助定位根因和验证解决方案。')

    add_heading3(doc, 'AI辅助开发的工作流')
    add_body_text(doc, '1. 人工编写Wiki文档（数据字典、API规范、业务规则）')
    add_body_text(doc, '2. 将Wiki文档作为上下文提供给AI')
    add_body_text(doc, '3. AI根据规范生成代码')
    add_body_text(doc, '4. 人工审查代码并运行测试')
    add_body_text(doc, '5. 发现问题后更新Wiki文档，重新指导AI修正')
    add_body_text(doc, '这种"文档驱动+AI执行"的模式，确保了代码与设计规范的高度一致性。')

    # 2.2 系统环境
    add_heading2(doc, '2.2 系统环境')

    add_table(doc,
        ['类别', '名称', '版本/说明'],
        [
            ['操作系统', 'Windows 10 IoT 企业版 LTSC', '21H2 (OS 内部版本 19044.7291)'],
            ['处理器', 'Intel Core i5-14600KF', '3.50 GHz, 14核20线程'],
            ['内存', '32.0 GB', 'DDR5'],
            ['显卡', 'NVIDIA GeForce RTX 3090', '24 GB显存, 用于AI推理加速'],
            ['存储', '932GB SSD (Samsung 990 PRO)', 'SQLite数据库文件存储于此'],
            ['JDK', 'OpenJDK', '21 (使用Record/Pattern Matching等特性)'],
            ['构建工具', 'Maven', '3.9+'],
            ['数据库', 'SQLite', '3.x (xerial JDBC驱动)'],
            ['ORM', 'Spring Data JPA + Hibernate', 'Spring Boot 4.0.6内置版本'],
            ['开发IDE', 'IntelliJ IDEA + Trae CN', 'AI辅助编程'],
            ['版本控制', 'Git + GitHub', '代码托管与协作'],
        ])
    doc.add_paragraph()

    # 2.3 开发中遇到的典型问题与解决方案
    add_heading2(doc, '2.3 开发中遇到的典型问题与解决方案')

    # 问题1
    add_heading3(doc, '问题1：Boolean vs Integer类型不匹配')

    add_body_text(doc, '现象：在ddl-auto=validate模式下，Entity中Boolean类型的字段（如user.status、question.img）被Hibernate映射为TINYINT，与DDL中INTEGER类型不一致，导致启动校验失败。')

    add_body_text(doc, '原因：Hibernate社区SQLite方言将Java Boolean映射为SQL TINYINT，而SQLite实际以INTEGER存储布尔值（0/1）。在validate模式下，Hibernate检查Java类型映射与DDL列类型的一致性，TINYINT≠INTEGER导致校验失败。')

    add_body_text(doc, '解决方案：将Entity中布尔字段类型由Boolean改为Integer（0/1），并显式声明columnDefinition="INTEGER"。这样Java类型与DDL列类型完全一致，validate模式通过。')

    add_body_text(doc, '影响范围：User.status（1=启用, 0=禁用）、Question.img（1=带图, 0=不带图）')

    add_body_text(doc, '参考：wiki/02-Data-Dictionary.md \u00a710.2 问题1')

    # 问题2
    add_heading3(doc, '问题2：LocalDateTime vs String时间字段')

    add_body_text(doc, '现象：Entity中LocalDateTime类型的字段（如exam.starttime、exam.endtime）在ddl-auto=validate模式下类型推断失败。')

    add_body_text(doc, '原因：Hibernate社区SQLite方言对LocalDateTime→TEXT的自动转换在validate模式下存在类型推断差异。方言将Java LocalDateTime映射为非TEXT类型（如TIMESTAMP），而SQLite实际存储为TEXT（ISO 8601格式字符串），导致类型不匹配。')

    add_body_text(doc, '解决方案：将时间字段类型由LocalDateTime改为String，存储ISO 8601格式字符串（如"2026-06-15T09:00:00"）。在需要时间比较时，通过LocalDateTime.parse()解析字符串后比较。')

    add_body_text(doc, '影响范围：Exam.starttime、Exam.endtime')

    add_body_text(doc, '参考：wiki/02-Data-Dictionary.md \u00a710.2 问题2')

    # 问题3
    add_heading3(doc, '问题3：Score外键映射方式')

    add_body_text(doc, '现象：初始文档建议Score实体中使用@JoinColumn标注外键字段，但在4表独立设计下，外键字段类型为Integer（而非实体对象），使用@JoinColumn会导致类型不匹配。')

    add_body_text(doc, '原因：@JoinColumn通常配合@ManyToOne使用，期望字段类型为实体对象（如User而非Integer）。4表独立设计下，Score.user和Score.exam字段类型为Integer，不使用JPA关联关系。')

    add_body_text(doc, '解决方案：Score实体中外键字段使用@Column(name="user")/@Column(name="exam")显式标注。物理外键语义由DDL的FOREIGN KEY约束承载，JPA层面不建立关联关系。')

    add_body_text(doc, '影响范围：Score.user、Score.exam')

    add_body_text(doc, '参考：wiki/02-Data-Dictionary.md \u00a710.2 问题3')

    # 问题4
    add_heading3(doc, '问题4：主键类型选择')

    add_body_text(doc, '现象：初始文档使用Long作为主键类型，但SQLite自增INTEGER的范围对于课程作业规模完全足够，且DDL中显式声明columnDefinition="INTEGER"。')

    add_body_text(doc, '原因：Long在JPA中映射为BIGINT，与DDL的INTEGER类型不一致（虽然SQLite不严格区分整数类型，但validate模式会检查）。')

    add_body_text(doc, '解决方案：将所有主键/外键Java类型由Long改为Integer。SQLite的INTEGER自增范围（最大9223372036854775807）对于课程作业场景完全足够。')

    add_body_text(doc, '影响范围：所有Entity的id字段、Score的外键字段')

    add_body_text(doc, '参考：wiki/02-Data-Dictionary.md \u00a710.2 问题4')

    # 问题5
    add_heading3(doc, '问题5：SQLite JDBC自动配置失败')

    add_body_text(doc, '现象：测试环境启动时，spring-boot-starter-data-jdbc依赖触发DataJdbcRepositoriesAutoConfiguration，而SQLite不支持JDBC方言，导致测试启动失败。')

    add_body_text(doc, '原因：Spring Boot自动检测到spring-data-jdbc依赖存在时，会尝试配置JDBC Repository支持。但本系统使用的是Spring Data JPA（而非JDBC），SQLite的JDBC驱动不支持Spring Data JDBC所需的方言特性。')

    add_body_text(doc, '解决方案：在application-test.yaml中通过spring.autoconfigure.exclude排除DataJdbcRepositoriesAutoConfiguration，确保测试环境只加载JPA自动配置。')

    add_body_text(doc, '影响范围：仅测试环境（application-test.yaml）')

    add_body_text(doc, '参考：wiki/02-Data-Dictionary.md \u00a710.2 问题5')

    # 2.4 实验感想
    add_heading2(doc, '2.4 实验感想')

    add_body_text(doc, '在本次面向对象软件设计与建模的课程实验中，我深刻体会到了理论与实践结合的重要性，以及AI辅助开发对软件工程实践带来的深远影响。')

    add_heading3(doc, 'AI辅助开发的体悟')
    add_body_text(doc, '本项目的71个Java文件全部由Trae CN生成，73个单元测试全部通过。这种"人负责架构设计和决策，AI负责代码实现"的协作模式极大地提高了开发效率——从数据模型设计到全模块实现，整个过程在极短时间内完成。但我也深刻认识到，AI并不能替代人的系统设计能力。如果缺乏清晰的架构决策（如4表独立设计、JSON多态字段、快照式组卷），AI生成的代码很容易陷入"能用但不可维护"的泥潭。Wiki文档先于代码的"文档驱动开发"模式，正是确保AI输出质量的关键——只有约束足够清晰，AI才能生成符合预期的代码。')

    add_heading3(doc, '面向对象设计思想的深化理解')
    add_body_text(doc, '在课堂上学习的封装、继承、多态等概念，在本项目中得到了鲜活的体现。封装不仅是数据隐藏，更体现在Entity-DTO隔离——User实体的password字段永远不会出现在API响应中，这种"接口与实现分离"的思想正是封装的本质。多态的体现更为直观：5种题型的答案字段使用同一个JSON字段存储，通过QuestionType枚举在运行时选择不同的解析和判分策略，这比传统的5张表方案简洁得多。4表独立设计则是对"组合优于继承"原则的实践——通过ID引用而非对象关联来建立实体间关系，消除了JPA关联带来的N+1查询风险。')

    add_heading3(doc, '软件工程的规范性认识')
    add_body_text(doc, 'Wiki文档作为项目的"唯一权威信息源"，在开发过程中发挥了不可替代的作用。02-Data-Dictionary.md中定义的字段类型、约束和默认值，是Entity代码的"宪法"；01-Global-Standards.md中定义的Result<T>统一响应和全局异常处理规范，是所有Controller必须遵循的契约。这种"文档驱动开发"的模式，使得AI生成的代码能够保持高度一致性——73个单元测试全部通过，正是规范约束力的体现。')

    add_heading3(doc, '对SQLite在JPA环境下限制的深刻认识')
    add_body_text(doc, '5个典型问题中有3个与SQLite+JPA的类型映射有关（Boolean→Integer、LocalDateTime→String、Long→Integer），这让我认识到：轻量级数据库虽然部署简单，但其与ORM框架的兼容性需要格外注意。Hibernate社区SQLite方言的类型映射与DDL定义的差异，是validate模式下启动失败的根源。解决这些问题的过程，也加深了我对JPA生命周期和Hibernate类型系统的理解。')

    # ============================================================
    # 第三章 完整AI提示词附录
    # ============================================================
    add_heading1(doc, '第三章 完整AI提示词附录')

    # 3.1 提示词使用说明
    add_heading2(doc, '3.1 提示词使用说明')

    add_body_text(doc, '本项目的AI辅助开发过程采用"分阶段、分任务"的提示词组织方式。每个阶段的提示词都包含三个核心要素：\u2460当前阶段的上下文（Wiki文档内容）\u2461具体任务描述\u2462约束条件（必须遵循的规范）。以下提示词根据实际开发过程整理，保留了核心指令结构。')

    # 3.2 核心提示词列表
    add_heading2(doc, '3.2 核心提示词列表')

    # 场景1
    add_heading3(doc, '场景1：数据模型设计阶段')

    add_body_text(doc, '场景说明：指导AI理解需求并设计4张核心表的ER关系和字段定义。')

    add_body_text(doc, '提示词正文（根据实际开发过程整理）：')

    prompt1 = '"你是一个数据库架构师。请为英语在线学习系统设计数据模型，系统包含4个核心实体：User（用户）、Question（题目）、Exam（考试）、Score（分数）。要求：1) 4表独立设计，不建立JPA关联关系；2) Question的answer字段使用JSON存储，支持5种题型（单选/多选/判断/填空/简答）；3) Exam的question_sum字段使用JSON存储组卷快照；4) Score的detail字段使用JSON存储答题明细；5) 主键使用Integer类型，布尔字段使用Integer(0/1)，时间字段使用String(ISO 8601)。请输出完整的DDL脚本和JPA Entity类定义。"'
    add_code_block(doc, prompt1)

    # 场景2
    add_heading3(doc, '场景2：全局规范制定阶段')

    add_body_text(doc, '场景说明：指导AI制定统一API契约、异常处理规范和代码分层标准。')

    add_body_text(doc, '提示词正文（根据实际开发过程整理）：')

    prompt2 = '"请为Spring Boot后端项目制定全局开发规范，包括：1) 统一API响应封装Result<T>，包含code/message/data三字段；2) 分页响应封装PageResult<T>；3) 全局异常处理器GlobalExceptionHandler，处理BusinessException/MethodArgumentNotValidException/Exception三种异常；4) 业务异常类BusinessException，含code和message；5) JWT认证工具类JwtUtil；6) 自定义权限注解@RequireRole。所有Controller方法必须返回Result<T>，Entity禁止直接暴露给API层。"'
    add_code_block(doc, prompt2)

    # 场景3
    add_heading3(doc, '场景3：后端模块开发阶段')

    add_body_text(doc, '场景说明：指导AI逐模块实现Controller/Service/Repository的完整代码。')

    add_body_text(doc, '提示词正文（根据实际开发过程整理）：')

    prompt3 = '"请根据Wiki文档中的模块规范，实现M0X模块的完整代码。要求：1) Controller层：每个端点使用@RequireRole注解声明权限，返回Result<T>；2) Service层：所有public方法标注@Transactional(rollbackFor=Exception.class)，Entity→DTO转换在私有toVO()方法中完成；3) Repository层：继承JpaRepository<Entity, Integer>，自定义查询使用@Query注解；4) DTO层：使用Java 21 Record类型定义。请严格按照Wiki文档中的字段定义和API规范生成代码，不要编造文档中不存在的功能。"'
    add_code_block(doc, prompt3)

    # 场景4
    add_heading3(doc, '场景4：单元测试编写阶段')

    add_body_text(doc, '场景说明：指导AI为每个模块编写单元测试。')

    add_body_text(doc, '提示词正文（根据实际开发过程整理）：')

    prompt4 = '"请为M0X模块编写单元测试，使用JUnit 5 + Spring Boot Test。要求：1) 使用@ExtendWith(SpringExtension.class)注解；2) Mock Repository层依赖；3) 覆盖所有public方法的正常流程和异常流程；4) 验证Result<T>的code和message；5) 验证@Transactional注解的回滚行为。测试类命名规范：XxxServiceTest。"'
    add_code_block(doc, prompt4)

    # 场景5
    add_heading3(doc, '场景5：Wiki文档生成阶段')

    add_body_text(doc, '场景说明：指导AI生成技术文档。')

    add_body_text(doc, '提示词正文（根据实际开发过程整理）：')

    prompt5 = '"请为英语在线学习系统生成Wiki技术文档，包含：1) 00-INDEX.md：项目总览、模块索引、AI协作规范；2) 01-Global-Standards.md：全局API契约、异常处理、JPA规范、代码分层；3) 02-Data-Dictionary.md：4张表的完整字段定义、约束、索引、JSON结构、JPA映射规则、开发问题记录。文档必须与代码实现完全一致，所有字段名、类型、约束必须与Entity代码和DDL脚本对齐。"'
    add_code_block(doc, prompt5)

    # 场景6
    add_heading3(doc, '场景6：问题排查阶段')

    add_body_text(doc, '场景说明：遇到JPA类型映射问题时的调试提示词。')

    add_body_text(doc, '提示词正文（根据实际开发过程整理）：')

    prompt6 = '"Spring Boot应用在ddl-auto=validate模式下启动失败，错误信息：[具体错误]。项目使用SQLite数据库+Hibernate社区方言。请分析原因并给出解决方案。已知问题模式：1) Boolean→TINYINT vs INTEGER不匹配；2) LocalDateTime→非TEXT vs TEXT不匹配；3) Long→BIGINT vs INTEGER不匹配。请检查Entity字段类型与DDL列定义是否一致。"'
    add_code_block(doc, prompt6)

    # ============================================================
    # 添加页码（底部居中）
    # ============================================================
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = p.add_run()
        set_run_font(run, 'FangSong', 10.5)

        fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run._element.append(fldChar1)

        run2 = p.add_run()
        set_run_font(run2, 'FangSong', 10.5)
        instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run2._element.append(instrText)

        run3 = p.add_run()
        set_run_font(run3, 'FangSong', 10.5)
        fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run3._element.append(fldChar2)

    # ============================================================
    # 保存文档
    # ============================================================
    output_dir = '/workspace/temp/report'
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, '实验总结和感想_黄泊凯_面向对象软件设计与建模.docx')
    doc.save(output_path)
    print(f'文档已生成: {output_path}')


if __name__ == '__main__':
    generate_report()
