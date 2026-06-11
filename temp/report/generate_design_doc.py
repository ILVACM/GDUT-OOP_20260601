#!/usr/bin/env python3
"""生成「系统设计」实验报告Word文档"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ============================================================
# 辅助函数（复用自 generate_doc.py）
# ============================================================

def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}>'
                          f'<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                          f'<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                          f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                          f'<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                          f'</w:tcBorders>')
    tcPr.append(tcBorders)


def set_table_borders(table):
    """设置表格所有单元格边框"""
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell)


def set_run_font(run, font_name, size_pt, bold=False):
    """设置run的字体、大小和加粗"""
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.name = font_name
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{font_name}"/>')
        rPr.append(rFonts)
    else:
        rFonts.set(qn('w:eastAsia'), font_name)


def add_heading1(doc, text):
    """添加一级标题：楷体，小二号(18pt)，加粗"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_run_font(run, 'KaiTi', 18, bold=True)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_heading2(doc, text):
    """添加二级标题：楷体，小三号(15pt)，加粗"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_run_font(run, 'KaiTi', 15, bold=True)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_heading3(doc, text):
    """添加三级标题：楷体，四号(14pt)，加粗"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_run_font(run, 'KaiTi', 14, bold=True)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_body_text(doc, text):
    """添加正文：仿宋，小四号(12pt)，1.5倍行距"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_run_font(run, 'FangSong', 12)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Cm(0.74)
    return p


def add_body_text_no_indent(doc, text):
    """添加正文（无缩进）：仿宋，小四号(12pt)，1.5倍行距"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_run_font(run, 'FangSong', 12)
    p.paragraph_format.line_spacing = 1.5
    return p


def add_mermaid_code(doc, code, figure_title):
    """添加Mermaid代码块和图题"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(code)
    set_run_font(run, 'Courier New', 12)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5" w:val="clear"/>')
    p.paragraph_format.element.get_or_add_pPr().append(shd)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(figure_title)
    set_run_font(run2, 'FangSong', 10.5)
    p2.paragraph_format.line_spacing = 1.5
    p2.paragraph_format.space_before = Pt(4)
    p2.paragraph_format.space_after = Pt(8)
    return p, p2


def add_table(doc, headers, rows):
    """添加表格：仿宋五号，1.0倍行距"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(h)
        set_run_font(run, 'FangSong', 10.5, bold=True)
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if c_idx > 0 else WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing = 1.0
            run = p.add_run(str(val))
            set_run_font(run, 'FangSong', 10.5)
    set_table_borders(table)
    return table


# ============================================================
# 主函数
# ============================================================

def generate_report():
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = 'FangSong'
    font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'FangSong')

    # 设置页面边距
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    # ============================================================
    # 封面页
    # ============================================================
    for _ in range(4):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

    cover_items = [
        ('广东工业大学', 26, True),
        ('面向对象软件设计与建模实验报告', 22, True),
        ('', 14, False),
        ('题目：系统设计', 18, False),
        ('', 14, False),
        ('指导教师：欧毓毅', 16, False),
        ('系别：计算机学院', 16, False),
        ('专业：软件工程', 16, False),
        ('学生姓名：黄泊凯', 16, False),
        ('班级/学号：软工3/3123004394', 16, False),
        ('实验日期：2026-06-01', 16, False),
    ]

    for text, size, bold in cover_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if text:
            run = p.add_run(text)
            set_run_font(run, 'KaiTi', size, bold=bold)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)

    doc.add_page_break()

    # ============================================================
    # 第一章 子系统架构设计
    # ============================================================
    add_heading1(doc, '第一章 子系统架构设计')

    # 1.1 整体架构图
    add_heading2(doc, '1.1 整体架构图')

    add_body_text(doc, '本系统基于Windows 10 IoT企业版LTSC操作系统进行开发与部署。开发硬件依托于高性能工作站，配备Intel Core i5-14600KF处理器与NVIDIA GeForce RTX 3090显卡，确保了AI辅助编程（Trae CN）与Spring Boot后端服务的流畅运行。系统利用本地32GB内存与932GB三星990 PRO SSD提供高速的数据读写能力，SQLite数据库文件存储于高性能SSD分区中，以保障题库与成绩数据的访问效率。')

    add_body_text(doc, '系统采用经典的三层架构模式，将系统职责清晰地划分为表示层、业务逻辑层和数据访问层，并辅以Common基础设施层提供横切关注点支持。')

    add_body_text(doc, '表示层（Controller层）负责接收HTTP请求、参数校验和统一响应封装，4个Controller类分别对应4个业务模块。前端层采用Vue 3 + Element Plus技术栈（当前为设计规划中），通过RESTful API与后端通信，使用JWT Bearer Token进行身份认证。')

    add_body_text(doc, '业务逻辑层（Service层）是系统的核心，封装了所有业务规则和事务控制，包括用户认证与权限管理、题库CRUD与答案校验、组卷逻辑与状态机控制、答题判分与统计报表等。Service层严格遵循DTO隔离原则，Entity不暴露给API层。')

    add_body_text(doc, '数据访问层（Repository层）继承Spring Data JPA的JpaRepository接口，提供基础CRUD操作和自定义查询方法。4个Repository分别对应4张核心数据库表。')

    add_body_text(doc, 'Common基础设施层提供统一响应封装（Result<T>/PageResult<T>）、全局异常处理（GlobalExceptionHandler）、JWT认证与权限控制（JwtUtil/JwtAuthenticationInterceptor/@RequireRole）和Web配置（WebMvcConfig）等横切关注点。')

    add_body_text(doc, '数据层使用SQLite嵌入式数据库，包含user、question、exam、score四张核心表，采用"单表+JSON"的灵活存储策略。')

    arch_code = """graph TD
    subgraph 运行环境
        OS["Windows 10 IoT LTSC<br/>Intel i5-14600KF / 32GB RAM<br/>NVIDIA RTX 3090 / 932GB SSD"]
    end
    subgraph 前端层-设计规划中
        FE["Vue 3 + Element Plus<br/>Vue Router / Pinia / Axios"]
    end
    subgraph 后端层-Spring Boot 4.0.6
        CTRL["Controller层<br/>UserController / QuestionController<br/>ExamController / ScoreController"]
        SVC["Service层<br/>UserService / QuestionService<br/>ExamService / ScoreService"]
        REPO["Repository层<br/>UserRepository / QuestionRepository<br/>ExamRepository / ScoreRepository"]
        COMMON["Common基础设施层<br/>Result / PageResult / BusinessException<br/>GlobalExceptionHandler / JwtUtil<br/>JwtAuthenticationInterceptor / RequireRole<br/>WebMvcConfig"]
    end
    subgraph 数据层-SQLite
        DB["user / question / exam / score"]
    end
    FE -->|"RESTful API<br/>JWT Bearer Token"| CTRL
    CTRL --> SVC
    SVC --> REPO
    SVC -.-> COMMON
    CTRL -.-> COMMON
    REPO --> DB
    OS -.-> FE
    OS -.-> CTRL"""

    add_mermaid_code(doc, arch_code, '图1-1 系统整体架构图')

    # 1.2 类包图
    add_heading2(doc, '1.2 类包图')

    add_body_text(doc, '下图展示了系统的包组织结构。com.cps.backend为根包，包含common和modules两个子包。common包提供横切关注点基础设施，包含api（统一响应）、exception（异常处理）、security（认证授权）和config（Web配置）四个子包。modules包按业务领域划分为四个模块包，每个模块包内部遵循Controller-Service-Repository-Entity-DTO-Enums的标准分层结构。')

    pkg_code = """graph LR
    subgraph com.cps.backend
        subgraph common
            api["api<br/>Result, PageResult"]
            exception["exception<br/>BusinessException<br/>GlobalExceptionHandler"]
            security["security<br/>JwtUtil<br/>JwtAuthenticationInterceptor<br/>RequireRole"]
            config["config<br/>WebMvcConfig"]
        end
        subgraph modules
            M01["M01userauth<br/>controller / service / repository<br/>entity / dto / enums"]
            M02["M02questionbank<br/>controller / service / repository<br/>entity / dto / enums"]
            M03["M03examassembly<br/>controller / service / repository<br/>entity / dto / enums"]
            M04["M04scorestatistics<br/>controller / service / repository<br/>entity / dto"]
        end
    end"""

    add_mermaid_code(doc, pkg_code, '图1-2 系统类包图')

    # ============================================================
    # 第二章 各层类文件结构
    # ============================================================
    add_heading1(doc, '第二章 各层类文件结构')

    # 2.1 Controller层
    add_heading2(doc, '2.1 Controller层')

    add_table(doc,
        ['类名', '所在包', '端点数', '主要职责'],
        [
            ['UserController', 'M01userauth.controller', '11', '用户认证与权限管理（登录/注册/登出/用户CRUD）'],
            ['QuestionController', 'M02questionbank.controller', '8', '题库管理（题目CRUD/批量导入/随机获取）'],
            ['ExamController', 'M03examassembly.controller', '10', '组卷与考试管理（手动/自动组卷/发布/撤回/删除）'],
            ['ScoreController', 'M04scorestatistics.controller', '9', '成绩统计（答题提交/评卷/成绩查询/统计报表）'],
        ])
    doc.add_paragraph()

    add_body_text(doc, 'Controller层是系统的表示层，负责接收HTTP请求、参数校验和统一响应封装。本系统共有4个Controller类，提供41个RESTful API端点。每个Controller方法均使用@RequireRole注解声明所需的用户角色权限，通过JwtAuthenticationInterceptor在请求拦截阶段进行Token验证和角色校验。Controller层不包含任何业务逻辑，仅负责请求路由、参数接收和Service方法调用，所有返回值统一使用Result<T>封装。这种设计确保了表示层与业务逻辑层的严格解耦，使得API契约的变更不会影响业务逻辑的实现。')

    # 2.2 Service层
    add_heading2(doc, '2.2 Service层')

    add_table(doc,
        ['类名', '所在包', '主要方法数', '核心职责'],
        [
            ['UserService', 'M01userauth.service', '11', '用户注册/登录/密码管理/用户CRUD/批量删除'],
            ['QuestionService', 'M02questionbank.service', '10', '题目CRUD/批量导入/答案校验/统计自维护/随机获取'],
            ['ExamService', 'M03examassembly.service', '12', '手动/自动组卷/考试状态机/发布撤回/学生预览'],
            ['ScoreService', 'M04scorestatistics.service', '12', '答题判分/教师评卷/成绩查询/错题集/统计报表'],
        ])
    doc.add_paragraph()

    add_body_text(doc, 'Service层是系统的业务逻辑核心，封装了所有业务规则和事务控制。本系统共有4个Service类，提供45个public方法。Service层严格遵循DTO隔离原则：所有查询方法均通过私有toVO()方法将Entity转换为VO/DTO对象返回，Entity绝不暴露给API层。关键业务逻辑均使用@Transactional注解保证事务一致性，包括用户注册、密码修改、组卷操作、答题提交等。Service层还负责跨Repository的协调工作，例如ExamService在组卷时需要调用QuestionRepository查询题目并更新use统计，ScoreService在判分时需要同时操作ScoreRepository和QuestionRepository。')

    # 2.3 Repository层
    add_heading2(doc, '2.3 Repository层')

    add_table(doc,
        ['接口名', '所在包', '继承', '自定义查询方法'],
        [
            ['UserRepository', 'M01userauth.repository', 'JpaRepository<User, Integer>', 'findByName, existsByName, findByTypeAndStatus, findByType'],
            ['QuestionRepository', 'M02questionbank.repository', 'JpaRepository<Question, Integer>', 'searchByKeyword, incrementUse, incrementCorrect, decrementUse, findRandomQuestion'],
            ['ExamRepository', 'M03examassembly.repository', 'JpaRepository<Exam, Integer>', 'findByStatus, findByStatusNot, existsByQuestionIdInQuestionSum'],
            ['ScoreRepository', 'M04scorestatistics.repository', 'JpaRepository<Score, Integer>', 'findByUserAndExam, findByUser, findByExam'],
        ])
    doc.add_paragraph()

    add_body_text(doc, 'Repository层是系统的数据访问层，继承Spring Data JPA的JpaRepository接口，自动获得save、findById、findAll、delete等基础CRUD方法。本系统共有4个Repository接口，除继承方法外还定义了17个自定义查询方法。自定义查询方法使用@Query注解编写JPQL或原生SQL，例如QuestionRepository.incrementUse()使用@Modifying+@Query实现原子性的use字段自增，QuestionRepository.searchByKeyword()使用多条件JPQL实现按题型和关键字搜索。本系统采用4表独立设计，无JPA @ManyToOne/@OneToMany关联关系，Repository之间的协调由Service层负责。')

    # 2.4 Entity层
    add_heading2(doc, '2.4 Entity层')

    add_table(doc,
        ['类名', '类型', '对应表/用途', '字段数'],
        [
            ['User', 'Entity', 'user表', '5'],
            ['Question', 'Entity', 'question表', '7'],
            ['Exam', 'Entity', 'exam表', '6'],
            ['Score', 'Entity', 'score表', '5'],
            ['UserType', 'Enum', '用户类型(student/teacher/admin)', '3值'],
            ['QuestionType', 'Enum', '题目类型(5种题型)', '5值'],
            ['ExamStatus', 'Enum', '考试状态(4态)', '4值'],
        ])
    doc.add_paragraph()

    add_body_text(doc, 'Entity层定义了系统的4个核心实体类和3个枚举类，与数据库表结构一一对应。本系统采用4表独立设计，Entity之间无JPA @ManyToOne/@OneToMany关联关系，外键字段（如Score.user和Score.exam）使用Integer类型而非实体对象引用，物理外键语义由DDL的FOREIGN KEY约束承载。这种设计避免了JPA的N+1查询问题和循环引用风险，同时保持了数据库层面的参照完整性。Entity类使用Lombok的@Data注解简化getter/setter代码，字段类型严格对齐DDL定义（如Boolean→Integer、LocalDateTime→String的适配）。')

    # 2.5 DTO层
    add_heading2(doc, '2.5 DTO层')

    add_heading3(doc, 'M01 DTO（7个）')
    add_table(doc,
        ['类名', '类型', '用途说明'],
        [
            ['LoginReq', 'Record', '登录请求(name, password)'],
            ['LoginResp', 'Record', '登录响应(token, userVO)'],
            ['RegisterReq', 'Record', '注册请求(name, password, type)'],
            ['ChangePasswordReq', 'Record', '修改密码请求(oldPassword, newPassword)'],
            ['UserStatusReq', 'Record', '用户状态请求(status)'],
            ['BatchDeleteReq', 'Record', '批量删除请求(ids)'],
            ['UserVO', 'Record', '用户视图对象(id, name, type, status)'],
        ])
    doc.add_paragraph()

    add_heading3(doc, 'M02 DTO（10个）')
    add_table(doc,
        ['类名', '类型', '用途说明'],
        [
            ['QuestionCreateReq', 'Record', '创建题目请求(type, context, img, answer)'],
            ['QuestionUpdateReq', 'Record', '更新题目请求(context, img, answer)'],
            ['QuestionQueryReq', 'Record', '查询题目请求(type, keyword, page, size)'],
            ['QuestionVO', 'Record', '题目视图对象(含统计accuracy)'],
            ['QuestionPreviewVO', 'Record', '题目预览(不含答案)'],
            ['BatchImportResult', 'Record', '批量导入结果(successCount, failCount, errors)'],
            ['Answer', 'Interface', '答案接口(多态)'],
            ['SingleChoiceAnswer', 'Record', '单选题答案(correctOption, options)'],
            ['MultipleChoiceAnswer', 'Record', '多选题答案(correctOptions, options)'],
            ['JudgeAnswer', 'Record', '判断题答案(correct)'],
            ['FillAnswer', 'Record', '填空题答案(blanks)'],
            ['EssayAnswer', 'Record', '简答题答案(reference)'],
        ])
    doc.add_paragraph()

    add_heading3(doc, 'M03 DTO（10个）')
    add_table(doc,
        ['类名', '类型', '用途说明'],
        [
            ['ExamCreateManualReq', 'Record', '手动组卷请求(exam, starttime, endtime, items)'],
            ['ExamCreateAutoReq', 'Record', '自动组卷请求(exam, starttime, endtime, autoRule)'],
            ['AutoRule', 'Record', '自动组卷规则(typeFilter, totalQuestions, totalScore, usePenalty)'],
            ['ExamQuestionItemReq', 'Record', '考试题目项请求(questionId, score)'],
            ['ExamVO', 'Record', '考试视图对象(含题目列表和统计)'],
            ['ExamQuestionVO', 'Record', '考试题目视图(questionId, type, score)'],
            ['ExamForStudentVO', 'Record', '学生视角考试视图(脱敏,不含答案)'],
            ['ExamQuestionForStudentVO', 'Record', '学生视角题目视图(含选项,不含答案)'],
            ['QuestionSum', 'Record', '题目快照(version, items, totalQuestions, totalScore)'],
            ['QuestionSumItem', 'Record', '快照项(questionId, score, type)'],
        ])
    doc.add_paragraph()

    add_heading3(doc, 'M04 DTO（13个）')
    add_table(doc,
        ['类名', '类型', '用途说明'],
        [
            ['ExamSubmitReq', 'Record', '答题提交请求(examId, answers)'],
            ['AnswerItem', 'Record', '答题项(questionId, userAnswer)'],
            ['EssayGradeReq', 'Record', '评卷请求(questionId, score)'],
            ['ScoreVO', 'Record', '分数视图(含逐题明细)'],
            ['ScoreListVO', 'Record', '分数列表视图(不含明细)'],
            ['ScoreDetail', 'Record', '答题明细(version, items, summary)'],
            ['DetailItem', 'Record', '明细项(questionId, userAnswer, correctAnswer, score, isCorrect)'],
            ['Summary', 'Record', '答题总结(correctCount, totalCount, accuracy)'],
            ['MistakeItemVO', 'Record', '错题视图(questionId, type, context, options, userAnswer, correctAnswer, examId, examName)'],
            ['ExamStatisticsVO', 'Record', '考试统计报表(参与人数/平均分/最高分/最低分/中位数/通过率/分布)'],
            ['QuestionStatisticsVO', 'class', '题目统计(id, type, use, correct, accuracyRate)'],
            ['DetailItemVO', 'Record', '明细视图(含题目context和maxScore)'],
            ['UserExamHistoryVO', 'Record', '用户考试历史视图(examName, score, maxScore, accuracy, submitTime)'],
            ['DraftCacheService', 'class', '草稿缓存服务'],
        ])
    doc.add_paragraph()

    add_body_text(doc, 'DTO层是系统数据传输对象的核心，本系统共定义了43个DTO/VO/Req类（含DraftCacheService），全部使用Java 21的Record类型定义（除QuestionStatisticsVO和DraftCacheService为class）。DTO层严格遵循"Entity禁止裸奔"原则：所有API请求使用Req对象接收参数，所有API响应使用VO对象返回数据，Entity与DTO之间的转换在Service层的私有toVO()方法中完成。这种隔离设计防止了JSON循环引用和敏感字段泄露（如User.password不会出现在UserVO中），同时使得API契约的变更不会影响数据库结构。DTO按职责分为三类：Req（请求参数）、VO（视图对象）、Item（嵌套项），命名清晰、职责单一。')

    # 2.6 Common层
    add_heading2(doc, '2.6 Common层')

    add_table(doc,
        ['类名', '子包', '职责说明'],
        [
            ['Result<T>', 'api', '统一API响应封装(code/message/data)，提供success()/error()静态方法'],
            ['PageResult<T>', 'api', '分页响应封装(content/totalElements/totalPages/page/size)，提供of()转换方法'],
            ['BusinessException', 'exception', '业务异常类(code/message)，继承RuntimeException'],
            ['GlobalExceptionHandler', 'exception', '全局异常处理器，@RestControllerAdvice，处理3种异常类型'],
            ['JwtUtil', 'security', 'JWT工具类，提供generateToken/parseToken/validateToken/getUserId等方法'],
            ['JwtAuthenticationInterceptor', 'security', 'JWT认证拦截器，preHandle验证Token+检查@RequireRole权限'],
            ['RequireRole', 'security', '自定义注解@Target(METHOD)，声明方法所需角色'],
            ['WebMvcConfig', 'config', 'WebMvc配置，注册拦截器+静态资源映射'],
        ])
    doc.add_paragraph()

    add_body_text(doc, 'Common基础设施层是系统的横切关注点支撑层，不依赖任何业务模块，被所有Controller和Service共享使用。Result<T>统一响应封装确保了所有API返回格式一致（{code, message, data}），前端只需按统一格式处理响应。GlobalExceptionHandler通过@RestControllerAdvice自动捕获BusinessException、MethodArgumentNotValidException和Exception三种异常，分别返回业务错误码、参数校验错误和500服务器错误。JWT认证体系由JwtUtil（Token生成/解析）、JwtAuthenticationInterceptor（请求拦截验证）和@RequireRole（方法级权限声明）三个组件协作完成，实现了无状态认证和注解式权限控制。WebMvcConfig负责注册拦截器并配置静态资源映射（如图片路径映射）。')

    # ============================================================
    # 第三章 用例详细时序图（VOPC）
    # ============================================================
    add_heading1(doc, '第三章 用例详细时序图（VOPC）')

    # 3.1 M01模块
    add_heading2(doc, '3.1 M01模块')

    m01_sequences = [
        {
            'title': '图3-1 UC-M01-01 用户登录时序图',
            'code': """sequenceDiagram
    participant U as 用户
    participant C as UserController
    participant S as UserService
    participant R as UserRepository
    participant J as JwtUtil
    U->>C: POST /auth/login {name, password}
    C->>S: login(LoginReq)
    S->>R: findByName(name)
    R-->>S: User
    S->>S: BCrypt密码校验
    S->>S: 检查status!=0
    S-->>C: LoginResp(userVO)
    C->>J: generateToken(userId, name, type)
    J-->>C: token
    C-->>U: Result<LoginResp>{token, userVO}""",
            'desc': '用户登录时，系统先通过UserRepository.findByName查询用户记录，然后使用BCrypt密码校验。系统兼容明文密码，首次登录成功后自动升级为BCrypt加密。若账号被禁用(status=0)则拒绝登录。Token签发在Controller层完成，JwtUtil.generateToken将userId、name和type写入JWT payload。',
        },
        {
            'title': '图3-2 UC-M01-02 用户注册时序图',
            'code': """sequenceDiagram
    participant U as 用户
    participant C as UserController
    participant S as UserService
    participant R as UserRepository
    U->>C: POST /auth/register {name, password, type}
    C->>S: register(RegisterReq)
    S->>R: existsByName(name)
    R-->>S: boolean
    S->>S: 校验仅学生可注册
    S->>S: BCrypt加密密码
    S->>R: save(User)
    R-->>S: User
    S-->>C: UserVO
    C-->>U: Result<UserVO>""",
            'desc': '注册时系统先校验用户名唯一性(UserRepository.existsByName)，然后校验仅学生可自助注册。密码使用BCrypt加密后存储，新用户默认status=1(启用)。管理员创建用户走createUser方法，不限制角色类型。',
        },
        {
            'title': '图3-3 UC-M01-03 用户登出时序图',
            'code': """sequenceDiagram
    participant U as 用户
    participant C as UserController
    U->>C: POST /auth/logout
    C-->>U: Result<Void> (JWT无状态,前端删Token)""",
            'desc': '用户登出为无状态操作，系统仅返回成功响应。由于JWT是无状态Token，服务端不维护会话，登出逻辑由前端完成——删除本地存储的Token即可。后续请求因不携带Token而被JwtAuthenticationInterceptor拦截。',
        },
        {
            'title': '图3-4 UC-M01-04 获取当前用户信息时序图',
            'code': """sequenceDiagram
    participant U as 用户
    participant C as UserController
    participant S as UserService
    participant R as UserRepository
    U->>C: GET /auth/me
    C->>C: 从request获取userId
    C->>S: getCurrentUser(userId)
    S->>R: findById(userId)
    R-->>S: User
    S-->>C: UserVO
    C-->>U: Result<UserVO>""",
            'desc': '获取当前用户信息时，Controller从request属性中获取userId（由JwtAuthenticationInterceptor在preHandle阶段解析Token并注入），然后调用UserService.getCurrentUser查询用户并转换为UserVO返回，确保password等敏感字段不暴露。',
        },
        {
            'title': '图3-5 UC-M01-05 修改密码时序图',
            'code': """sequenceDiagram
    participant U as 用户
    participant C as UserController
    participant S as UserService
    participant R as UserRepository
    U->>C: POST /auth/password {oldPassword, newPassword}
    C->>C: 从request获取userId
    C->>S: changePassword(userId, req)
    S->>R: findById(userId)
    R-->>S: User
    S->>S: BCrypt校验旧密码
    S->>S: BCrypt加密新密码
    S->>R: save(User)
    S-->>C: void
    C-->>U: Result<Void>""",
            'desc': '修改密码时，系统先通过BCrypt校验旧密码是否正确，兼容明文密码自动升级机制。校验通过后，新密码使用BCrypt加密并更新到数据库。整个操作使用@Transactional保证事务一致性。',
        },
        {
            'title': '图3-6 UC-M01-06 分页查询用户列表时序图',
            'code': """sequenceDiagram
    participant A as 管理员
    participant C as UserController
    participant S as UserService
    participant R as UserRepository
    A->>C: GET /users?page=&size=&type=&status=
    C->>S: listUsers(page, size, type, status)
    S->>R: findAll/findByTypeAndStatus/findByType
    R-->>S: Page<User>
    S->>S: 转换为UserVO
    S-->>C: PageResult<UserVO>
    C-->>A: Result<PageResult<UserVO>>""",
            'desc': '分页查询用户列表支持按type和status筛选。Service层根据筛选条件组合选择不同的Repository查询方法（findAll/findByTypeAndStatus/findByType），查询结果通过toVO()转换为UserVO分页对象返回，密码字段被过滤。',
        },
        {
            'title': '图3-7 UC-M01-07 创建用户时序图',
            'code': """sequenceDiagram
    participant A as 管理员
    participant C as UserController
    participant S as UserService
    participant R as UserRepository
    A->>C: POST /users {name, password, type}
    C->>S: createUser(RegisterReq)
    S->>R: existsByName(name)
    R-->>S: boolean
    S->>S: BCrypt加密密码
    S->>R: save(User)
    R-->>S: User
    S-->>C: UserVO
    C-->>A: Result<UserVO>""",
            'desc': '管理员创建用户时不限制角色类型（可创建student/teacher/admin），但需校验用户名唯一性。密码使用BCrypt加密存储，新用户默认status=1。与自助注册的区别在于无"仅学生可注册"的限制。',
        },
        {
            'title': '图3-8 UC-M01-08 更新用户时序图',
            'code': """sequenceDiagram
    participant A as 管理员
    participant C as UserController
    participant S as UserService
    participant R as UserRepository
    A->>C: PUT /users/{id} {name, password, type}
    C->>S: updateUser(id, req)
    S->>R: findById(id)
    R-->>S: User
    S->>R: existsByName(name) [若name变更]
    S->>S: 更新字段
    S->>R: save(User)
    R-->>S: User
    S-->>C: UserVO
    C-->>A: Result<UserVO>""",
            'desc': '更新用户时，系统先查询目标用户是否存在，若修改了name则需校验新用户名的唯一性。密码字段若非空则BCrypt加密后更新。所有字段更新后通过Repository.save持久化，并转换为UserVO返回。',
        },
        {
            'title': '图3-9 UC-M01-09 更新用户状态时序图',
            'code': """sequenceDiagram
    participant A as 管理员
    participant C as UserController
    participant S as UserService
    participant R as UserRepository
    A->>C: PATCH /users/{id}/status {status}
    C->>S: updateUserStatus(id, req)
    S->>R: findById(id)
    R-->>S: User
    S->>S: 检查最后一个admin保护
    S->>R: save(User)
    S-->>C: void
    C-->>A: Result<Void>""",
            'desc': '更新用户状态时，系统内置"最后一个admin保护"机制：若禁用的用户是admin类型且是最后一个启用状态的admin，则拒绝操作并抛出BusinessException。这确保系统始终至少有一个可用的管理员账号。',
        },
        {
            'title': '图3-10 UC-M01-10 删除用户时序图',
            'code': """sequenceDiagram
    participant A as 管理员
    participant C as UserController
    participant S as UserService
    participant R as UserRepository
    participant SR as ScoreRepository
    A->>C: DELETE /users/{id}
    C->>C: 从request获取currentUserId
    C->>S: deleteUser(id, currentUserId)
    S->>S: 校验不能删除自己
    S->>R: findById(id)
    R-->>S: User
    S->>S: 检查最后一个admin保护
    S->>SR: findByUser(id)
    SR-->>S: List<Score>
    alt 有考试记录
        S->>R: save(User{status=0})
    else 无考试记录
        S->>R: delete(User)
    end
    S-->>C: void
    C-->>A: Result<Void>""",
            'desc': '删除用户时有三重保护：不能删除自己、不能删除最后一个admin、有考试记录的用户禁用而非硬删除。系统通过ScoreRepository.findByUser检查是否存在关联成绩记录，若存在则将status设为0（软禁用），保证成绩数据的参照完整性。',
        },
        {
            'title': '图3-11 UC-M01-11 批量删除用户时序图',
            'code': """sequenceDiagram
    participant A as 管理员
    participant C as UserController
    participant S as UserService
    A->>C: DELETE /users/batch {ids}
    C->>C: 从request获取currentUserId
    C->>S: batchDeleteUsers(ids, currentUserId)
    loop 逐个ID
        S->>S: deleteUser(id, currentUserId)
    end
    S-->>C: void
    C-->>A: Result<Void>""",
            'desc': '批量删除用户通过循环调用deleteUser方法实现，每个ID独立执行删除逻辑（包括自删除保护、admin保护、软禁用策略）。批量操作不使用@Transactional整体事务，单条失败不影响其他记录的处理。',
        },
    ]

    for seq in m01_sequences:
        add_mermaid_code(doc, seq['code'], seq['title'])
        add_body_text(doc, seq['desc'])

    # 3.2 M02模块
    add_heading2(doc, '3.2 M02模块')

    m02_sequences = [
        {
            'title': '图3-12 UC-M02-01 创建题目时序图',
            'code': """sequenceDiagram
    participant T as 教师/管理员
    participant C as QuestionController
    participant S as QuestionService
    participant R as QuestionRepository
    T->>C: POST /questions {type, context, img, answer}
    C->>S: create(QuestionCreateReq)
    S->>S: 校验答案JSON与题型匹配
    S->>R: save(Question{use=0, correct=0})
    R-->>S: Question
    S-->>C: QuestionVO
    C-->>T: Result<QuestionVO>""",
            'desc': '创建题目时，系统根据QuestionType使用switch表达式选择对应的Answer实现类（SingleChoiceAnswer/MultipleChoiceAnswer/JudgeAnswer/FillAnswer/EssayAnswer）进行JSON反序列化和业务校验，确保答案结构与题型匹配。新建题目use和correct均初始化为0。',
        },
        {
            'title': '图3-13 UC-M02-02 批量导入题目时序图',
            'code': """sequenceDiagram
    participant T as 教师/管理员
    participant C as QuestionController
    participant S as QuestionService
    participant R as QuestionRepository
    T->>C: POST /questions/batch [QuestionCreateReq...]
    C->>S: batchCreate(reqs)
    loop 逐题
        S->>S: 校验答案JSON
        alt 校验通过
            S->>R: save(Question)
        else 校验失败
            S->>S: 记录错误信息
        end
    end
    S-->>C: BatchImportResult
    C-->>T: Result<BatchImportResult>""",
            'desc': '批量导入限制不超过100题，采用逐题容错处理策略：每题独立校验答案JSON与题型匹配，校验通过则保存，失败则跳过并记录错误信息。最终返回BatchImportResult包含successCount、failCount和errors列表，方便用户定位问题。',
        },
        {
            'title': '图3-14 UC-M02-03 查询题目详情时序图',
            'code': """sequenceDiagram
    participant T as 教师/管理员
    participant C as QuestionController
    participant S as QuestionService
    participant R as QuestionRepository
    T->>C: GET /questions/{id}
    C->>S: findById(id)
    S->>R: findById(id)
    R-->>S: Question
    S->>S: 转换为QuestionVO
    S-->>C: QuestionVO
    C-->>T: Result<QuestionVO>""",
            'desc': '查询题目详情时，Service层通过toVO()方法将Entity转换为QuestionVO，其中包含答案JSON的解析和统计信息（accuracy=correct/use）。教师和管理员可查看完整答案，学生视角则使用QuestionPreviewVO过滤答案字段。',
        },
        {
            'title': '图3-15 UC-M02-04 分页查询题目列表时序图',
            'code': """sequenceDiagram
    participant T as 教师/管理员
    participant C as QuestionController
    participant S as QuestionService
    participant R as QuestionRepository
    T->>C: GET /questions?type=&keyword=&page=&size=
    C->>S: search(QuestionQueryReq)
    S->>R: searchByKeyword/findByType/findAll
    R-->>S: Page<Question>
    S->>S: 转换为QuestionVO
    S-->>C: Page<QuestionVO>
    C-->>T: Result<PageResult<QuestionVO>>""",
            'desc': '分页查询题目列表支持按type和keyword筛选。searchByKeyword使用@Query注解编写JPQL，实现按题型和关键字（模糊匹配context字段）的组合查询。查询结果逐条转换为QuestionVO，包含统计accuracy字段。',
        },
        {
            'title': '图3-16 UC-M02-05 更新题目时序图',
            'code': """sequenceDiagram
    participant T as 教师/管理员
    participant C as QuestionController
    participant S as QuestionService
    participant R as QuestionRepository
    T->>C: PUT /questions/{id} {context, img, answer}
    C->>S: update(id, QuestionUpdateReq)
    S->>R: findById(id)
    R-->>S: Question
    S->>S: 校验答案JSON
    S->>R: save(Question)
    R-->>S: Question
    S-->>C: QuestionVO
    C-->>T: Result<QuestionVO>""",
            'desc': '更新题目时，系统先查询目标题目是否存在，然后校验新的答案JSON与原题型是否匹配（题型不可修改）。校验通过后更新context、img、answer等字段并保存。use和correct统计字段不受更新操作影响。',
        },
        {
            'title': '图3-17 UC-M02-06 删除题目时序图',
            'code': """sequenceDiagram
    participant T as 教师/管理员
    participant C as QuestionController
    participant S as QuestionService
    participant R as QuestionRepository
    T->>C: DELETE /questions/{id}
    C->>S: delete(id)
    S->>R: findById(id)
    R-->>S: Question
    S->>R: delete(Question)
    S-->>C: void
    C-->>T: Result<Void>""",
            'desc': '删除题目采用硬删除策略，直接从数据库中移除记录。由于考试通过question_sum JSON快照引用题目，删除题目不会影响已组卷考试的数据完整性，但系统会记录警告日志提示该题目可能被考试快照引用。',
        },
        {
            'title': '图3-18 UC-M02-07 批量删除题目时序图',
            'code': """sequenceDiagram
    participant T as 教师/管理员
    participant C as QuestionController
    participant S as QuestionService
    T->>C: DELETE /questions/batch [ids]
    loop 逐个ID
        C->>S: delete(id)
    end
    C-->>T: Result<Void>""",
            'desc': '批量删除题目通过Controller层循环调用QuestionService.delete()方法实现，每个ID独立执行删除逻辑。与单条删除相同，采用硬删除策略并记录警告日志。',
        },
        {
            'title': '图3-19 UC-M02-08 随机获取题目时序图',
            'code': """sequenceDiagram
    participant T as 教师/管理员
    participant C as QuestionController
    participant S as QuestionService
    participant R as QuestionRepository
    T->>C: GET /questions/random?type=&excludedIds=
    C->>S: getRandomQuestion(type, excludedIds)
    S->>R: findRandomByType(type, excludedIds)
    R-->>S: Question
    S->>S: 转换为QuestionPreviewVO
    S-->>C: QuestionPreviewVO
    C-->>T: Result<QuestionPreviewVO>""",
            'desc': '随机获取题目用于自动组卷场景，支持按题型过滤和排除已选题目ID列表。Repository层使用原生SQL的ORDER BY RANDOM() LIMIT 1实现随机抽取。返回QuestionPreviewVO不含答案，防止答案泄露。',
        },
    ]

    for seq in m02_sequences:
        add_mermaid_code(doc, seq['code'], seq['title'])
        add_body_text(doc, seq['desc'])

    # 3.3 M03模块
    add_heading2(doc, '3.3 M03模块')

    m03_sequences = [
        {
            'title': '图3-20 UC-M03-01 创建手动组卷考试时序图',
            'code': """sequenceDiagram
    participant T as 教师/管理员
    participant C as ExamController
    participant S as ExamService
    participant ER as ExamRepository
    participant QR as QuestionRepository
    T->>C: POST /exams/manual {exam, starttime, endtime, items[]}
    C->>S: createManualExam(ExamCreateManualReq)
    S->>QR: findById(questionId) [逐题]
    QR-->>S: Question
    S->>S: 构造question_sum JSON快照
    S->>QR: incrementUse(questionId) [逐题]
    S->>ER: save(Exam{status=draft})
    ER-->>S: Exam
    S-->>C: ExamVO
    C-->>T: Result<ExamVO>""",
            'desc': '手动组卷时，系统逐题校验questionId是否存在，然后构造QuestionSum JSON快照（包含version、items列表、totalQuestions和totalScore）。组卷完成后为每个被抽中题目调用incrementUse原子自增use字段。考试初始状态为draft。',
        },
        {
            'title': '图3-21 UC-M03-02 创建自动组卷考试时序图',
            'code': """sequenceDiagram
    participant T as 教师/管理员
    participant C as ExamController
    participant S as ExamService
    participant ER as ExamRepository
    participant QR as QuestionRepository
    T->>C: POST /exams/auto {exam, starttime, endtime, rule}
    C->>S: createAutoExam(ExamCreateAutoReq)
    S->>QR: findByType(typeFilter)
    QR-->>S: List<Question>
    S->>S: 按rule随机抽题(usePenalty加权)
    S->>S: 构造question_sum JSON快照
    S->>QR: incrementUse(questionId) [逐题]
    S->>ER: save(Exam{status=draft})
    ER-->>S: Exam
    S-->>C: ExamVO
    C-->>T: Result<ExamVO>""",
            'desc': '自动组卷使用usePenalty加权随机算法：权重=1/(1+use)，use值越大的热点题抽中概率越低，促进题目均衡使用。系统按typeFilter筛选候选题目，再按totalQuestions和totalScore规则抽取，最后构造QuestionSum快照并自增use统计。',
        },
        {
            'title': '图3-22 UC-M03-03 获取可参加考试列表时序图',
            'code': """sequenceDiagram
    participant S as 学生
    participant C as ExamController
    participant ES as ExamService
    participant R as ExamRepository
    S->>C: GET /exams/available
    C->>ES: listAvailableExams()
    ES->>R: findByStatus(publish/running)
    R-->>ES: List<Exam>
    ES->>ES: 实时计算状态(时间窗)
    ES->>ES: 转换为ExamForStudentVO(脱敏)
    ES-->>C: List<ExamForStudentVO>
    C-->>S: Result<List<ExamForStudentVO>>""",
            'desc': '获取可参加考试列表时，系统查询publish和running状态的考试，然后通过resolveCurrentStatus()方法按当前时间与starttime/endtime的对比实时判定考试状态。返回ExamForStudentVO进行答案脱敏处理，剔除correctOption等判分关键字段。',
        },
        {
            'title': '图3-23 UC-M03-04 获取考试详情时序图',
            'code': """sequenceDiagram
    participant T as 教师/管理员
    participant C as ExamController
    participant S as ExamService
    participant R as ExamRepository
    T->>C: GET /exams/{id}
    C->>S: getExamById(id)
    S->>R: findById(id)
    R-->>S: Exam
    S->>S: 解析question_sum JSON
    S-->>C: ExamVO
    C-->>T: Result<ExamVO>""",
            'desc': '获取考试详情时，Service层解析exam.question_sum JSON字段，还原题目列表和分值信息。教师和管理员视角返回ExamVO包含完整答案，同时通过resolveCurrentStatus()实时计算考试当前状态。',
        },
        {
            'title': '图3-24 UC-M03-05 学生预览考试时序图',
            'code': """sequenceDiagram
    participant S as 学生
    participant C as ExamController
    participant ES as ExamService
    participant R as ExamRepository
    S->>C: GET /exams/{id}/preview
    C->>ES: getExamForStudent(id)
    ES->>R: findById(id)
    R-->>ES: Exam
    ES->>ES: 转换为ExamForStudentVO(脱敏)
    ES-->>C: ExamForStudentVO
    C-->>S: Result<ExamForStudentVO>""",
            'desc': '学生预览考试时，系统通过extractOptionsOnly()方法对答案进行脱敏处理：保留options选项列表，但剔除correctOption、correctOptions、correct、blanks、reference等判分关键字段，防止答案泄露。返回ExamForStudentVO包含题目选项但不含答案。',
        },
        {
            'title': '图3-25 UC-M03-06 修改考试时序图',
            'code': """sequenceDiagram
    participant T as 教师/管理员
    participant C as ExamController
    participant S as ExamService
    participant R as ExamRepository
    T->>C: PUT /exams/{id} {exam, starttime, endtime, items[]}
    C->>S: updateExam(id, req)
    S->>R: findById(id)
    R-->>S: Exam
    S->>S: 校验status=draft
    S->>R: save(Exam)
    R-->>S: Exam
    S-->>C: ExamVO
    C-->>T: Result<ExamVO>""",
            'desc': '修改考试仅允许draft状态，系统校验状态后更新考试信息。若题目列表变更，需要处理use统计的增量/回退：新增题目incrementUse，移除题目decrementUse，保留题目不变。最后重构question_sum JSON快照。',
        },
        {
            'title': '图3-26 UC-M03-07 发布考试时序图',
            'code': """sequenceDiagram
    participant T as 教师/管理员
    participant C as ExamController
    participant S as ExamService
    participant R as ExamRepository
    T->>C: POST /exams/{id}/publish
    C->>S: publishExam(id)
    S->>R: findById(id)
    R-->>S: Exam
    S->>S: 校验status=draft
    S->>R: save(Exam{status=publish})
    S-->>C: void
    C-->>T: Result<Void>""",
            'desc': '发布考试实现draft→publish的状态流转。系统校验当前状态必须为draft，然后将status更新为publish。发布后考试对学生可见，学生可通过listAvailableExams接口看到该考试。实际运行状态由resolveCurrentStatus()按时间窗实时判定。',
        },
        {
            'title': '图3-27 UC-M03-08 撤回考试时序图',
            'code': """sequenceDiagram
    participant T as 教师/管理员
    participant C as ExamController
    participant S as ExamService
    participant R as ExamRepository
    T->>C: POST /exams/{id}/withdraw
    C->>S: withdrawExam(id)
    S->>R: findById(id)
    R-->>S: Exam
    S->>S: 校验status=publish
    S->>R: save(Exam{status=draft})
    S-->>C: void
    C-->>T: Result<Void>""",
            'desc': '撤回考试实现publish→draft的逆向状态流转。系统校验当前状态必须为publish，然后将status回退为draft。撤回后考试对学生不可见。注意running和done状态不可撤回，保证考试过程的不可篡改性。',
        },
        {
            'title': '图3-28 UC-M03-09 删除考试时序图',
            'code': """sequenceDiagram
    participant T as 教师/管理员
    participant C as ExamController
    participant S as ExamService
    participant R as ExamRepository
    T->>C: DELETE /exams/{id}
    C->>S: deleteExam(id)
    S->>R: findById(id)
    R-->>S: Exam
    S->>S: 校验status=draft
    S->>R: delete(Exam)
    S-->>C: void
    C-->>T: Result<Void>""",
            'desc': '删除考试仅允许draft状态，删除前系统解析question_sum JSON，为每个被引用的题目调用decrementUse回退use统计，确保题目使用次数的准确性。非draft状态的考试不可删除，防止误删正在进行或已结束的考试。',
        },
        {
            'title': '图3-29 UC-M03-10 分页查询考试列表时序图',
            'code': """sequenceDiagram
    participant T as 教师/管理员
    participant C as ExamController
    participant S as ExamService
    participant R as ExamRepository
    T->>C: GET /exams?status=&page=&size=
    C->>S: listExams(page, size, status)
    S->>R: findByStatusNot(status, pageable)
    R-->>S: Page<Exam>
    S->>S: 转换为ExamVO
    S-->>C: PageResult<ExamVO>
    C-->>T: Result<PageResult<ExamVO>>""",
            'desc': '分页查询考试列表支持按status筛选，使用findByStatusNot排除指定状态。查询结果通过toVOWithResolvedStatus()转换，该方法调用resolveCurrentStatus()实时计算考试当前状态（根据starttime/endtime与当前时间的对比判定draft/publish/running/done）。',
        },
    ]

    for seq in m03_sequences:
        add_mermaid_code(doc, seq['code'], seq['title'])
        add_body_text(doc, seq['desc'])

    # 3.4 M04模块
    add_heading2(doc, '3.4 M04模块')

    m04_sequences = [
        {
            'title': '图3-30 UC-M04-01 提交答卷时序图',
            'code': """sequenceDiagram
    participant S as 学生
    participant C as ScoreController
    participant SS as ScoreService
    participant SR as ScoreRepository
    participant ER as ExamRepository
    participant QR as QuestionRepository
    S->>C: POST /exams/{examId}/submit {answers[]}
    C->>C: 从request获取userId
    C->>SS: submitExam(ExamSubmitReq, userId)
    SS->>ER: findById(examId)
    ER-->>SS: Exam
    SS->>SS: 校验考试状态running
    SS->>SS: 解析question_sum确定题序分值
    loop 逐题判分
        SS->>QR: findById(questionId)
        QR-->>SS: Question
        SS->>SS: 解析answer JSON并判分
    end
    SS->>SS: 计算总分+构造detail JSON
    SS->>SR: upsertScore(userId, examId, all, detail)
    loop isCorrect=true的题目
        SS->>QR: incrementCorrect(questionId)
    end
    SS-->>C: ScoreVO
    C-->>S: Result<ScoreVO>""",
            'desc': '提交答卷时，系统先校验考试状态为running，然后解析question_sum确定题序和分值。gradeOne()方法按QuestionType分别实现五种判分逻辑：单选(精确匹配)、多选(集合等价)、判断(布尔比对)、填空(逐空忽略大小写)、简答(待评卷，score=0，isCorrect=null)。判分完成后构造ScoreDetail JSON并UPSERT写入score表。',
        },
        {
            'title': '图3-31 UC-M04-02 教师评卷时序图',
            'code': """sequenceDiagram
    participant T as 教师/管理员
    participant C as ScoreController
    participant S as ScoreService
    participant SR as ScoreRepository
    participant QR as QuestionRepository
    T->>C: POST /scores/{scoreId}/grade-essay {items[]}
    C->>S: gradeEssay(scoreId, EssayGradeReq)
    S->>SR: findById(scoreId)
    SR-->>S: Score
    S->>S: 解析detail JSON
    loop 逐题评卷
        S->>S: 更新简答题score和isCorrect
        alt isCorrect变为true
            S->>QR: incrementCorrect(questionId)
        end
    end
    S->>S: 重新计算总分和准确率
    S->>SR: save(Score)
    S-->>C: ScoreVO
    C-->>T: Result<ScoreVO>""",
            'desc': '教师评卷时，系统解析score.detail JSON，逐题更新简答题的得分和isCorrect状态。关键设计：仅当isCorrect从非true变为true时才调用incrementCorrect更新question.correct统计，避免重复计数。评卷完成后重新计算总分和准确率并保存。',
        },
        {
            'title': '图3-32 UC-M04-03 查询我的成绩时序图',
            'code': """sequenceDiagram
    participant U as 用户
    participant C as ScoreController
    participant S as ScoreService
    participant R as ScoreRepository
    U->>C: GET /scores/me?page=&size=
    C->>C: 从request获取userId
    C->>S: getMyScores(userId, page, size)
    S->>R: findByUser(userId, pageable)
    R-->>S: Page<Score>
    S-->>C: PageResult<ScoreListVO>
    C-->>U: Result<PageResult<ScoreListVO>>""",
            'desc': '查询我的成绩时，系统从请求属性获取userId，然后分页查询该用户的所有成绩记录。返回ScoreListVO不含逐题明细，仅包含考试名称、总分等概要信息，减少数据传输量。需要关联Exam信息获取考试名称。',
        },
        {
            'title': '图3-33 UC-M04-04 查询我的错题集时序图',
            'code': """sequenceDiagram
    participant S as 学生
    participant C as ScoreController
    participant SS as ScoreService
    participant SR as ScoreRepository
    S->>C: GET /scores/me/mistakes?page=&size=
    C->>C: 从request获取userId
    C->>SS: getMyMistakes(userId, page, size)
    SS->>SR: findByUser(userId)
    SR-->>SS: List<Score>
    SS->>SS: 解析detail,筛选isCorrect=false
    SS-->>C: PageResult<MistakeItemVO>
    C-->>S: Result<PageResult<MistakeItemVO>>""",
            'desc': '查询错题集时，系统获取用户所有成绩记录，解析每条记录的detail JSON，筛选isCorrect为false或null的题目。使用LinkedHashMap去重策略，同一题只保留最近一次错误记录，保证错题集的简洁性。返回MistakeItemVO包含题目内容和正确答案。',
        },
        {
            'title': '图3-34 UC-M04-05 查询分数详情时序图',
            'code': """sequenceDiagram
    participant U as 用户
    participant C as ScoreController
    participant S as ScoreService
    participant R as ScoreRepository
    U->>C: GET /scores/{id}
    C->>S: findById(id)
    S->>R: findById(id)
    R-->>S: Score
    S->>S: 解析detail JSON
    S-->>C: ScoreVO
    C-->>U: Result<ScoreVO>""",
            'desc': '查询分数详情时，系统解析score.detail JSON，还原逐题答题明细（包括用户答案、正确答案、得分、是否正确）。同时关联Question表获取题目context和maxScore，构造完整的ScoreVO返回，方便学生查看每道题的作答情况。',
        },
        {
            'title': '图3-35 UC-M04-06 查询考试所有考生分数时序图',
            'code': """sequenceDiagram
    participant T as 教师/管理员
    participant C as ScoreController
    participant S as ScoreService
    participant R as ScoreRepository
    T->>C: GET /exams/{examId}/scores?page=&size=
    C->>S: getExamScores(examId, page, size)
    S->>R: findByExam(examId, pageable)
    R-->>S: Page<Score>
    S-->>C: PageResult<ScoreListVO>
    C-->>T: Result<PageResult<ScoreListVO>>""",
            'desc': '查询考试所有考生分数时，系统按examId分页查询成绩记录，返回ScoreListVO列表。每条记录关联User表获取考生姓名，关联Exam表获取考试名称，构造ScoreListVO包含考生信息和分数概要。',
        },
        {
            'title': '图3-36 UC-M04-07 考试统计报表时序图',
            'code': """sequenceDiagram
    participant T as 教师/管理员
    participant C as ScoreController
    participant S as ScoreService
    participant SR as ScoreRepository
    participant ER as ExamRepository
    T->>C: GET /statistics/exams/{examId}
    C->>S: getExamStatistics(examId)
    S->>ER: findById(examId)
    ER-->>S: Exam
    S->>SR: findByExam(examId)
    SR-->>S: List<Score>
    S->>S: 计算参与人数/平均分/最高分/最低分/通过率
    S-->>C: ExamStatisticsVO
    C-->>T: Result<ExamStatisticsVO>""",
            'desc': '考试统计报表计算参与人数、平均分、最高分、最低分、中位数和通过率（60%及格线）。分数分布按10分一档统计（0-10, 10-20, ..., 90-100），返回ExamStatisticsVO包含完整的统计信息和分布数据，辅助教师评估考试难度和教学效果。',
        },
        {
            'title': '图3-37 UC-M04-08 题目统计列表时序图',
            'code': """sequenceDiagram
    participant T as 教师/管理员
    participant C as ScoreController
    participant S as ScoreService
    participant R as QuestionRepository
    T->>C: GET /statistics/questions?page=&size=&sortBy=
    C->>S: getQuestionStatisticsPaginated(page, size, sortBy)
    S->>R: findAll(pageable)
    R-->>S: Page<Question>
    S->>S: 计算每题正确率=correct/use
    S-->>C: PageResult<QuestionStatisticsVO>
    C-->>T: Result<PageResult<QuestionStatisticsVO>>""",
            'desc': '题目统计列表查询所有题目的使用次数、正确次数和正确率(correct/use)，支持按accuracyRate、use、correct等字段排序。正确率为0的情况（use=0时）特殊处理为N/A。返回QuestionStatisticsVO分页结果，辅助教师评估题目质量和难度。',
        },
        {
            'title': '图3-38 UC-M04-09 单题详细统计时序图',
            'code': """sequenceDiagram
    participant T as 教师/管理员
    participant C as ScoreController
    participant S as ScoreService
    participant R as QuestionRepository
    T->>C: GET /statistics/questions/{id}
    C->>S: getQuestionStatisticById(id)
    S->>R: findById(id)
    R-->>S: Question
    S->>S: 计算正确率=correct/use
    S-->>C: QuestionStatisticsVO
    C-->>T: Result<QuestionStatisticsVO>""",
            'desc': '单题详细统计查询指定题目的使用次数、正确次数和正确率。系统从Question实体中读取use和correct字段，计算accuracyRate=correct/use（use为0时返回0）。返回QuestionStatisticsVO包含题目类型和完整统计信息。',
        },
    ]

    for seq in m04_sequences:
        add_mermaid_code(doc, seq['code'], seq['title'])
        add_body_text(doc, seq['desc'])

    # ============================================================
    # 第四章 核心类详细定义
    # ============================================================
    add_heading1(doc, '第四章 核心类详细定义')

    # 4.1 UserService
    add_heading2(doc, '4.1 UserService')

    add_mermaid_code(doc, """classDiagram
    class UserService {
        -UserRepository userRepository
        -ScoreRepository scoreRepository
        -BCryptPasswordEncoder passwordEncoder
        +register(RegisterReq) UserVO @Transactional
        +login(LoginReq) LoginResp @Transactional
        +getCurrentUser(Integer) UserVO
        +changePassword(Integer, ChangePasswordReq) void @Transactional
        +listUsers(Integer, Integer, UserType, Integer) PageResult~UserVO~
        +createUser(RegisterReq) UserVO @Transactional
        +updateUser(Integer, RegisterReq) UserVO @Transactional
        +updateUserStatus(Integer, UserStatusReq) void @Transactional
        +deleteUser(Integer, Integer) void @Transactional
        +batchDeleteUsers(List~Integer~, Integer) void @Transactional
        +updateStatus(Integer, UserStatusReq) void
        +findById(Integer) UserVO
        +delete(Integer, Integer) void
        -toVO(User) UserVO
    }""", '图4-1 UserService类图')

    add_body_text(doc, 'UserService是用户认证与管理的核心服务，注入了UserRepository和ScoreRepository两个数据访问组件。关键设计包括：（1）BCrypt密码校验兼容明文自动升级；（2）仅学生可自助注册的权限控制；（3）删除用户时的自删除保护和最后admin保护；（4）有考试记录的用户禁用而非删除的级联策略。所有写操作均标注@Transactional保证事务一致性。')

    # 4.2 QuestionService
    add_heading2(doc, '4.2 QuestionService')

    add_mermaid_code(doc, """classDiagram
    class QuestionService {
        -QuestionRepository questionRepository
        -ObjectMapper objectMapper
        -ExamRepository examRepository
        +create(QuestionCreateReq) QuestionVO @Transactional
        +batchCreate(List~QuestionCreateReq~) BatchImportResult @Transactional
        +findById(Integer) QuestionVO
        +search(QuestionQueryReq) Page~QuestionVO~
        +update(Integer, QuestionUpdateReq) QuestionVO @Transactional
        +delete(Integer) void @Transactional
        +incrementUse(Integer) void @Transactional
        +incrementCorrect(Integer) void @Transactional
        +decrementUse(Integer) void @Transactional
        +getRandomQuestion(QuestionType, List~Integer~) QuestionPreviewVO
        -answerToJsonString(Object) String
        -validateAnswerJson(QuestionType, String) void
        -toVO(Question) QuestionVO
    }""", '图4-2 QuestionService类图')

    add_body_text(doc, 'QuestionService是题库管理的核心服务，注入了QuestionRepository、ObjectMapper和ExamRepository三个组件。关键设计包括：（1）答案JSON多态校验——根据QuestionType使用switch表达式选择对应的Answer实现类进行反序列化和业务校验；（2）批量导入限制（不超过100题）和逐题容错处理；（3）题内统计自维护——incrementUse/incrementCorrect/decrementUse三个原子操作方法；（4）删除被引用题目时的警告日志。')

    # 4.3 ExamService
    add_heading2(doc, '4.3 ExamService')

    add_mermaid_code(doc, """classDiagram
    class ExamService {
        -ExamRepository examRepository
        -QuestionRepository questionRepository
        -QuestionService questionService
        -ObjectMapper objectMapper
        -String userDir
        +createManualExam(ExamCreateManualReq) ExamVO @Transactional
        +createManual(ExamCreateManualReq) ExamVO @Transactional
        +createAutoExam(ExamCreateAutoReq) ExamVO @Transactional
        +createAuto(ExamCreateAutoReq) ExamVO @Transactional
        +getExamById(Integer) ExamVO
        +getExamForStudent(Integer) ExamForStudentVO
        +publishExam(Integer) void @Transactional
        +publish(Integer) ExamVO @Transactional
        +withdrawExam(Integer) void @Transactional
        +withdraw(Integer) ExamVO
        +updateExam(Integer, ExamCreateManualReq) ExamVO @Transactional
        +deleteExam(Integer) void @Transactional
        +delete(Integer) void
        +findById(Integer) ExamVO
        +edit(Integer, ExamCreateManualReq) ExamVO
        +listExams(Integer, Integer, ExamStatus) PageResult~ExamVO~
        +listAvailableExams() List~ExamForStudentVO~
        +resolveCurrentStatus(Exam, LocalDateTime) ExamStatus
        -validateTimeWindow(String, String) void
        -weightedRandomPick(List~Question~, int) List~Question~
        -extractOptionsOnly(Question) Object
        -resolveImageUrl(Integer, Integer) String
        -parseQuestionSum(String) QuestionSum
        -toJson(QuestionSum) String
        -toVO(Exam) ExamVO
        -toVOWithResolvedStatus(Exam) ExamVO
        -toPageResult(List~ExamVO~, int, int) PageResult~ExamVO~
    }""", '图4-3 ExamService类图')

    add_body_text(doc, 'ExamService是组卷与考试管理的核心服务，是系统中最复杂的Service。关键设计包括：（1）快照式组卷——构造QuestionSum JSON一次性快照，题目后续修改不影响已组卷考试；（2）考试状态机——draft↔publish→running→done四态流转，通过resolveCurrentStatus()按时间窗实时判定；（3）自动组卷的usePenalty加权随机算法——权重=1/(1+use)，热点题抽中概率降低；（4）更新考试时的use统计增量/回退处理；（5）学生预览时的答案脱敏（extractOptionsOnly剔除correctOption/correctOptions等判分关键字段）。')

    # 4.4 ScoreService
    add_heading2(doc, '4.4 ScoreService')

    add_mermaid_code(doc, """classDiagram
    class ScoreService {
        -ScoreRepository scoreRepository
        -ExamRepository examRepository
        -QuestionRepository questionRepository
        -QuestionService questionService
        -ExamService examService
        -ObjectMapper objectMapper
        -UserRepository userRepository
        -DraftCacheService draftCacheService
        +submitExam(ExamSubmitReq, Integer) ScoreVO @Transactional
        +gradeEssay(Integer, EssayGradeReq) ScoreVO @Transactional
        +findByUser(Integer) List~ScoreVO~
        +findByExam(Integer) List~ScoreVO~
        +findById(Integer) ScoreVO
        +findMistakes(Integer) List~MistakeItemVO~
        +getExamStatistics(Integer) ExamStatisticsVO
        +getQuestionStatistics() List~QuestionStatisticsVO~
        +getMyScores(Integer, int, int) PageResult~ScoreListVO~
        +getExamScores(Integer, int, int) PageResult~ScoreListVO~
        +getMyMistakes(Integer, int, int) PageResult~MistakeItemVO~
        +getQuestionStatisticsPaginated(int, int, String) PageResult~QuestionStatisticsVO~
        +getQuestionStatisticById(Integer) QuestionStatisticsVO
        -gradeOne(Question, Object, int) GradingResult
        -extractCorrectAnswer(Question) Object
        -extractOptions(Question) List~String~
        -parseQuestionSum(String) QuestionSum
        -parseScoreDetail(String) ScoreDetail
        -toJson(Object) String
        -toScoreListVO(Score, Exam) ScoreListVO
        -toPageResult(List~T~, int, int, int) PageResult~T~
        -toScoreVO(Score, Exam, Map, QuestionSum) ScoreVO
    }""", '图4-4 ScoreService类图')

    add_body_text(doc, 'ScoreService是成绩统计的核心服务，注入了8个依赖组件，是依赖最多的Service。关键设计包括：（1）答题判分的多态处理——gradeOne()方法使用switch表达式按QuestionType分别实现单选(精确匹配)、多选(集合等价)、判断(布尔比对)、填空(逐空忽略大小写)、简答(待评卷)五种判分逻辑；（2）教师评卷时的isCorrect变更检测——仅当isCorrect从非true变为true时才更新question.correct；（3）考试统计报表的通过率计算（60%及格线）和分数分布（10分一档）；（4）错题集的去重策略——使用LinkedHashMap保证同一题只保留最近一次错误记录。')

    # 4.5 JwtUtil
    add_heading2(doc, '4.5 JwtUtil')

    add_mermaid_code(doc, """classDiagram
    class JwtUtil {
        -String secret
        -long expiration
        +generateToken(Integer userId, String name, String type) String
        +parseToken(String token) Claims
        +validateToken(String token) boolean
        +getUserId(String token) Integer
        +getName(String token) String
        +getType(String token) String
        -getSigningKey() SecretKey
    }""", '图4-5 JwtUtil类图')

    add_body_text(doc, 'JwtUtil是JWT无状态认证的核心工具类，使用HMAC-SHA签名算法。Token的payload包含subject(userId)、name和type三个声明，过期时间由application.yaml的jwt.expiration配置。generateToken()签发Token，parseToken()解析并验证签名，validateToken()通过尝试解析判断Token有效性。getUserId/getName/getType三个方法从Token中提取用户信息，供JwtAuthenticationInterceptor使用。')

    # 4.6 Result<T>
    add_heading2(doc, '4.6 Result<T>')

    add_mermaid_code(doc, """classDiagram
    class Result~T~ {
        -int code
        -String message
        -T data
        +success() Result~Void~
        +success(T data) Result~T~
        +error(int code, String message) Result~T~
        +getCode() int
        +getMessage() String
        +getData() T
    }""", '图4-6 Result<T>类图')

    add_body_text(doc, 'Result<T>是统一API响应封装类，所有Controller方法均返回Result<T>类型。成功响应code=200，message="success"；失败响应code=业务错误码（如4101用户名已存在、4201题目不存在、4301考试不存在、4401分数记录不存在等），message=错误描述。泛型T允许data字段承载任意类型的响应数据，包括UserVO、QuestionVO、PageResult<ExamVO>等。')

    # 4.7 GlobalExceptionHandler
    add_heading2(doc, '4.7 GlobalExceptionHandler')

    add_mermaid_code(doc, """classDiagram
    class GlobalExceptionHandler {
        +handleBusiness(BusinessException) Result~Void~
        +handleIllegalArgument(MethodArgumentNotValidException) Result~Void~
        +handleAny(Exception) Result~Void~
    }""", '图4-7 GlobalExceptionHandler类图')

    add_body_text(doc, 'GlobalExceptionHandler通过@RestControllerAdvice注解实现全局异常处理，统一将异常转换为Result<Void>响应。handleBusiness()处理BusinessException，返回业务错误码和消息；handleIllegalArgument()处理MethodArgumentNotValidException，提取参数校验错误信息；handleAny()兜底处理所有未捕获异常，返回500服务器错误。这种设计确保了系统不会向客户端暴露堆栈信息，同时保持了API响应格式的一致性。')

    # ============================================================
    # 第五章 数据库表结构设计
    # ============================================================
    add_heading1(doc, '第五章 数据库表结构设计')

    # 5.1 user表
    add_heading2(doc, '5.1 user表')

    add_table(doc,
        ['字段名', 'Java类型', 'SQLite类型', '约束', '默认值', '业务含义'],
        [
            ['id', 'Integer', 'INTEGER', 'PK, AUTOINCREMENT', '—', '主键'],
            ['name', 'String', 'TEXT', 'NOT NULL, UNIQUE', '—', '用户名'],
            ['password', 'String', 'TEXT', 'NOT NULL', '—', '密码(BCrypt哈希)'],
            ['type', 'UserType(String)', 'TEXT', "NOT NULL, CHECK(IN('student','teacher','admin'))", "'student'", '用户类型'],
            ['status', 'Integer', 'INTEGER', 'NOT NULL, CHECK(IN(0,1))', '1', '状态(1=启用,0=禁用)'],
        ])
    doc.add_paragraph()

    add_body_text(doc, '索引：UNIQUE INDEX idx_user_name(name), INDEX idx_user_type(type)')

    # 5.2 question表
    add_heading2(doc, '5.2 question表')

    add_table(doc,
        ['字段名', 'Java类型', 'SQLite类型', '约束', '默认值', '业务含义'],
        [
            ['id', 'Integer', 'INTEGER', 'PK, AUTOINCREMENT', '—', '主键'],
            ['type', 'QuestionType(String)', 'TEXT', "NOT NULL, CHECK(IN('SingleChoice','MultipleChoice','Judge','Fill','Essay'))", '—', '题目类型'],
            ['context', 'String', 'TEXT', 'NOT NULL', '—', '题干'],
            ['img', 'Integer', 'INTEGER', 'NOT NULL, CHECK(IN(0,1))', '0', '是否带图(1=是,0=否)'],
            ['answer', 'String', 'TEXT', 'NOT NULL', '—', '答案JSON(按type适配5种结构)'],
            ['use', 'Integer', 'INTEGER', 'NOT NULL, CHECK(>=0)', '0', '被抽中次数'],
            ['correct', 'Integer', 'INTEGER', 'NOT NULL, CHECK(>=0 AND <=use)', '0', '被回答正确次数'],
        ])
    doc.add_paragraph()

    add_body_text(doc, '索引：INDEX idx_question_type(type), INDEX idx_question_use(use)')

    add_body_text(doc, 'answer JSON 5种结构：')
    add_body_text(doc, '（1）SingleChoice: {"version":1, "correctOption":"A", "options":["A","B","C","D"]}')
    add_body_text(doc, '（2）MultipleChoice: {"version":1, "correctOptions":["A","C"], "options":["A","B","C","D"]}')
    add_body_text(doc, '（3）Judge: {"version":1, "correct":true}')
    add_body_text(doc, '（4）Fill: {"version":1, "blanks":["went","goes"]}')
    add_body_text(doc, '（5）Essay: {"version":1, "reference":"参考答案"}')

    # 5.3 exam表
    add_heading2(doc, '5.3 exam表')

    add_table(doc,
        ['字段名', 'Java类型', 'SQLite类型', '约束', '默认值', '业务含义'],
        [
            ['id', 'Integer', 'INTEGER', 'PK, AUTOINCREMENT', '—', '主键'],
            ['exam', 'String', 'TEXT', 'NOT NULL', '—', '考试名称'],
            ['status', 'ExamStatus(String)', 'TEXT', "NOT NULL, CHECK(IN('draft','publish','running','done'))", "'draft'", '考试状态'],
            ['starttime', 'String', 'TEXT', 'NOT NULL', '—', '开始时间(ISO 8601)'],
            ['endtime', 'String', 'TEXT', 'NOT NULL, CHECK(>starttime)', '—', '结束时间(ISO 8601)'],
            ['question_sum', 'String', 'TEXT', 'NOT NULL', '—', '题目汇总JSON快照'],
        ])
    doc.add_paragraph()

    add_body_text(doc, '索引：INDEX idx_exam_status(status), INDEX idx_exam_time(starttime, endtime)')

    add_body_text(doc, 'question_sum JSON结构：')
    add_body_text_no_indent(doc, '{')
    add_body_text_no_indent(doc, '  "version": 1,')
    add_body_text_no_indent(doc, '  "items": [{"questionId": 1, "score": 5, "type": "SingleChoice"}],')
    add_body_text_no_indent(doc, '  "totalQuestions": 1,')
    add_body_text_no_indent(doc, '  "totalScore": 5')
    add_body_text_no_indent(doc, '}')

    add_body_text(doc, '状态机：draft ↔ publish → running → done')

    # 5.4 score表
    add_heading2(doc, '5.4 score表')

    add_table(doc,
        ['字段名', 'Java类型', 'SQLite类型', '约束', '默认值', '业务含义'],
        [
            ['id', 'Integer', 'INTEGER', 'PK, AUTOINCREMENT', '—', '主键'],
            ['user', 'Integer', 'INTEGER', 'NOT NULL, FK→user.id', '—', '考生ID'],
            ['exam', 'Integer', 'INTEGER', 'NOT NULL, FK→exam.id', '—', '考试ID'],
            ['all', 'Integer', 'INTEGER', 'NOT NULL, CHECK(>=0)', '—', '总分'],
            ['detail', 'String', 'TEXT', 'NOT NULL', '—', '答题明细JSON'],
        ])
    doc.add_paragraph()

    add_body_text(doc, '约束：UNIQUE(user, exam) — 一人一考仅一条记录')
    add_body_text(doc, '索引：INDEX idx_score_user(user), INDEX idx_score_exam(exam)')

    add_body_text(doc, 'detail JSON结构：')
    add_body_text_no_indent(doc, '{')
    add_body_text_no_indent(doc, '  "version": 1,')
    add_body_text_no_indent(doc, '  "items": [{"questionId": 1, "userAnswer": "B", "correctAnswer": "B", "score": 5, "isCorrect": true}],')
    add_body_text_no_indent(doc, '  "summary": {"correctCount": 1, "totalCount": 1, "accuracy": 1.0}')
    add_body_text_no_indent(doc, '}')

    # 5.5 表间关系说明
    add_heading2(doc, '5.5 表间关系说明')

    add_body_text(doc, '（1）物理外键关系：score.user → user.id（一对多，一个用户可有多条成绩记录）；score.exam → exam.id（一对多，一个考试可有多条考生成绩）。这两个外键由DDL的FOREIGN KEY约束定义，但JPA层面不使用@ManyToOne关联，而是使用Integer类型字段+@Column注解。')

    add_body_text(doc, '（2）逻辑引用关系：exam.question_sum → question.id（多对多快照引用），通过JSON字段中的questionId列表实现，组卷时一次性快照，题目后续修改不影响已组卷考试。user → question（隐式一对多），question表不存储creator_id，通过score表反向追溯出题教师。')

    add_body_text(doc, '（3）删除策略：删除用户时，若存在关联score记录则禁用(status=0)而非硬删除，保证成绩数据完整性；删除题目时直接硬删除，但会记录警告日志（题目可能被考试快照引用）；删除考试时仅允许draft状态删除，删除前回退所有题目的use统计；删除score记录无特殊限制。')

    # ============================================================
    # 添加页码（底部居中）
    # ============================================================
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = p.add_run()
        set_run_font(run, 'FangSong', 10.5)

        fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run._element.append(fldChar1)

        run2 = p.add_run()
        set_run_font(run2, 'FangSong', 10.5)
        instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run2._element.append(instrText)

        run3 = p.add_run()
        set_run_font(run3, 'FangSong', 10.5)
        fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run3._element.append(fldChar2)

    # ============================================================
    # 保存文档
    # ============================================================
    output_dir = '/workspace/temp/report'
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, '系统设计_黄泊凯_面向对象软件设计与建模.docx')
    doc.save(output_path)
    print(f'文档已生成: {output_path}')


if __name__ == '__main__':
    generate_report()
