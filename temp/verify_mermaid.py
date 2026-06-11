#!/usr/bin/env python3
"""深入检查文档中Mermaid内容的存储方式"""

import re
from docx import Document
from docx.oxml.ns import qn

DOC_PATH = "/workspace/temp/report/需求分析_黄泊凯_面向对象软件设计与建模.docx"
doc = Document(DOC_PATH)

print("=" * 70)
print("检查文档中所有包含Mermaid关键词的段落")
print("=" * 70)

mermaid_keywords = ["erDiagram", "classDiagram", "sequenceDiagram", "```mermaid", "```"]

for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if any(kw in text for kw in mermaid_keywords):
        style = p.style.name if p.style else "None"
        # 检查段落中的run格式
        run_info = []
        for run in p.runs:
            fn = run.font.name or ""
            rPr = run._element.find(qn('w:rPr'))
            ea = ""
            if rPr is not None:
                rFonts = rPr.find(qn('w:rFonts'))
                if rFonts is not None:
                    ea = rFonts.get(qn('w:eastAsia'), "")
            run_info.append(f"font={fn}/ea={ea}")
        print(f"\n段落#{i} [style={style}]:")
        print(f"  文本: {text[:200]}")
        print(f"  Runs: {run_info[:5]}")

print("\n\n" + "=" * 70)
print("检查文档中所有代码块/特殊格式段落")
print("=" * 70)

# 检查是否有嵌入的对象或图片中的Mermaid
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    style = p.style.name if p.style else "None"
    if "code" in style.lower() or "source" in style.lower() or "pre" in style.lower():
        print(f"\n代码段落#{i} [style={style}]:")
        print(f"  文本: {text[:200]}")

# 检查所有段落样式
print("\n\n" + "=" * 70)
print("所有段落样式统计")
print("=" * 70)
styles = {}
for p in doc.paragraphs:
    sname = p.style.name if p.style else "None"
    styles[sname] = styles.get(sname, 0) + 1
for s, c in sorted(styles.items(), key=lambda x: -x[1]):
    print(f"  {s}: {c}")

# 统计全文中erDiagram/classDiagram/sequenceDiagram出现次数
print("\n\n" + "=" * 70)
print("Mermaid关键词统计（全文段落）")
print("=" * 70)
all_text = "\n".join(p.text for p in doc.paragraphs)
for kw in ["erDiagram", "classDiagram", "sequenceDiagram"]:
    count = all_text.count(kw)
    print(f"  {kw}: {count}次")

# 检查表格中的Mermaid
print("\n\n" + "=" * 70)
print("Mermaid关键词统计（表格）")
print("=" * 70)
table_text = ""
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            table_text += cell.text + "\n"
for kw in ["erDiagram", "classDiagram", "sequenceDiagram"]:
    count = table_text.count(kw)
    print(f"  {kw}: {count}次")
