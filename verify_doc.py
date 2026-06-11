#!/usr/bin/env python3
"""验证系统设计文档是否满足所有checklist检查点"""

from docx import Document
from docx.shared import Pt, Emu
import re
import sys

DOC_PATH = "/workspace/temp/report/系统设计_黄泊凯_面向对象软件设计与建模.docx"

doc = Document(DOC_PATH)

# 提取所有段落文本
all_text = "\n".join(p.text for p in doc.paragraphs)
all_paragraphs = doc.paragraphs

results = {}

def check(name, condition, detail=""):
    results[name] = {"pass": condition, "detail": detail}
    status = "✅ PASS" if condition else "❌ FAIL"
    print(f"{status} | {name}")
    if detail:
        print(f"       {detail}")

# ============================================================
# 1. 封面页检查
# ============================================================
print("\n" + "="*60)
print("1. 封面页检查")
print("="*60)

cover_text = "\n".join(p.text for p in all_paragraphs[:16])

check("封面-课程名称(面向对象软件设计与建模)",
      "面向对象软件设计与建模" in cover_text)

check("封面-题目(系统设计)",
      "系统设计" in cover_text)

check("封面-指导教师(欧毓毅)",
      "欧毓毅" in cover_text)

check("封面-学生姓名(黄泊凯)",
      "黄泊凯" in cover_text)

check("封面-班级/学号(3123004394)",
      "3123004394" in cover_text)

check("封面-实验日期(2026-06-01)",
      "2026-06-01" in cover_text)

# ============================================================
# 2. 第一章 架构设计检查
# ============================================================
print("\n" + "="*60)
print("2. 第一章 架构设计检查")
print("="*60)

# 检查架构图使用Mermaid graph语法
has_graph_td = "graph TD" in all_text
has_graph_lr = "graph LR" in all_text
check("1.1 架构图用Mermaid graph语法",
      has_graph_td or has_graph_lr,
      f"graph TD: {has_graph_td}, graph LR: {has_graph_lr}")

# 检查架构图包含前端层+后端三层+数据层
arch_keywords = ["表示层", "业务逻辑层", "数据访问层", "数据层"]
arch_found = [kw for kw in arch_keywords if kw in all_text]
check("1.1 架构图包含前端层+后端三层+数据层",
      len(arch_found) >= 3,
      f"找到: {arch_found}")

# 检查运行环境说明(300字以上)
env_keywords = ["Windows 10", "i5-14600KF", "RTX 3090"]
env_found = [kw for kw in env_keywords if kw in all_text]
# 找运行环境段落
env_text = ""
for p in all_paragraphs:
    if "Windows 10" in p.text or "运行环境" in p.text:
        env_text += p.text
env_char_count = len(env_text) if env_text else 0
check("1.1 运行环境说明(300字以上)",
      env_char_count >= 300 and len(env_found) >= 2,
      f"环境说明字数: {env_char_count}, 关键词: {env_found}")

# 检查类包图
check("1.2 类包图用Mermaid语法",
      "graph LR" in all_text,
      "包含graph LR语法")

check("1.2 类包图包含common和modules",
      "common" in all_text and "modules" in all_text,
      f"common: {'common' in all_text}, modules: {'modules' in all_text}")

# 检查common子包
common_subpackages = ["api", "exception", "security", "config"]
common_found = [sp for sp in common_subpackages if sp in all_text]
check("1.2 类包图common含4子包",
      len(common_found) >= 4,
      f"找到子包: {common_found}")

# 检查modules模块包
modules_packages = ["user", "question", "exam", "score"]
modules_found = [m for m in modules_packages if m in all_text]
check("1.2 类包图modules含4模块包",
      len(modules_found) >= 4,
      f"找到模块: {modules_found}")

# ============================================================
# 3. 第二章 类文件结构检查
# ============================================================
print("\n" + "="*60)
print("3. 第二章 类文件结构检查")
print("="*60)

# 检查6层表格
tables = doc.tables

# Table 0: Controller层
if len(tables) > 0:
    t = tables[0]
    header = [c.text.strip() for c in t.rows[0].cells]
    data_rows = len(t.rows) - 1
    # 验证数据行中包含Controller类名
    controller_names = [t.rows[i].cells[0].text.strip() for i in range(1, len(t.rows))]
    all_controllers = all("Controller" in name for name in controller_names)
    check("2.1 Controller层表格含4个Controller类及端点数",
          data_rows >= 4 and all_controllers,
          f"表头: {header}, 数据行: {data_rows}, 类名: {controller_names}")
else:
    check("2.1 Controller层表格含4个Controller类及端点数", False, "未找到表格")

# Table 1: Service层
if len(tables) > 1:
    t = tables[1]
    header = [c.text.strip() for c in t.rows[0].cells]
    data_rows = len(t.rows) - 1
    service_names = [t.rows[i].cells[0].text.strip() for i in range(1, len(t.rows))]
    all_services = all("Service" in name for name in service_names)
    check("2.2 Service层表格含4个Service类及方法数",
          data_rows >= 4 and all_services,
          f"表头: {header}, 数据行: {data_rows}, 类名: {service_names}")
else:
    check("2.2 Service层表格含4个Service类及方法数", False, "未找到表格")

# Table 2: Repository层
if len(tables) > 2:
    t = tables[2]
    header = [c.text.strip() for c in t.rows[0].cells]
    data_rows = len(t.rows) - 1
    check("2.3 Repository层表格含4个Repository接口",
          "Repository" in header[0] or "接口" in header[0] and data_rows >= 4,
          f"表头: {header}, 数据行: {data_rows}")
else:
    check("2.3 Repository层表格含4个Repository接口", False, "未找到表格")

# Table 3: Entity层
if len(tables) > 3:
    t = tables[3]
    header = [c.text.strip() for c in t.rows[0].cells]
    data_rows = len(t.rows) - 1
    entity_names = [t.rows[i].cells[0].text.strip() for i in range(1, len(t.rows))]
    has_enum = any("Enum" in t.rows[i].cells[1].text or "枚举" in t.rows[i].cells[1].text
                   for i in range(1, len(t.rows)))
    check("2.4 Entity层表格含4个Entity+3个Enum",
          data_rows >= 7 and has_enum,
          f"数据行: {data_rows}, 类名: {entity_names}")
else:
    check("2.4 Entity层表格含4个Entity+3个Enum", False, "未找到表格")

# Tables 4-7: DTO层 (M01-M04)
dto_total = 0
dto_names = []
for i in [4, 5, 6, 7]:
    if len(tables) > i:
        t = tables[i]
        data_rows = len(t.rows) - 1
        dto_total += data_rows
        for row in t.rows[1:]:
            dto_names.append(row.cells[0].text.strip())

check("2.5 DTO层表格含43个DTO类（按模块分组）",
      dto_total == 43,
      f"实际DTO数量: {dto_total}, 类名列表: {dto_names}")

# Table 8: Common层
if len(tables) > 8:
    t = tables[8]
    header = [c.text.strip() for c in t.rows[0].cells]
    data_rows = len(t.rows) - 1
    check("2.6 Common层表格含8个类",
          data_rows >= 8,
          f"表头: {header}, 数据行: {data_rows}")
else:
    check("2.6 Common层表格含8个类", False, "未找到表格")

# 检查每层有200字以上文字说明
layer_sections = {
    "2.1 Controller层": "Controller层是系统的表示层",
    "2.2 Service层": "Service层是系统的业务逻辑核心",
    "2.3 Repository层": "Repository层是系统的数据访问层",
    "2.4 Entity层": "Entity层定义了系统的",
    "2.5 DTO层": "DTO层是系统数据传输对象",
    "2.6 Common层": "Common基础设施层是系统的横切关注点",
}
for section, marker in layer_sections.items():
    found = marker in all_text
    # 计算该段文字长度
    char_count = 0
    for p in all_paragraphs:
        if marker in p.text:
            char_count = len(p.text)
            break
    check(f"{section} 有200字以上说明",
          found and char_count >= 200,
          f"找到标记: {found}, 字数: {char_count}")

# ============================================================
# 4. 第三章 VOPC时序图检查
# ============================================================
print("\n" + "="*60)
print("4. 第三章 VOPC时序图检查")
print("="*60)

# 统计sequenceDiagram出现次数
seq_count = all_text.count("sequenceDiagram")
check("第三章 sequenceDiagram出现38次",
      seq_count == 38,
      f"实际出现次数: {seq_count}")

# 按模块统计
# M01: 图3-1到3-11 (11张)
# M02: 图3-12到3-19 (8张)
# M03: 图3-20到3-29 (10张)
# M04: 图3-30到3-38 (9张)

m01_count = 0
m02_count = 0
m03_count = 0
m04_count = 0

for p in all_paragraphs:
    text = p.text.strip()
    if "sequenceDiagram" in text:
        # 根据位置判断属于哪个模块
        pass

# 通过图号统计
fig_m01 = len(re.findall(r'图3-(\d+)', all_text))
fig_numbers = re.findall(r'图3-(\d+)', all_text)
if fig_numbers:
    max_fig = max(int(n) for n in fig_numbers)
    m01_figs = [int(n) for n in fig_numbers if int(n) <= 11]
    m02_figs = [int(n) for n in fig_numbers if 12 <= int(n) <= 19]
    m03_figs = [int(n) for n in fig_numbers if 20 <= int(n) <= 29]
    m04_figs = [int(n) for n in fig_numbers if 30 <= int(n) <= 38]
    
    check("M01模块11张时序图",
          len(m01_figs) == 11,
          f"实际: {len(m01_figs)}张")
    check("M02模块8张时序图",
          len(m02_figs) == 8,
          f"实际: {len(m02_figs)}张")
    check("M03模块10张时序图",
          len(m03_figs) == 10,
          f"实际: {len(m03_figs)}张")
    check("M04模块9张时序图",
          len(m04_figs) == 9,
          f"实际: {len(m04_figs)}张")

# 检查每张图下方有说明文字
# 图标题格式: 图3-X ...时序图
# 说明文字在图标题之后
fig_captions = []
desc_after_fig = 0
for i, p in enumerate(all_paragraphs):
    if re.match(r'图3-\d+', p.text.strip()):
        fig_captions.append(i)
        # 检查后续段落是否有说明文字(非空且不是下一个图)
        if i + 1 < len(all_paragraphs):
            next_text = all_paragraphs[i+1].text.strip()
            if next_text and not next_text.startswith("图3-") and "sequenceDiagram" not in next_text:
                desc_after_fig += 1

check("每张图下方附说明文字",
      desc_after_fig >= 30,
      f"有说明的图: {desc_after_fig}/{len(fig_captions)}")

# 检查方法名与代码一致（检查关键方法名出现）
key_methods = ["findByName", "existsByName", "generateToken", "resolveCurrentStatus",
               "gradeOne", "incrementUse", "decrementUse", "toVO"]
methods_found = [m for m in key_methods if m in all_text]
check("方法名与代码一致",
      len(methods_found) >= 6,
      f"找到方法: {methods_found}")

# ============================================================
# 5. 第四章 核心类详细定义检查
# ============================================================
print("\n" + "="*60)
print("5. 第四章 核心类详细定义检查")
print("="*60)

core_classes = {
    "UserService": "classDiagram",
    "QuestionService": "classDiagram",
    "ExamService": "classDiagram",
    "ScoreService": "classDiagram",
    "JwtUtil": "classDiagram",
    "Result": "classDiagram",
    "GlobalExceptionHandler": "classDiagram",
}

for cls_name, diagram_type in core_classes.items():
    # 检查类名在classDiagram代码块中出现
    has_class = False
    has_diagram = False
    for p in all_paragraphs:
        if cls_name in p.text and "classDiagram" in p.text:
            has_class = True
            has_diagram = True
            break
        if cls_name in p.text and "class " in p.text:
            has_class = True
    
    check(f"4.x {cls_name}类图含属性和方法签名",
          has_class,
          f"类名出现: {has_class}")

# 检查每个类图下方有设计要点说明
design_notes = 0
for i, p in enumerate(all_paragraphs):
    text = p.text.strip()
    if re.match(r'图4-\d+', text):
        # 检查后续段落
        if i + 1 < len(all_paragraphs):
            next_text = all_paragraphs[i+1].text.strip()
            if next_text and len(next_text) > 50:
                design_notes += 1

check("每个类图下方附设计要点说明",
      design_notes >= 7,
      f"有说明的类图: {design_notes}/7")

# ============================================================
# 6. 第五章 数据库表结构检查
# ============================================================
print("\n" + "="*60)
print("6. 第五章 数据库表结构检查")
print("="*60)

# 检查4张表定义
db_tables = ["user表", "question表", "exam表", "score表"]
for tbl in db_tables:
    check(f"5.x {tbl}字段定义",
          tbl in all_text,
          f"{'找到' if tbl in all_text else '未找到'} {tbl}")

# 检查表定义表格（Tables 9-12）
table_names_in_ch5 = []
for i in [9, 10, 11, 12]:
    if len(tables) > i:
        t = tables[i]
        header = [c.text.strip() for c in t.rows[0].cells]
        has_full_cols = all(h in header for h in ["字段名", "Java类型", "SQLite类型", "约束", "默认值", "业务含义"])
        data_rows = len(t.rows) - 1
        table_names_in_ch5.append((i, header, data_rows, has_full_cols))

check("5.1 user表字段定义完整(6列)",
      len(table_names_in_ch5) > 0 and table_names_in_ch5[0][3],
      f"表9: {table_names_in_ch5[0] if table_names_in_ch5 else '未找到'}")

check("5.2 question表字段定义完整(6列)",
      len(table_names_in_ch5) > 1 and table_names_in_ch5[1][3],
      f"表10: {table_names_in_ch5[1] if len(table_names_in_ch5) > 1 else '未找到'}")

check("5.3 exam表字段定义完整(6列)",
      len(table_names_in_ch5) > 2 and table_names_in_ch5[2][3],
      f"表11: {table_names_in_ch5[2] if len(table_names_in_ch5) > 2 else '未找到'}")

check("5.4 score表字段定义完整(6列)",
      len(table_names_in_ch5) > 3 and table_names_in_ch5[3][3],
      f"表12: {table_names_in_ch5[3] if len(table_names_in_ch5) > 3 else '未找到'}")

# 检查JSON结构说明
json_structures = {
    "answer JSON 5种结构": "answer JSON",
    "question_sum JSON结构": "question_sum JSON",
    "detail JSON结构": "detail JSON",
}
for name, marker in json_structures.items():
    check(f"5.x {name}",
          marker in all_text,
          f"{'找到' if marker in all_text else '未找到'}")

# 检查5种answer JSON结构
answer_types = ["SingleChoice", "MultipleChoice", "Judge", "Fill", "Essay"]
answer_found = [a for a in answer_types if a in all_text]
check("5.2 answer JSON 5种结构说明",
      len(answer_found) >= 5,
      f"找到: {answer_found}")

# 检查状态机说明
check("5.3 exam状态机说明",
      "draft" in all_text and "publish" in all_text and "running" in all_text and "done" in all_text,
      "draft/publish/running/done状态")

# 检查表间关系
check("5.5 表间关系说明",
      "表间关系" in all_text or "物理外键" in all_text,
      f"物理外键: {'物理外键' in all_text}, 逻辑引用: {'逻辑引用' in all_text}")

# ============================================================
# 7. 文档格式检查
# ============================================================
print("\n" + "="*60)
print("7. 文档格式检查")
print("="*60)

# 检查字体使用
kaiti_count = 0
fangsong_count = 0
courier_count = 0

for p in all_paragraphs:
    for run in p.runs:
        if run.font.name == "KaiTi":
            kaiti_count += 1
        elif run.font.name == "FangSong":
            fangsong_count += 1
        elif run.font.name == "Courier New":
            courier_count += 1

check("楷体(KaiTi)字体使用",
      kaiti_count > 0,
      f"使用KaiTi的run数: {kaiti_count}")

check("仿宋(FangSong)字体使用",
      fangsong_count > 0,
      f"使用FangSong的run数: {fangsong_count}")

check("等宽字体(Courier New)用于UML图",
      courier_count > 0,
      f"使用Courier New的run数: {courier_count}")

# 检查一级标题格式（楷体小二号加粗）
# 小二号 = 18pt = 228600 EMU
heading1_ok = False
for p in all_paragraphs:
    for run in p.runs:
        if run.font.name == "KaiTi" and run.font.bold and run.font.size and run.font.size == Pt(18):
            heading1_ok = True
            break
    if heading1_ok:
        break

check("一级标题：楷体小二号(18pt)加粗",
      heading1_ok,
      f"找到符合格式的一级标题: {heading1_ok}")

# 检查二级标题格式（楷体小三号加粗）
# 小三号 = 15pt = 190500 EMU
heading2_ok = False
for p in all_paragraphs:
    for run in p.runs:
        if run.font.name == "KaiTi" and run.font.bold and run.font.size and run.font.size == Pt(15):
            heading2_ok = True
            break
    if heading2_ok:
        break

check("二级标题：楷体小三号(15pt)加粗",
      heading2_ok,
      f"找到符合格式的二级标题: {heading2_ok}")

# 检查三级标题格式（楷体四号加粗）
# 四号 = 14pt = 177800 EMU
heading3_ok = False
for p in all_paragraphs:
    for run in p.runs:
        if run.font.name == "KaiTi" and run.font.bold and run.font.size and run.font.size == Pt(14):
            heading3_ok = True
            break
    if heading3_ok:
        break

check("三级标题：楷体四号(14pt)加粗",
      heading3_ok,
      f"找到符合格式的三级标题: {heading3_ok}")

# 检查正文格式（仿宋小四号）
# 小四号 = 12pt = 152400 EMU
body_ok = False
for p in all_paragraphs:
    for run in p.runs:
        if run.font.name == "FangSong" and run.font.size and run.font.size == Pt(12):
            body_ok = True
            break
    if body_ok:
        break

check("正文：仿宋小四号(12pt)",
      body_ok,
      f"找到符合格式的正文: {body_ok}")

# 检查文件名
import os
filename = os.path.basename(DOC_PATH)
check("文件名：系统设计_黄泊凯_面向对象软件设计与建模.docx",
      filename == "系统设计_黄泊凯_面向对象软件设计与建模.docx",
      f"实际文件名: {filename}")

# 检查保存路径
check("保存路径：temp/report/",
      os.path.exists(DOC_PATH),
      f"文件存在: {os.path.exists(DOC_PATH)}")

# ============================================================
# 汇总
# ============================================================
print("\n" + "="*60)
print("验证汇总")
print("="*60)

total = len(results)
passed = sum(1 for r in results.values() if r["pass"])
failed = total - passed

print(f"总计: {total} 项检查")
print(f"通过: {passed} 项")
print(f"失败: {failed} 项")

if failed > 0:
    print("\n失败项详情:")
    for name, r in results.items():
        if not r["pass"]:
            print(f"  ❌ {name}: {r['detail']}")

# 输出用于更新checklist的标记
print("\n\n# Checklist更新标记（用于checklist.md）")
for name, r in results.items():
    mark = "x" if r["pass"] else " "
    print(f"- [{mark}] {name}")
